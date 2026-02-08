from __future__ import annotations

import json
import os
import re
import sqlite3
import uuid
import contextvars
import hashlib
import shutil
from pathlib import Path
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Any, Literal

import httpx
from pypdf import PdfReader
from docx import Document
from fastapi import FastAPI, Form, HTTPException, Request, UploadFile, File
from fastapi.responses import HTMLResponse, RedirectResponse, JSONResponse, FileResponse
from starlette.middleware.base import BaseHTTPMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

Status = Literal["draft", "submitted", "returned", "confirmed", "deprecated"]
Role = Literal["viewer", "author", "assessment_author", "content_publisher", "reviewer", "audit", "admin"]

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")
DB_DEMO_PATH = os.path.join(DATA_DIR, "lcs_demo.db")
DB_BLANK_PATH = os.path.join(DATA_DIR, "lcs_blank.db")
UPLOADS_DIR = os.path.join(DATA_DIR, "uploads")
EXPORTS_DIR = os.path.join(DATA_DIR, "exports")

# Per-request DB selection (MVP): use a cookie. Default to demo.
DB_KEY_COOKIE = "lcs_db"
DB_KEY_DEMO = "demo"
DB_KEY_BLANK = "blank"
DB_PATH_CTX: contextvars.ContextVar[str] = contextvars.ContextVar("lcs_db_path", default=DB_DEMO_PATH)
DB_KEY_CTX: contextvars.ContextVar[str] = contextvars.ContextVar("lcs_db_key", default=DB_KEY_DEMO)

DB_PROFILE_KEY_RE = re.compile(r"^[a-z][a-z0-9_-]{0,63}$")

LMSTUDIO_BASE_URL = os.environ.get("LCS_LMSTUDIO_BASE_URL", "http://127.0.0.1:1234").rstrip("/")
LMSTUDIO_MODEL = os.environ.get("LCS_LMSTUDIO_MODEL", "mistralai/mistral-7b-instruct-v0.3")

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")
templates = Jinja2Templates(directory=TEMPLATES_DIR)

app = FastAPI(title="Learning Content System MVP")


@app.exception_handler(HTTPException)
async def _http_exception_handler(request: Request, exc: HTTPException):
    """Prefer HTML error details for browser flows.

    FastAPI's default is JSON, which often shows up as a generic "Internal Server Error" page
    in a normal browser POST flow.
    """
    accept = (request.headers.get("accept") or "").lower()
    if "text/html" in accept and (str(request.url.path).startswith("/import/pdf") or str(request.url.path).startswith("/import/json")):
        # Render the import form again, but include the error detail.
        template = "import_json.html" if str(request.url.path).startswith("/import/json") else "import_pdf.html"
        ctx = {"error": str(exc.detail)}
        if template == "import_pdf.html":
            ctx.update({"lmstudio_base_url": LMSTUDIO_BASE_URL, "lmstudio_model": LMSTUDIO_MODEL})
        return templates.TemplateResponse(
            request,
            template,
            ctx,
            status_code=exc.status_code,
        )
    return JSONResponse(status_code=exc.status_code, content={"detail": exc.detail})


@app.exception_handler(Exception)
async def _unhandled_exception_handler(request: Request, exc: Exception):
    """Catch-all so browser users get *some* error text instead of a blank 500 page."""
    import traceback

    # Always log the traceback to stdout so uvicorn logs capture it.
    traceback.print_exc()

    accept = (request.headers.get("accept") or "").lower()
    if "text/html" in accept and (str(request.url.path).startswith("/import/pdf") or str(request.url.path).startswith("/import/json")):
        template = "import_json.html" if str(request.url.path).startswith("/import/json") else "import_pdf.html"
        ctx = {"error": f"Unhandled error: {type(exc).__name__}: {exc}"}
        if template == "import_pdf.html":
            ctx.update({"lmstudio_base_url": LMSTUDIO_BASE_URL, "lmstudio_model": LMSTUDIO_MODEL})
        return templates.TemplateResponse(
            request,
            template,
            ctx,
            status_code=500,
        )

    return JSONResponse(
        status_code=500,
        content={"detail": f"Unhandled error: {type(exc).__name__}: {exc}"},
    )


static_dir = os.path.join(os.path.dirname(__file__), "static")
app.mount("/static", StaticFiles(directory=static_dir), name="static")


# --- Auth + RBAC (MVP) ---

# Demo-friendly local auth:
# - Users are stored in the selected SQLite DB.
# - Login issues a server-side session token (cookie).
# - Roles are not client-controlled.
DEFAULT_ROLE: Role = "viewer"
SESSION_COOKIE = "lcs_session"

# --- Lightweight enterprise fields (MVP+) ---
# Stored as JSON blobs to avoid migrations churn in the prototype.
DEFAULT_TAGS: list[str] = []
DEFAULT_META: dict[str, str] = {}

ROLE_ORDER: dict[Role, int] = {
    "viewer": 0,
    "author": 1,
    "assessment_author": 2,
    "content_publisher": 3,
    "reviewer": 4,
    "audit": 5,
    "admin": 6,
}


def _is_public_path(path: str) -> bool:
    return path.startswith("/static/") or path in ("/login", "/logout")


def can(role: Role, action: str) -> bool:
    """Very small RBAC matrix.

    Design rule (important):
    - Domain entitlements are used to gate *authoring* and *review/confirm* operations.
    - Domain entitlements must NOT gate read-only viewing/browsing. Any authenticated user
      may view records across domains.
    - Delivery/publishing is confirmed-only. Export authorization is role-based, not domain-based.

    Actions:
      - task:create, task:revise, task:submit, task:confirm
      - workflow:create, workflow:revise, workflow:submit, workflow:confirm
      - assessment:create, assessment:revise, assessment:submit, assessment:confirm
      - delivery:view, delivery:export
      - import:pdf
      - import:json
      - db:switch
      - audit:view
      - task:force_submit, task:force_confirm
      - workflow:force_submit, workflow:force_confirm
    """
    if role == "admin":
        return True

    if action == "audit:view":
        return role in ("audit",)

    if action == "delivery:view":
        return role in ("viewer", "author", "assessment_author", "content_publisher", "reviewer")

    if action == "delivery:export":
        return role in ("content_publisher",)

    if action.endswith(":force_confirm") or action.endswith(":force_submit"):
        return role in ("admin",)

    if action.endswith(":confirm"):
        # Review firewall: reviewers can review/confirm *everything*.
        # They do not revise content/assessments; they only confirm or return.
        return role in ("reviewer",)

    if action.startswith("assessment:"):
        # Content/assessment firewall:
        # - assessment authors create/revise/submit assessments
        # - reviewers can confirm/return (handled by :confirm)
        # - content authors do not author assessments
        if action.endswith(":submit"):
            return role in ("assessment_author",)
        if action.endswith(":create") or action.endswith(":revise"):
            return role in ("assessment_author",)

    if action.endswith(":submit"):
        return role in ("author",)

    if action.endswith(":create"):
        return role in ("author",)

    if action.endswith(":revise"):
        # Review firewall: reviewers do not revise.
        return role in ("author",)

    if action in ("import:pdf", "import:json"):
        # Keep ingestion with content authoring, not assessment.
        return role in ("author",)

    if action == "db:switch":
        return role in ("admin",)

    return False


def require(role: Role, action: str) -> None:
    if not can(role, action):
        raise HTTPException(status_code=403, detail=f"Forbidden: requires permission {action}")


def require_admin(request: Request) -> None:
    if request.state.role != "admin":
        raise HTTPException(status_code=403, detail="Forbidden: admin only")


def _user_id(conn: sqlite3.Connection, username: str) -> int | None:
    row = conn.execute("SELECT id FROM users WHERE username=? AND disabled_at IS NULL", (username,)).fetchone()
    return int(row["id"]) if row else None


def _user_has_domain(conn: sqlite3.Connection, username: str, domain: str) -> bool:
    if not domain:
        return False

    # Admin is a break-glass role: treat as authorized for all domains.
    r = conn.execute("SELECT role FROM users WHERE username=? AND disabled_at IS NULL", (username,)).fetchone()
    if r and str(r["role"]) == "admin":
        return True

    uid = _user_id(conn, username)
    if uid is None:
        return False
    row = conn.execute("SELECT 1 FROM user_domains WHERE user_id=? AND domain=?", (uid, domain)).fetchone()
    return bool(row)


def _user_domains(conn: sqlite3.Connection, username: str) -> list[str]:
    """Return active domain entitlements for the user (sorted).

    Admin is treated as implicitly authorized for all domains.
    """
    # Admin break-glass
    r = conn.execute("SELECT role FROM users WHERE username=? AND disabled_at IS NULL", (username,)).fetchone()
    if r and str(r["role"]) == "admin":
        return _active_domains(conn)

    uid = _user_id(conn, username)
    if uid is None:
        return []
    rows = conn.execute(
        "SELECT domain FROM user_domains WHERE user_id=? ORDER BY domain ASC",
        (uid,),
    ).fetchall()
    return [str(x["domain"]) for x in rows if x and x["domain"]]


def _active_domains(conn: sqlite3.Connection) -> list[str]:
    rows = conn.execute("SELECT name FROM domains WHERE disabled_at IS NULL ORDER BY name ASC").fetchall()
    return [str(r["name"]) for r in rows]


def _workflow_domains(conn: sqlite3.Connection, refs: list[tuple[str, int]]) -> list[str]:
    if not refs:
        return []
    domains: set[str] = set()
    for rid, ver in refs:
        r = conn.execute("SELECT domain FROM tasks WHERE record_id=? AND version=?", (rid, ver)).fetchone()
        d = (str(r["domain"]) if r else "").strip()
        if d:
            domains.add(d)
    return sorted(domains)


def _backfill_workflow_domains(conn: sqlite3.Connection) -> None:
    """Best-effort migration: populate workflows.domains_json for existing rows."""
    if not _column_exists(conn, "workflows", "domains_json"):
        return

    rows = conn.execute("SELECT record_id, version, domains_json FROM workflows").fetchall()
    for r in rows:
        try:
            existing = (r["domains_json"] or "").strip()
        except Exception:
            existing = ""
        if existing and existing != "[]":
            continue

        refs = conn.execute(
            "SELECT task_record_id, task_version FROM workflow_task_refs WHERE workflow_record_id=? AND workflow_version=? ORDER BY order_index",
            (r["record_id"], int(r["version"])),
        ).fetchall()
        pairs = [(x["task_record_id"], int(x["task_version"])) for x in refs]
        doms = _workflow_domains(conn, pairs)
        conn.execute(
            "UPDATE workflows SET domains_json=? WHERE record_id=? AND version=?",
            (_json_dump(doms), r["record_id"], int(r["version"])),
        )


def _hash_password(password: str, salt_hex: str) -> str:
    import hashlib

    pw = (password or "").encode("utf-8")
    salt = bytes.fromhex(salt_hex)
    dk = hashlib.pbkdf2_hmac("sha256", pw, salt, 200_000)
    return dk.hex()


def _verify_password(password: str, salt_hex: str, hash_hex: str) -> bool:
    import hmac

    return hmac.compare_digest(_hash_password(password, salt_hex), hash_hex)


def _new_session_token() -> str:
    import secrets

    return secrets.token_urlsafe(32)


def _require_login(request: Request) -> bool:
    # Login page should be reachable without a session.
    return not _is_public_path(str(request.url.path))


class AuthMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        # DB selection (cookie). Default to demo DB.
        key = _selected_db_key(request)
        request.state.db_key = key
        request.state.db_path = _db_path_for_key(key)
        DB_KEY_CTX.set(key)
        DB_PATH_CTX.set(request.state.db_path)

        # Default unauth state (used by /login rendering).
        request.state.user = ""
        request.state.role = DEFAULT_ROLE

        if _require_login(request):
            token = (request.cookies.get(SESSION_COOKIE) or "").strip()
            if token:
                with db() as conn:
                    row = conn.execute(
                        """
                        SELECT u.username, u.role
                        FROM sessions s
                        JOIN users u ON u.id = s.user_id
                        WHERE s.token=? AND s.revoked_at IS NULL
                        """,
                        (token,),
                    ).fetchone()
                if row:
                    request.state.user = str(row["username"])
                    request.state.role = str(row["role"])  # type: ignore[assignment]

            if not request.state.user:
                accept = (request.headers.get("accept") or "").lower()
                if "text/html" in accept:
                    return RedirectResponse(url="/login", status_code=303)
                # Don't raise inside middleware (can produce noisy exception groups).
                return JSONResponse(status_code=401, content={"detail": "Unauthorized"})

        return await call_next(request)


app.add_middleware(AuthMiddleware)

# make permission checks available in templates
templates.env.globals["can"] = can


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def _selected_db_key(request: Request | None = None) -> str:
    if request is None:
        return DB_KEY_CTX.get()
    k = (request.cookies.get(DB_KEY_COOKIE) or DB_KEY_DEMO).strip().lower()
    if k not in _available_db_keys():
        return DB_KEY_DEMO
    return k


def _db_path_for_key(key: str) -> str:
    if key == DB_KEY_DEMO:
        return DB_DEMO_PATH
    if key == DB_KEY_BLANK:
        return DB_BLANK_PATH
    return os.path.join(DATA_DIR, f"lcs_{key}.db")


def _available_db_keys() -> set[str]:
    return {DB_KEY_DEMO, DB_KEY_BLANK, *_list_custom_db_keys()}


def _list_custom_db_keys() -> list[str]:
    """Return custom db profile keys from files in DATA_DIR.

    Custom DB files are named: lcs_<key>.db
    Reserved keys demo/blank are excluded.
    """
    os.makedirs(DATA_DIR, exist_ok=True)
    keys: list[str] = []
    for name in os.listdir(DATA_DIR):
        if not name.startswith("lcs_") or not name.endswith(".db"):
            continue
        if name in (os.path.basename(DB_DEMO_PATH), os.path.basename(DB_BLANK_PATH)):
            continue
        key = name[len("lcs_") : -len(".db")].strip().lower()
        if not key:
            continue
        if key in (DB_KEY_DEMO, DB_KEY_BLANK):
            continue
        if not DB_PROFILE_KEY_RE.match(key):
            # ignore weird files
            continue
        keys.append(key)
    keys.sort()
    return keys


def _create_custom_db_profile(key: str) -> None:
    key = (key or "").strip().lower()
    if key in (DB_KEY_DEMO, DB_KEY_BLANK):
        raise HTTPException(status_code=400, detail="Reserved profile key")
    if not DB_PROFILE_KEY_RE.match(key):
        raise HTTPException(status_code=400, detail="Invalid profile key (use: a-z, 0-9, _, -)")

    os.makedirs(DATA_DIR, exist_ok=True)
    dst = _db_path_for_key(key)
    if os.path.exists(dst):
        raise HTTPException(status_code=409, detail="Profile already exists")

    # Copy the blank DB as a template (schema + seeded users/domains; no content).
    # This keeps the MVP flow simple (switching DBs will still allow logging in).
    if not os.path.exists(DB_BLANK_PATH):
        init_db_path(DB_BLANK_PATH)
        DB_PATH_CTX.set(DB_BLANK_PATH)
        with db() as conn:
            _seed_demo_users(conn)
            _seed_demo_domains(conn)
            _seed_demo_entitlements(conn)

    shutil.copyfile(DB_BLANK_PATH, dst)


def db() -> sqlite3.Connection:
    """Open a connection to the currently selected DB (via context var).

    Notes:
    - FastAPI sync routes run in a threadpool; we open a fresh connection per request.
    - SQLite is single-writer; WAL + a busy timeout avoids spurious "database is locked"
      errors under light concurrency.
    """
    os.makedirs(DATA_DIR, exist_ok=True)
    path = DB_PATH_CTX.get()

    # timeout: how long sqlite3 waits on a locked DB before raising OperationalError.
    conn = sqlite3.connect(path, timeout=10.0)
    conn.row_factory = sqlite3.Row

    # Pragmas are per-connection.
    conn.execute("PRAGMA foreign_keys = ON")
    conn.execute("PRAGMA journal_mode = WAL")
    conn.execute("PRAGMA busy_timeout = 10000")
    return conn


def _column_exists(conn: sqlite3.Connection, table: str, column: str) -> bool:
    rows = conn.execute(f"PRAGMA table_info({table})").fetchall()
    return any(r["name"] == column for r in rows)


def init_db_path(db_path: str) -> None:
    os.makedirs(DATA_DIR, exist_ok=True)
    os.makedirs(UPLOADS_DIR, exist_ok=True)
    os.makedirs(EXPORTS_DIR, exist_ok=True)

    # Temporarily point context to this path for schema/migrations.
    DB_PATH_CTX.set(db_path)

    with db() as conn:
        conn.executescript(
            """
            CREATE TABLE IF NOT EXISTS tasks (
              record_id TEXT NOT NULL,
              version INTEGER NOT NULL,
              status TEXT NOT NULL,

              title TEXT NOT NULL,
              outcome TEXT NOT NULL,
              facts_json TEXT NOT NULL,
              concepts_json TEXT NOT NULL,
              procedure_name TEXT NOT NULL,
              steps_json TEXT NOT NULL,
              dependencies_json TEXT NOT NULL,
              irreversible_flag INTEGER NOT NULL,
              task_assets_json TEXT NOT NULL,

              domain TEXT NOT NULL DEFAULT '',

              tags_json TEXT NOT NULL DEFAULT '[]',
              meta_json TEXT NOT NULL DEFAULT '{}',

              created_at TEXT NOT NULL,
              updated_at TEXT NOT NULL,
              created_by TEXT NOT NULL,
              updated_by TEXT NOT NULL,
              reviewed_at TEXT,
              reviewed_by TEXT,
              change_note TEXT,

              needs_review_flag INTEGER NOT NULL DEFAULT 0,
              needs_review_note TEXT,

              PRIMARY KEY (record_id, version)
            );

            CREATE TABLE IF NOT EXISTS workflows (
              record_id TEXT NOT NULL,
              version INTEGER NOT NULL,
              status TEXT NOT NULL,

              title TEXT NOT NULL,
              objective TEXT NOT NULL,

              domains_json TEXT NOT NULL DEFAULT '[]',
              tags_json TEXT NOT NULL DEFAULT '[]',
              meta_json TEXT NOT NULL DEFAULT '{}',

              created_at TEXT NOT NULL,
              updated_at TEXT NOT NULL,
              created_by TEXT NOT NULL,
              updated_by TEXT NOT NULL,
              reviewed_at TEXT,
              reviewed_by TEXT,
              change_note TEXT,

              needs_review_flag INTEGER NOT NULL DEFAULT 0,
              needs_review_note TEXT,

              PRIMARY KEY (record_id, version)
            );

            CREATE TABLE IF NOT EXISTS workflow_task_refs (
              workflow_record_id TEXT NOT NULL,
              workflow_version INTEGER NOT NULL,
              order_index INTEGER NOT NULL,
              task_record_id TEXT NOT NULL,
              task_version INTEGER NOT NULL,

              PRIMARY KEY (workflow_record_id, workflow_version, order_index),
              FOREIGN KEY (workflow_record_id, workflow_version)
                REFERENCES workflows(record_id, version)
                ON DELETE CASCADE,
              FOREIGN KEY (task_record_id, task_version)
                REFERENCES tasks(record_id, version)
                ON DELETE RESTRICT
            );

            -- ---- Assessments (MVP) ----

            CREATE TABLE IF NOT EXISTS assessment_items (
              record_id TEXT NOT NULL,
              version INTEGER NOT NULL,
              status TEXT NOT NULL,

              stem TEXT NOT NULL,
              options_json TEXT NOT NULL,
              correct_key TEXT NOT NULL,
              rationale TEXT NOT NULL DEFAULT '',

              claim TEXT NOT NULL DEFAULT 'fact_probe',
              domains_json TEXT NOT NULL DEFAULT '[]',
              lint_json TEXT NOT NULL DEFAULT '[]',
              refs_json TEXT NOT NULL DEFAULT '[]',

              tags_json TEXT NOT NULL DEFAULT '[]',
              meta_json TEXT NOT NULL DEFAULT '{}',

              created_at TEXT NOT NULL,
              updated_at TEXT NOT NULL,
              created_by TEXT NOT NULL,
              updated_by TEXT NOT NULL,
              reviewed_at TEXT,
              reviewed_by TEXT,
              change_note TEXT,

              needs_review_flag INTEGER NOT NULL DEFAULT 0,
              needs_review_note TEXT,

              PRIMARY KEY (record_id, version)
            );

            CREATE TABLE IF NOT EXISTS assessment_refs (
              assessment_record_id TEXT NOT NULL,
              assessment_version INTEGER NOT NULL,
              order_index INTEGER NOT NULL,
              ref_type TEXT NOT NULL,
              ref_record_id TEXT NOT NULL,
              ref_version INTEGER NOT NULL,

              PRIMARY KEY (assessment_record_id, assessment_version, order_index),
              FOREIGN KEY (assessment_record_id, assessment_version)
                REFERENCES assessment_items(record_id, version)
                ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS audit_log (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              entity_type TEXT NOT NULL,
              record_id TEXT NOT NULL,
              version INTEGER NOT NULL,
              action TEXT NOT NULL,
              actor TEXT NOT NULL,
              at TEXT NOT NULL,
              note TEXT
            );

            -- ---- Export artifacts (v0) ----
            CREATE TABLE IF NOT EXISTS export_artifacts (
              id TEXT PRIMARY KEY,
              kind TEXT NOT NULL,
              filename TEXT NOT NULL,
              path TEXT NOT NULL,
              sha256 TEXT NOT NULL,

              workflow_record_id TEXT NOT NULL,
              workflow_version INTEGER NOT NULL,
              task_refs_json TEXT NOT NULL,

              exported_at TEXT NOT NULL,
              exported_by TEXT NOT NULL,
              retention_days INTEGER NOT NULL DEFAULT 30
            );

            CREATE INDEX IF NOT EXISTS idx_export_artifacts_workflow ON export_artifacts(workflow_record_id, workflow_version);

            -- Local auth (demo-friendly; real auth may be added later)
            CREATE TABLE IF NOT EXISTS users (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              username TEXT NOT NULL UNIQUE,
              role TEXT NOT NULL,
              password_salt_hex TEXT NOT NULL,
              password_hash_hex TEXT NOT NULL,
              demo_password TEXT,
              created_at TEXT NOT NULL,
              created_by TEXT NOT NULL,
              disabled_at TEXT
            );

            CREATE TABLE IF NOT EXISTS sessions (
              token TEXT PRIMARY KEY,
              user_id INTEGER NOT NULL,
              created_at TEXT NOT NULL,
              revoked_at TEXT,
              FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
            );

            -- Domains (admin-managed registry) + per-user domain entitlements
            CREATE TABLE IF NOT EXISTS domains (
              name TEXT PRIMARY KEY,
              created_at TEXT NOT NULL,
              created_by TEXT NOT NULL,
              disabled_at TEXT
            );

            CREATE TABLE IF NOT EXISTS user_domains (
              user_id INTEGER NOT NULL,
              domain TEXT NOT NULL,
              created_at TEXT NOT NULL,
              created_by TEXT NOT NULL,
              PRIMARY KEY (user_id, domain),
              FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
              FOREIGN KEY (domain) REFERENCES domains(name) ON DELETE RESTRICT
            );

            CREATE TABLE IF NOT EXISTS ingestions (
              id TEXT PRIMARY KEY,
              source_type TEXT NOT NULL,
              source_sha256 TEXT NOT NULL,
              filename TEXT NOT NULL,
              created_by TEXT NOT NULL,
              created_at TEXT NOT NULL,
              status TEXT NOT NULL DEFAULT 'draft',
              cursor_chunk INTEGER NOT NULL DEFAULT 0,
              max_tasks_per_run INTEGER NOT NULL DEFAULT 10,
              note TEXT
            );

            CREATE TABLE IF NOT EXISTS ingestion_chunks (
              ingestion_id TEXT NOT NULL,
              chunk_index INTEGER NOT NULL,
              pages_json TEXT NOT NULL,
              text TEXT NOT NULL,
              llm_result_json TEXT,
              created_at TEXT NOT NULL,
              PRIMARY KEY (ingestion_id, chunk_index),
              FOREIGN KEY (ingestion_id) REFERENCES ingestions(id) ON DELETE CASCADE
            );

            CREATE INDEX IF NOT EXISTS idx_ingestions_sha ON ingestions(source_sha256);

            CREATE INDEX IF NOT EXISTS idx_tasks_status ON tasks(status);
            CREATE INDEX IF NOT EXISTS idx_workflows_status ON workflows(status);
            """
        )

        # lightweight migrations (prototype-friendly)
        if not _column_exists(conn, "tasks", "tags_json"):
            conn.execute("ALTER TABLE tasks ADD COLUMN tags_json TEXT NOT NULL DEFAULT '[]'")
        if not _column_exists(conn, "tasks", "meta_json"):
            conn.execute("ALTER TABLE tasks ADD COLUMN meta_json TEXT NOT NULL DEFAULT '{}'")
        if not _column_exists(conn, "workflows", "domains_json"):
            conn.execute("ALTER TABLE workflows ADD COLUMN domains_json TEXT NOT NULL DEFAULT '[]'")
        if not _column_exists(conn, "workflows", "tags_json"):
            conn.execute("ALTER TABLE workflows ADD COLUMN tags_json TEXT NOT NULL DEFAULT '[]'")
        if not _column_exists(conn, "workflows", "meta_json"):
            conn.execute("ALTER TABLE workflows ADD COLUMN meta_json TEXT NOT NULL DEFAULT '{}'")

        if not _column_exists(conn, "tasks", "domain"):
            conn.execute("ALTER TABLE tasks ADD COLUMN domain TEXT NOT NULL DEFAULT ''")

        # Backfill derived workflow domains when the column is introduced.
        _backfill_workflow_domains(conn)


def _seed_demo_users(conn: sqlite3.Connection) -> None:
    now = utc_now_iso()

    # Demo credentials are intentionally obvious.
    # Passwords are stored hashed, but also displayed on the login page via demo_password.
    demo = [
        ("jhendrix", "reviewer", "password1"),
        ("jjoplin", "author", "password2"),
        ("wcarlos", "assessment_author", "password5"),
        ("awinehouse", "content_publisher", "password6"),
        ("fmercury", "viewer", "password3"),
        ("bspringsteen", "audit", "password4"),
        ("kcobain", "admin", "admin"),
    ]

    # Best-effort renames from older demo sets (idempotent).
    # Only run when the source username exists and the target username does not.
    def _rename(old: str, new: str):
        src = conn.execute("SELECT 1 FROM users WHERE username=?", (old,)).fetchone()
        dst = conn.execute("SELECT 1 FROM users WHERE username=?", (new,)).fetchone()
        if src and not dst:
            conn.execute("UPDATE users SET username=? WHERE username=?", (new, old))

    _rename("mcury", "fmercury")
    _rename("dspringsteen", "bspringsteen")
    _rename("admin", "kcobain")
    _rename("mcarey", "wcarlos")
    _rename("publisher", "awinehouse")

    # NOTE: we intentionally do NOT auto-migrate roles.
    # assessment_author and content_publisher are separate roles.

    import secrets

    # Ensure each demo user exists (idempotent).
    for username, role, pw in demo:
        row = conn.execute(
            "SELECT id FROM users WHERE username=?",
            (username,),
        ).fetchone()
        if row:
            # Keep role + demo_password aligned for consistent demos.
            conn.execute(
                "UPDATE users SET role=?, demo_password=?, disabled_at=NULL WHERE username=?",
                (role, pw, username),
            )
            continue

        salt = secrets.token_bytes(16).hex()
        conn.execute(
            """
            INSERT INTO users(username, role, password_salt_hex, password_hash_hex, demo_password, created_at, created_by)
            VALUES (?,?,?,?,?,?,?)
            """,
            (username, role, salt, _hash_password(pw, salt), pw, now, "seed"),
        )


def _seed_demo_domains(conn: sqlite3.Connection) -> None:
    now = utc_now_iso()
    # Minimal initial registry (admin can manage later)
    initial = [
        "linux",
        "kubernetes",
        "postgres",
        "aws",
    ]
    for d in initial:
        conn.execute(
            "INSERT OR IGNORE INTO domains(name, created_at, created_by) VALUES (?,?,?)",
            (d, now, "seed"),
        )


def _seed_demo_entitlements(conn: sqlite3.Connection) -> None:
    now = utc_now_iso()

    def uid(name: str) -> int | None:
        r = conn.execute("SELECT id FROM users WHERE username=? AND disabled_at IS NULL", (name,)).fetchone()
        return int(r["id"]) if r else None

    # Keep this conservative: authors/reviewers only get linux by default.
    assignments: dict[str, list[str]] = {
        "jhendrix": ["linux"],
        "jjoplin": ["linux"],
    }

    for username, domains in assignments.items():
        u = uid(username)
        if not u:
            continue
        for d in domains:
            conn.execute(
                "INSERT OR IGNORE INTO user_domains(user_id, domain, created_at, created_by) VALUES (?,?,?,?)",
                (u, d, now, "seed"),
            )


def init_db() -> None:
    # Ensure both demo and blank DBs exist and are migrated.
    init_db_path(DB_DEMO_PATH)
    init_db_path(DB_BLANK_PATH)

    # Seed demo-friendly users for both DBs (safe for MVP).
    for p in (DB_DEMO_PATH, DB_BLANK_PATH):
        DB_PATH_CTX.set(p)
        with db() as conn:
            _seed_demo_users(conn)
            _seed_demo_domains(conn)
            _seed_demo_entitlements(conn)


@app.on_event("startup")
def _startup() -> None:
    init_db()


# --- Linting (warnings-only) ---

ABSTRACT_VERBS = {
    "edit",
    "configure",
    "set up",
    "setup",
    "manage",
    "ensure",
    "handle",
    "prepare",
    "troubleshoot",
}

STATE_CHANGE_VERBS = {
    "install",
    "mount",
    "enable",
    "add",
    "update",
    "remove",
    "create",
    "delete",
}


def _normalize_steps(raw: Any) -> list[dict[str, Any]]:
    """Return steps as list of {text, completion, actions?}.

    Canonical meaning:
      - text: what you are doing (intent)
      - completion: how you prove the Step is complete (required)
      - actions: optional sub-instructions for how to perform the Step in a specific tool/environment

    Backward compatible with legacy storage:
      - steps as list[str]
      - steps as list[{text, completion}]
    """
    if raw is None:
        return []
    if isinstance(raw, list):
        out: list[dict[str, Any]] = []
        for item in raw:
            if isinstance(item, str):
                out.append({"text": item, "completion": "", "actions": []})
            elif isinstance(item, dict):
                actions_raw = item.get("actions")
                actions: list[str] = []
                if isinstance(actions_raw, list):
                    actions = [str(x) for x in actions_raw if str(x).strip()]
                elif isinstance(actions_raw, str):
                    # allow a single multi-line string
                    actions = [ln.strip() for ln in actions_raw.splitlines() if ln.strip()]

                out.append(
                    {
                        "text": str(item.get("text", "")),
                        "completion": str(item.get("completion", "")),
                        "actions": actions,
                    }
                )
        # Drop empty rows
        return [s for s in out if (s.get("text") or "").strip() or (s.get("completion") or "").strip()]
    return []


def lint_steps(steps: Any) -> list[str]:
    warnings: list[str] = []

    normalized = _normalize_steps(steps)

    for i, step in enumerate(normalized, start=1):
        s = step.get("text", "")
        low = s.strip().lower()

        # Abstract/bundling verbs
        for v in ABSTRACT_VERBS:
            if low.startswith(v + " ") or low == v:
                if not re.search(r"`.+?`", s) and not re.search(
                    r"\b(confirm|verify|check)\b", low
                ):
                    warnings.append(
                        f"Step {i}: starts with abstract verb '{v}'. Prefer decomposed steps with explicit method + completion check."
                    )
                break

        # Multi-action detector (refined): only warn when conjunctions likely hide multiple procedural operations.
        if re.search(r"\b(and|then|also|as well as)\b", low):
            verb_markers = (
                list(ABSTRACT_VERBS)
                + list(STATE_CHANGE_VERBS)
                + [
                    "run",
                    "open",
                    "copy",
                    "move",
                    "create",
                    "delete",
                    "set",
                    "insert",
                    "save",
                    "restart",
                    "reload",
                    "verify",
                    "confirm",
                    "record",
                    "list",
                    "check",
                    "edit",
                ]
            )
            # Count verb-like tokens appearing as word starts.
            hits = 0
            for v in set(verb_markers):
                if re.search(rf"\b{re.escape(v)}\b", low):
                    hits += 1
                if hits >= 2:
                    break
            if hits >= 2:
                warnings.append(
                    f"Step {i}: may include multiple procedural operations (conjunction + multiple verbs). Consider splitting."
                )

        # Verification expectation
        if any(low.startswith(v + " ") or low == v for v in STATE_CHANGE_VERBS):
            if not re.search(r"\b(confirm|verify|check)\b", low) and not re.search(r"`.+?`", s):
                warnings.append(
                    f"Step {i}: appears to change state; include an explicit confirmation check (command/UI observable) or follow with a check step."
                )

    return warnings


def _zip_steps(step_text: list[str], step_completion: list[str], step_actions: list[str]) -> list[dict[str, Any]]:
    # Keep ordering.
    out: list[dict[str, Any]] = []
    for t, c, a in zip(step_text, step_completion, step_actions, strict=False):
        actions = [ln.strip() for ln in (a or "").splitlines() if ln.strip()]
        out.append({"text": (t or "").strip(), "completion": (c or "").strip(), "actions": actions})

    # If lists are mismatched, extend with remaining text.
    longest = max(len(step_text), len(step_completion), len(step_actions))
    for i in range(len(out), longest):
        t = step_text[i] if i < len(step_text) else ""
        c = step_completion[i] if i < len(step_completion) else ""
        a = step_actions[i] if i < len(step_actions) else ""
        actions = [ln.strip() for ln in (a or "").splitlines() if ln.strip()]
        out.append({"text": (t or "").strip(), "completion": (c or "").strip(), "actions": actions})

    # Drop empty rows
    return [s for s in out if (s.get("text") or "").strip() or (s.get("completion") or "").strip()]


def _validate_steps_required(steps: list[dict[str, Any]]) -> None:
    """Enforce step contract: step text + completion are required; actions are optional."""
    if not steps:
        raise HTTPException(status_code=400, detail="At least one step is required")
    for idx, st in enumerate(steps, start=1):
        if not (st.get("text") or "").strip():
            raise HTTPException(status_code=400, detail=f"Step {idx}: step text is required")
        if not (st.get("completion") or "").strip():
            raise HTTPException(status_code=400, detail=f"Step {idx}: completion text is required")


# --- Helpers ---


def _pdf_extract_pages(pdf_path: str) -> list[dict[str, Any]]:
    reader = PdfReader(pdf_path)
    pages: list[dict[str, Any]] = []
    for idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append({"page": idx, "text": text})
    return pages


def _chunk_text(pages: list[dict[str, Any]], max_chars: int = 12000) -> list[dict[str, Any]]:
    """Chunk by character count, preserving page numbers."""
    chunks: list[dict[str, Any]] = []
    buf: list[str] = []
    buf_pages: list[int] = []
    size = 0

    def flush():
        nonlocal buf, buf_pages, size
        if not buf:
            return
        chunks.append({"pages": sorted(set(buf_pages)), "text": "\n\n".join(buf).strip()})
        buf, buf_pages, size = [], [], 0

    for p in pages:
        t = (p.get("text") or "").strip()
        if not t:
            continue
        header = f"[PAGE {p['page']}]"
        block = header + "\n" + t
        if size + len(block) > max_chars and buf:
            flush()
        buf.append(block)
        buf_pages.append(int(p["page"]))
        size += len(block)

    flush()
    return chunks


def _lmstudio_probe(base_url: str | None = None) -> dict[str, Any]:
    """Fast health probe for LM Studio.

    Returns {ok: bool, detail: str, base_url: str}.
    """
    bu = (base_url or LMSTUDIO_BASE_URL).rstrip("/")
    try:
        with httpx.Client(timeout=httpx.Timeout(2.0, connect=1.0)) as client:
            # Prefer OpenAI-compatible endpoint; HEAD is not always supported, so use GET.
            r = client.get(f"{bu}/v1/models")
            if r.status_code < 400:
                return {"ok": True, "detail": "ok", "base_url": bu}
            # Fallback probe
            r2 = client.get(f"{bu}/api/v1/models")
            if r2.status_code < 400:
                return {"ok": True, "detail": "ok", "base_url": bu}
            return {"ok": False, "detail": f"HTTP {r.status_code}", "base_url": bu}
    except Exception as e:
        return {"ok": False, "detail": str(e), "base_url": bu}


def _lmstudio_chat(messages: list[dict[str, str]], temperature: float = 0.2, max_tokens: int = 2000, base_url: str | None = None) -> str:
    """Call LM Studio local server.

    NOTE: Some LM Studio model prompt templates only support `user` + `assistant` roles.
    We normalize away `system` by prepending it to the first user message.

    Supports both:
      - OpenAI-compatible: POST /v1/chat/completions
      - LM Studio API: POST /api/v1/chat

    Returns the assistant content.
    """

    # Normalize roles: merge system content into first user message.
    sys_parts: list[str] = []
    norm: list[dict[str, str]] = []
    for m in messages:
        role = (m.get("role") or "").strip()
        content = m.get("content") or ""
        if role == "system":
            if content.strip():
                sys_parts.append(content.strip())
            continue
        if role not in ("user", "assistant"):
            # drop unknown roles in MVP
            continue
        norm.append({"role": role, "content": content})

    if sys_parts:
        sys_blob = "\n\n".join(sys_parts)
        if norm and norm[0]["role"] == "user":
            norm[0]["content"] = f"SYSTEM INSTRUCTIONS:\n{sys_blob}\n\n" + (norm[0]["content"] or "")
        else:
            norm.insert(0, {"role": "user", "content": f"SYSTEM INSTRUCTIONS:\n{sys_blob}"})

    payload = {
        "model": LMSTUDIO_MODEL,
        "messages": norm,
        "temperature": temperature,
        "max_tokens": max_tokens,
    }

    bu = (base_url or LMSTUDIO_BASE_URL).rstrip("/")

    try:
        with httpx.Client(timeout=httpx.Timeout(300.0, connect=30.0)) as client:
            r = client.post(f"{bu}/v1/chat/completions", json=payload)
            if r.status_code == 404:
                # Fallback to LM Studio API (only when OpenAI-compatible endpoint is absent)
                r2 = client.post(f"{bu}/api/v1/chat", json=payload)
                if r2.status_code >= 400:
                    raise HTTPException(
                        status_code=502,
                        detail=f"LM Studio API error {r2.status_code}: {r2.text[:500]}",
                    )
                data2 = r2.json()
                if isinstance(data2, dict) and "choices" in data2:
                    return data2["choices"][0]["message"]["content"]
                if isinstance(data2, dict) and "message" in data2 and isinstance(data2["message"], dict):
                    return data2["message"].get("content", "")
                return json.dumps(data2)

            if r.status_code >= 400:
                raise HTTPException(
                    status_code=502,
                    detail=f"LM Studio OpenAI-compatible error {r.status_code}: {r.text[:500]}",
                )

            data = r.json()
            return data["choices"][0]["message"]["content"]
    except httpx.ReadTimeout as e:
        raise HTTPException(status_code=504, detail=f"LM Studio timed out: {e}")
    except httpx.ConnectError as e:
        raise HTTPException(status_code=502, detail=f"LM Studio connection error: {e}")
    except httpx.HTTPError as e:
        # Catch-all for other httpx failures
        raise HTTPException(status_code=502, detail=f"LM Studio HTTP error: {e}")


def _json_load(s: str) -> Any:
    return json.loads(s) if s else None


def _json_dump(v: Any) -> str:
    return json.dumps(v, ensure_ascii=False)


def _sha256_bytes(b: bytes) -> str:
    h = hashlib.sha256()
    h.update(b)
    return h.hexdigest()


def _short_code(prefix: str, record_id: str) -> str:
    """Deterministic short display id (for human-visible trace tags)."""
    h = hashlib.sha256((record_id or "").encode("utf-8", errors="ignore")).hexdigest().upper()
    return f"{prefix}-{h[:6]}"


def _norm_text(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip().lower())


def _task_fingerprint(task: dict[str, Any]) -> str:
    """Deterministic fingerprint for exact-ish dedupe."""
    title = _norm_text(str(task.get("title", "")))
    outcome = _norm_text(str(task.get("outcome", "")))
    steps = task.get("steps") or []
    steps_norm = _normalize_steps(steps)
    parts: list[str] = [title, outcome]
    for st in steps_norm:
        parts.append(_norm_text(str(st.get("text", ""))))
        parts.append(_norm_text(str(st.get("completion", ""))))
    raw = "\n".join(parts).encode("utf-8", errors="ignore")
    return _sha256_bytes(raw)


def _extract_step_targets(steps: list[dict[str, Any]]) -> set[str]:
    """Extract rough targets for near-duplicate hints (paths, services, packages)."""
    targets: set[str] = set()
    path_re = re.compile(r"(/etc/[^\s]+|/var/[^\s]+|/usr/[^\s]+|/opt/[^\s]+)")
    svc_re = re.compile(r"\b(systemctl)\s+(restart|reload|enable|disable)\s+([a-zA-Z0-9_.@-]+)")
    pkg_re = re.compile(r"\bapt(-get)?\s+install\s+(-y\s+)?([a-zA-Z0-9+_.:-]+)")

    for st in steps or []:
        t = (st.get("text") or "") + "\n" + (st.get("completion") or "")
        for m in path_re.findall(t):
            targets.add(m.lower())
        for m in svc_re.findall(t):
            targets.add(f"service:{m[2].lower()}")
        for m in pkg_re.findall(t):
            targets.add(f"pkg:{m[2].lower()}")
    return targets


def _near_duplicate_score(a: dict[str, Any], b: dict[str, Any]) -> float:
    """Heuristic similarity score in [0,1]."""
    a_steps = _normalize_steps(a.get("steps") or [])
    b_steps = _normalize_steps(b.get("steps") or [])

    a_title = set(_norm_text(str(a.get("title", ""))).split())
    b_title = set(_norm_text(str(b.get("title", ""))).split())
    a_out = set(_norm_text(str(a.get("outcome", ""))).split())
    b_out = set(_norm_text(str(b.get("outcome", ""))).split())

    def jacc(x: set[str], y: set[str]) -> float:
        if not x and not y:
            return 0.0
        return len(x & y) / max(1, len(x | y))

    title_sim = jacc(a_title, b_title)
    out_sim = jacc(a_out, b_out)

    a_tgt = _extract_step_targets(a_steps)
    b_tgt = _extract_step_targets(b_steps)
    tgt_sim = jacc(a_tgt, b_tgt)

    # Weighted: outcome + targets matter more than title.
    return 0.20 * title_sim + 0.45 * out_sim + 0.35 * tgt_sim


def audit(
    entity_type: str,
    record_id: str,
    version: int,
    action: str,
    actor: str,
    note: str | None = None,
    *,
    conn: sqlite3.Connection | None = None,
) -> None:
    """Write an audit_log row.

    Important: if you're already inside `with db() as conn:` and have performed writes,
    pass that same `conn=` here. Otherwise audit() will open a second connection and
    SQLite can throw `OperationalError: database is locked`.
    """
    if conn is None:
        with db() as conn2:
            conn2.execute(
                "INSERT INTO audit_log(entity_type, record_id, version, action, actor, at, note) VALUES (?,?,?,?,?,?,?)",
                (entity_type, record_id, version, action, actor, utc_now_iso(), note),
            )
        return

    conn.execute(
        "INSERT INTO audit_log(entity_type, record_id, version, action, actor, at, note) VALUES (?,?,?,?,?,?,?)",
        (entity_type, record_id, version, action, actor, utc_now_iso(), note),
    )


def get_latest_version(conn: sqlite3.Connection, table: str, record_id: str) -> int | None:
    row = conn.execute(
        f"SELECT MAX(version) AS v FROM {table} WHERE record_id=?", (record_id,)
    ).fetchone()
    return int(row["v"]) if row and row["v"] is not None else None


def require_can_edit(status: str) -> None:
    # Legacy guardrail (we now treat *all* records as immutable and always create new versions).
    if status == "confirmed":
        raise HTTPException(
            status_code=409,
            detail="Confirmed records are immutable. Create a new version.",
        )


def parse_lines(text: str) -> list[str]:
    # Accept newline-separated list
    lines = [ln.strip() for ln in (text or "").splitlines()]
    return [ln for ln in lines if ln]


def parse_tags(text: str) -> list[str]:
    raw = (text or "").strip()
    if not raw:
        return []
    parts = [p.strip() for p in raw.split(",")]
    return [p for p in parts if p]


def parse_meta(text: str) -> dict[str, str]:
    meta: dict[str, str] = {}
    for ln in parse_lines(text or ""):
        if "=" not in ln:
            # ignore malformed lines in MVP
            continue
        k, v = ln.split("=", 1)
        k = k.strip()
        v = v.strip()
        if not k:
            continue
        meta[k] = v
    return meta


# --- Routes ---


@app.get("/login", response_class=HTMLResponse)
def login_form(request: Request):
    # Show demo users and their passwords (MVP demo convenience).
    with db() as conn:
        users = conn.execute(
            "SELECT username, role, COALESCE(demo_password, '') AS demo_password FROM users WHERE disabled_at IS NULL ORDER BY role DESC, username ASC"
        ).fetchall()

    return templates.TemplateResponse(request, "login.html", {"users": [dict(u) for u in users]})


@app.post("/login")
def login_run(request: Request, username: str = Form(""), password: str = Form("")):
    username = (username or "").strip()
    password = password or ""
    if not username:
        raise HTTPException(status_code=400, detail="username is required")

    with db() as conn:
        u = conn.execute(
            """
            SELECT id, username, role, password_salt_hex, password_hash_hex
            FROM users
            WHERE username=? AND disabled_at IS NULL
            """,
            (username,),
        ).fetchone()
        if not u:
            raise HTTPException(status_code=403, detail="Invalid credentials")
        if not _verify_password(password, str(u["password_salt_hex"]), str(u["password_hash_hex"])):
            raise HTTPException(status_code=403, detail="Invalid credentials")

        token = _new_session_token()
        conn.execute(
            "INSERT INTO sessions(token, user_id, created_at) VALUES (?,?,?)",
            (token, int(u["id"]), utc_now_iso()),
        )

    resp = RedirectResponse(url="/", status_code=303)
    resp.set_cookie(SESSION_COOKIE, token, httponly=True, samesite="lax")
    return resp


@app.post("/logout")
def logout(request: Request):
    token = (request.cookies.get(SESSION_COOKIE) or "").strip()
    if token:
        with db() as conn:
            conn.execute("UPDATE sessions SET revoked_at=? WHERE token=? AND revoked_at IS NULL", (utc_now_iso(), token))

    resp = RedirectResponse(url="/login", status_code=303)
    resp.delete_cookie(SESSION_COOKIE)
    return resp


@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    # Reviewer default: show their work queue.
    if request.state.role == "reviewer":
        return RedirectResponse(url="/review", status_code=303)

    return templates.TemplateResponse(request, "home.html", {})


@app.get("/explainer", response_class=HTMLResponse)
def explainer(request: Request):
    """Plain-language explainer page.

    Auth required via middleware.
    """
    return templates.TemplateResponse(request, "explainer.html", {})


# --- DB switching (MVP) ---

@app.get("/_pulse")
def pulse(request: Request):
    """Return small operational counters for the UI pulse strip.

    Auth is required via middleware.
    """
    role = request.state.role
    user = request.state.user

    with db() as conn:
        # Task counts
        task_counts = {
            r["status"]: int(r["c"])
            for r in conn.execute("SELECT status, COUNT(*) AS c FROM tasks GROUP BY status").fetchall()
        }

        # Reviewer-scoped counts (domain entitlements)
        reviewer_pending = None
        reviewer_domains: list[str] = []
        if role in ("reviewer", "admin"):
            reviewer_domains = _user_domains(conn, user)

            if role == "admin":
                # admin sees everything
                reviewer_pending = int(
                    conn.execute(
                        "SELECT ("
                        " (SELECT COUNT(*) FROM tasks WHERE status='submitted') +"
                        " (SELECT COUNT(*) FROM workflows WHERE status='submitted') +"
                        " (SELECT COUNT(*) FROM assessment_items WHERE status='submitted')"
                        ") AS c"
                    ).fetchone()["c"]
                )
            else:
                if reviewer_domains:
                    qmarks = ",".join(["?"] * len(reviewer_domains))

                    t = int(
                        conn.execute(
                            f"SELECT COUNT(*) AS c FROM tasks WHERE status='submitted' AND domain IN ({qmarks})",
                            reviewer_domains,
                        ).fetchone()["c"]
                    )
                    # Workflows: domain match is derived via domains_json; filter in Python for portability.
                    w_rows = conn.execute(
                        "SELECT domains_json FROM workflows WHERE status='submitted'"
                    ).fetchall()
                    w = 0
                    rdset = {d.strip().lower() for d in reviewer_domains if d}
                    for wr in w_rows:
                        doms = [str(x).strip().lower() for x in (_json_load(wr["domains_json"]) or [])]
                        if rdset.intersection(doms):
                            w += 1

                    # Assessments: same.
                    a_rows = conn.execute(
                        "SELECT domains_json FROM assessment_items WHERE status='submitted'"
                    ).fetchall()
                    a = 0
                    for ar in a_rows:
                        doms = [str(x).strip().lower() for x in (_json_load(ar["domains_json"]) or [])]
                        if rdset.intersection(doms):
                            a += 1

                    reviewer_pending = t + w + a
                else:
                    reviewer_pending = 0

        # Workflow counts
        wf_counts = {
            r["status"]: int(r["c"])
            for r in conn.execute("SELECT status, COUNT(*) AS c FROM workflows GROUP BY status").fetchall()
        }

        # Assessment counts
        as_counts = {
            r["status"]: int(r["c"])
            for r in conn.execute("SELECT status, COUNT(*) AS c FROM assessment_items GROUP BY status").fetchall()
        }

        last_audit = conn.execute("SELECT at, actor, action FROM audit_log ORDER BY at DESC LIMIT 1").fetchone()

    return {
        "tasks": {
            "draft": task_counts.get("draft", 0),
            "submitted": task_counts.get("submitted", 0),
            "confirmed": task_counts.get("confirmed", 0),
        },
        "workflows": {
            "draft": wf_counts.get("draft", 0),
            "submitted": wf_counts.get("submitted", 0),
            "confirmed": wf_counts.get("confirmed", 0),
        },
        "assessments": {
            "draft": as_counts.get("draft", 0),
            "submitted": as_counts.get("submitted", 0),
            "confirmed": as_counts.get("confirmed", 0),
        },
        "review": {
            "pending": reviewer_pending,
            "domains": reviewer_domains,
        },
        "audit": {
            "last": dict(last_audit) if last_audit else None,
        },
    }


@app.get("/db", response_class=HTMLResponse)
def db_switch_form(request: Request):
    require(request.state.role, "db:switch")
    profiles = [{"key": k, "label": k} for k in [DB_KEY_DEMO, DB_KEY_BLANK] + _list_custom_db_keys()]
    return templates.TemplateResponse(request, "db_switch.html", {"profiles": profiles})


@app.post("/db/switch")
def db_switch(request: Request, db_key: str = Form(DB_KEY_DEMO)):
    require(request.state.role, "db:switch")
    key = (db_key or DB_KEY_DEMO).strip().lower()
    if key not in _available_db_keys():
        raise HTTPException(status_code=400, detail="Invalid db_key")

    resp = RedirectResponse(url="/", status_code=303)
    resp.set_cookie(DB_KEY_COOKIE, key, httponly=False, samesite="lax")
    return resp


@app.post("/db/create")
def db_create(request: Request, db_key: str = Form("")):
    require(request.state.role, "db:switch")
    key = (db_key or "").strip().lower()
    if not key:
        raise HTTPException(status_code=400, detail="db_key is required")

    _create_custom_db_profile(key)

    # Switch to it immediately.
    resp = RedirectResponse(url="/db", status_code=303)
    resp.set_cookie(DB_KEY_COOKIE, key, httponly=False, samesite="lax")
    return resp


@app.get("/admin/users", response_class=HTMLResponse)
def admin_users(request: Request, error: str | None = None):
    require_admin(request)
    with db() as conn:
        rows = conn.execute(
            "SELECT id, username, role, COALESCE(demo_password, '') AS demo_password, disabled_at FROM users ORDER BY disabled_at IS NOT NULL, role DESC, username ASC"
        ).fetchall()

        # Attach per-user domains
        users: list[dict[str, Any]] = []
        for r in rows:
            u = dict(r)
            doms = conn.execute(
                "SELECT domain FROM user_domains WHERE user_id=? ORDER BY domain ASC",
                (int(r["id"]),),
            ).fetchall()
            u["domains"] = [str(x["domain"]) for x in doms]
            users.append(u)

    return templates.TemplateResponse(request, "admin/users.html", {"users": users, "error": error})


@app.post("/admin/users/create")
def admin_users_create(request: Request, username: str = Form(""), role: str = Form("viewer"), password: str = Form("")):
    require_admin(request)
    username = (username or "").strip()
    password = password or ""
    if not username:
        raise HTTPException(status_code=400, detail="username is required")
    if role not in ROLE_ORDER:
        raise HTTPException(status_code=400, detail="invalid role")
    if not password:
        raise HTTPException(status_code=400, detail="password is required")

    import secrets

    salt = secrets.token_bytes(16).hex()
    with db() as conn:
        try:
            conn.execute(
                """
                INSERT INTO users(username, role, password_salt_hex, password_hash_hex, demo_password, created_at, created_by)
                VALUES (?,?,?,?,?,?,?)
                """,
                (username, role, salt, _hash_password(password, salt), password, utc_now_iso(), request.state.user),
            )
            audit("user", username, 1, "create", request.state.user, note=f"role={role}", conn=conn)
        except sqlite3.IntegrityError:
            raise HTTPException(status_code=409, detail="username already exists")

    return RedirectResponse(url="/admin/users", status_code=303)


@app.post("/admin/users/reset")
def admin_users_reset(request: Request, username: str = Form(""), password: str = Form("")):
    require_admin(request)
    username = (username or "").strip()
    password = password or ""
    if not username or not password:
        raise HTTPException(status_code=400, detail="username and password required")

    import secrets

    salt = secrets.token_bytes(16).hex()
    with db() as conn:
        row = conn.execute("SELECT id FROM users WHERE username=?", (username,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="user not found")
        conn.execute(
            "UPDATE users SET password_salt_hex=?, password_hash_hex=?, demo_password=? WHERE username=?",
            (salt, _hash_password(password, salt), password, username),
        )
        # Revoke sessions
        conn.execute("UPDATE sessions SET revoked_at=? WHERE user_id=? AND revoked_at IS NULL", (utc_now_iso(), int(row["id"])))
        audit("user", username, 1, "reset_password", request.state.user, conn=conn)

    return RedirectResponse(url="/admin/users", status_code=303)


@app.post("/admin/users/disable")
def admin_users_disable(request: Request, username: str = Form("")):
    require_admin(request)
    username = (username or "").strip()
    if not username:
        raise HTTPException(status_code=400, detail="username required")

    with db() as conn:
        row = conn.execute("SELECT id FROM users WHERE username=?", (username,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="user not found")
        conn.execute("UPDATE users SET disabled_at=? WHERE username=?", (utc_now_iso(), username))
        conn.execute("UPDATE sessions SET revoked_at=? WHERE user_id=? AND revoked_at IS NULL", (utc_now_iso(), int(row["id"])))
        audit("user", username, 1, "disable", request.state.user, conn=conn)

    # If you disabled yourself, you'll be bounced to /login on next request.
    return RedirectResponse(url="/admin/users", status_code=303)


@app.post("/admin/users/enable")
def admin_users_enable(request: Request, username: str = Form("")):
    require_admin(request)
    username = (username or "").strip()
    if not username:
        raise HTTPException(status_code=400, detail="username required")

    with db() as conn:
        row = conn.execute("SELECT 1 FROM users WHERE username=?", (username,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="user not found")
        conn.execute("UPDATE users SET disabled_at=NULL WHERE username=?", (username,))
        audit("user", username, 1, "enable", request.state.user, conn=conn)

    return RedirectResponse(url="/admin/users", status_code=303)


@app.post("/admin/users/delete")
def admin_users_delete(request: Request, username: str = Form("")):
    require_admin(request)
    username = (username or "").strip()
    if not username:
        raise HTTPException(status_code=400, detail="username required")
    if username == request.state.user:
        raise HTTPException(status_code=400, detail="cannot delete the current user")

    with db() as conn:
        row = conn.execute("SELECT id FROM users WHERE username=?", (username,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="user not found")
        uid = int(row["id"])
        conn.execute("DELETE FROM sessions WHERE user_id=?", (uid,))
        conn.execute("DELETE FROM users WHERE id=?", (uid,))
        audit("user", username, 1, "delete", request.state.user, conn=conn)

    return RedirectResponse(url="/admin/users", status_code=303)


@app.post("/admin/users/domains")
def admin_user_domains_form(request: Request, username: str = Form("")):
    require_admin(request)
    username = (username or "").strip()
    if not username:
        raise HTTPException(status_code=400, detail="username required")

    with db() as conn:
        u = conn.execute("SELECT id, username, role FROM users WHERE username=?", (username,)).fetchone()
        if not u:
            raise HTTPException(status_code=404, detail="user not found")

        domains = _active_domains(conn)
        selected_rows = conn.execute("SELECT domain FROM user_domains WHERE user_id=?", (int(u["id"]),)).fetchall()
        selected = {str(r["domain"]) for r in selected_rows}

    return templates.TemplateResponse(
        request,
        "admin/user_domains.html",
        {"user": dict(u), "domains": domains, "selected": selected},
    )


@app.post("/admin/users/domains/save")
def admin_user_domains_save(request: Request, username: str = Form(""), domain: list[str] = Form([])):
    require_admin(request)
    username = (username or "").strip()
    selected = sorted({(d or "").strip().lower() for d in (domain or []) if (d or "").strip()})

    with db() as conn:
        u = conn.execute("SELECT id FROM users WHERE username=?", (username,)).fetchone()
        if not u:
            raise HTTPException(status_code=404, detail="user not found")

        allowed = set(_active_domains(conn))
        for d in selected:
            if d not in allowed:
                raise HTTPException(status_code=400, detail=f"Invalid domain '{d}'")

        uid = int(u["id"])
        conn.execute("DELETE FROM user_domains WHERE user_id=?", (uid,))
        now = utc_now_iso()
        for d in selected:
            conn.execute(
                "INSERT INTO user_domains(user_id, domain, created_at, created_by) VALUES (?,?,?,?)",
                (uid, d, now, request.state.user),
            )

        audit("user", username, 1, "set_domains", request.state.user, note=",".join(selected), conn=conn)

    return RedirectResponse(url="/admin/users", status_code=303)


@app.get("/admin/domains", response_class=HTMLResponse)
def admin_domains(request: Request, error: str | None = None):
    require_admin(request)
    with db() as conn:
        rows = conn.execute("SELECT name, created_at, created_by, disabled_at FROM domains ORDER BY name ASC").fetchall()
    return templates.TemplateResponse(request, "admin/domains.html", {"domains": [dict(r) for r in rows], "error": error})


@app.post("/admin/domains/create")
def admin_domains_create(request: Request, name: str = Form("")):
    require_admin(request)
    name_norm = (name or "").strip().lower()
    if not name_norm:
        raise HTTPException(status_code=400, detail="name required")
    if not re.fullmatch(r"[a-z0-9][a-z0-9_-]*", name_norm):
        raise HTTPException(status_code=400, detail="invalid domain name")

    with db() as conn:
        try:
            conn.execute(
                "INSERT INTO domains(name, created_at, created_by) VALUES (?,?,?)",
                (name_norm, utc_now_iso(), request.state.user),
            )
            audit("domain", name_norm, 1, "create", request.state.user, conn=conn)
        except sqlite3.IntegrityError:
            raise HTTPException(status_code=409, detail="domain already exists")

    return RedirectResponse(url="/admin/domains", status_code=303)


@app.post("/admin/domains/disable")
def admin_domains_disable(request: Request, name: str = Form("")):
    require_admin(request)
    name_norm = (name or "").strip().lower()
    with db() as conn:
        row = conn.execute("SELECT 1 FROM domains WHERE name=?", (name_norm,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="domain not found")
        conn.execute("UPDATE domains SET disabled_at=? WHERE name=?", (utc_now_iso(), name_norm))
        audit("domain", name_norm, 1, "disable", request.state.user, conn=conn)

    return RedirectResponse(url="/admin/domains", status_code=303)


@app.post("/admin/domains/enable")
def admin_domains_enable(request: Request, name: str = Form("")):
    require_admin(request)
    name_norm = (name or "").strip().lower()
    with db() as conn:
        row = conn.execute("SELECT 1 FROM domains WHERE name=?", (name_norm,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="domain not found")
        conn.execute("UPDATE domains SET disabled_at=NULL WHERE name=?", (name_norm,))
        audit("domain", name_norm, 1, "enable", request.state.user, conn=conn)

    return RedirectResponse(url="/admin/domains", status_code=303)


@app.post("/admin/domains/delete")
def admin_domains_delete(request: Request, name: str = Form("")):
    require_admin(request)
    name_norm = (name or "").strip().lower()
    with db() as conn:
        row = conn.execute("SELECT 1 FROM domains WHERE name=?", (name_norm,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="domain not found")
        try:
            conn.execute("DELETE FROM domains WHERE name=?", (name_norm,))
        except sqlite3.IntegrityError:
            raise HTTPException(status_code=409, detail="domain is referenced by user entitlements; disable it instead")
        audit("domain", name_norm, 1, "delete", request.state.user, conn=conn)

    return RedirectResponse(url="/admin/domains", status_code=303)


@app.get("/review", response_class=HTMLResponse)
def review_queue(request: Request):
    # Reviewers and admins only. (Admin implicitly has all domains.)
    if request.state.role not in ("reviewer", "admin"):
        raise HTTPException(status_code=403, detail="Forbidden: reviewer/admin only")

    with db() as conn:
        # Determine authorized domains
        if request.state.role == "admin":
            doms = _active_domains(conn)
        else:
            uid = _user_id(conn, request.state.user)
            dom_rows = conn.execute("SELECT domain FROM user_domains WHERE user_id=?", (uid,)).fetchall() if uid else []
            doms = [str(r["domain"]) for r in dom_rows]

        items: list[dict[str, Any]] = []
        if doms:
            dset = {d.strip().lower() for d in doms if d}
            qmarks = ",".join(["?"] * len(doms))

            # Tasks
            t_rows = conn.execute(
                f"SELECT record_id, version, title, status, domain FROM tasks WHERE status='submitted' AND domain IN ({qmarks}) ORDER BY domain ASC, title ASC",
                doms,
            ).fetchall()
            for r in t_rows:
                items.append({"type": "task", "record_id": r["record_id"], "version": int(r["version"]), "title": r["title"], "status": r["status"], "domains": [str(r["domain"])],})

            # Workflows (domain derived)
            w_rows = conn.execute(
                "SELECT record_id, version, title, status, domains_json FROM workflows WHERE status='submitted' ORDER BY title ASC"
            ).fetchall()
            for r in w_rows:
                wdoms = [str(x).strip().lower() for x in (_json_load(r["domains_json"]) or []) if str(x).strip()]
                if dset.intersection(wdoms):
                    items.append({"type": "workflow", "record_id": r["record_id"], "version": int(r["version"]), "title": r["title"], "status": r["status"], "domains": wdoms,})

            # Assessments
            a_rows = conn.execute(
                "SELECT record_id, version, stem, status, domains_json FROM assessment_items WHERE status='submitted' ORDER BY stem ASC"
            ).fetchall()
            for r in a_rows:
                adoms = [str(x).strip().lower() for x in (_json_load(r["domains_json"]) or []) if str(x).strip()]
                if dset.intersection(adoms):
                    title = str(r["stem"])[:80]
                    items.append({"type": "assessment", "record_id": r["record_id"], "version": int(r["version"]), "title": title, "status": r["status"], "domains": adoms,})

            # stable sort: domain then type
            items.sort(key=lambda it: ("".join(it.get("domains") or []), it.get("type"), it.get("title")))

    return templates.TemplateResponse(request, "review_queue.html", {"items": items, "domains": doms})


@app.get("/audit", response_class=HTMLResponse)
def audit_list(
    request: Request,
    entity_type: str | None = None,
    record_id: str | None = None,
    limit: int = 200,
):
    require(request.state.role, "audit:view")

    entity_type = (entity_type or "").strip() or None
    record_id = (record_id or "").strip() or None
    limit = max(1, min(int(limit or 200), 1000))

    where: list[str] = []
    params: list[Any] = []
    if entity_type:
        where.append("entity_type=?")
        params.append(entity_type)
    if record_id:
        where.append("record_id=?")
        params.append(record_id)

    sql = "SELECT id, entity_type, record_id, version, action, actor, at, note FROM audit_log"
    if where:
        sql += " WHERE " + " AND ".join(where)
    sql += " ORDER BY at DESC, id DESC LIMIT ?"
    params.append(limit)

    with db() as conn:
        rows = conn.execute(sql, params).fetchall()

    items = [dict(r) for r in rows]

    return templates.TemplateResponse(
        request,
        "audit_list.html",
        {"items": items, "entity_type": entity_type, "record_id": record_id, "limit": limit},
    )


@app.get("/_lmstudio/status")
def lmstudio_status(request: Request, base_url: str | None = None):
    require(request.state.role, "import:pdf")
    probe = _lmstudio_probe(base_url)
    return {"ok": bool(probe.get("ok")), "base_url": str(probe.get("base_url")), "detail": str(probe.get("detail"))}


@app.get("/import/pdf", response_class=HTMLResponse)
def import_pdf_form(request: Request):
    require(request.state.role, "import:pdf")
    actor = request.state.user

    with db() as conn:
        rows = conn.execute(
            "SELECT id, filename, created_at, status, cursor_chunk FROM ingestions WHERE source_type='pdf' AND created_by=? ORDER BY created_at DESC LIMIT 50",
            (actor,),
        ).fetchall()

    # Base URL is client-overridable (browser-local). Server default shown as fallback.
    probe = _lmstudio_probe()

    return templates.TemplateResponse(
        request,
        "import_pdf.html",
        {
            "lmstudio_base_url": str(probe.get("base_url")) or LMSTUDIO_BASE_URL,
            "lmstudio_model": LMSTUDIO_MODEL,
            "lmstudio_ok": bool(probe.get("ok")),
            "lmstudio_detail": str(probe.get("detail")),
            "ingestions": [dict(r) for r in rows],
        },
    )


@app.post("/import/pdf/prepare")
def import_pdf_prepare(
    request: Request,
    pdf: UploadFile = File(...),
    max_tasks: int = Form(10),
    max_chunks: int = Form(8),
    actor_note: str = Form("Imported from PDF"),
    lmstudio_base_url: str = Form(""),
):
    require(request.state.role, "import:pdf")
    actor = request.state.user

    probe = _lmstudio_probe(lmstudio_base_url)
    if not probe.get("ok"):
        raise HTTPException(status_code=502, detail=f"LM Studio is not reachable at {probe.get('base_url')} ({probe.get('detail')})")

    # Save upload + compute hash identity
    os.makedirs(UPLOADS_DIR, exist_ok=True)
    safe_name = re.sub(r"[^a-zA-Z0-9._-]+", "_", pdf.filename or "upload.pdf")
    file_id = str(uuid.uuid4())
    out_path = os.path.join(UPLOADS_DIR, f"{file_id}__{safe_name}")
    file_bytes = pdf.file.read()
    with open(out_path, "wb") as f:
        f.write(file_bytes)

    sha = _sha256_bytes(file_bytes)

    max_tasks = max(1, min(int(max_tasks), 50))
    max_chunks = max(1, min(int(max_chunks), 50))

    with db() as conn:
        existing = conn.execute(
            "SELECT * FROM ingestions WHERE source_type='pdf' AND source_sha256=? AND created_by=? ORDER BY created_at DESC LIMIT 1",
            (sha, actor),
        ).fetchone()

        if existing:
            ingestion_id = str(existing["id"])
        else:
            ingestion_id = str(uuid.uuid4())
            conn.execute(
                "INSERT INTO ingestions(id, source_type, source_sha256, filename, created_by, created_at, status, cursor_chunk, max_tasks_per_run, note) VALUES (?,?,?,?,?,?,?,?,?,?)",
                (ingestion_id, "pdf", sha, safe_name, actor, utc_now_iso(), "draft", 0, max_tasks, actor_note.strip() or "Imported from PDF"),
            )

        # If we don't have cached chunks yet, extract + store now.
        cached = conn.execute(
            "SELECT 1 FROM ingestion_chunks WHERE ingestion_id=? LIMIT 1",
            (ingestion_id,),
        ).fetchone()

        if not cached:
            pages = _pdf_extract_pages(out_path)
            chunks = _chunk_text(pages, max_chars=12000)
            if not chunks:
                raise HTTPException(status_code=400, detail="No extractable text found in PDF")
            now = utc_now_iso()
            for idx, ch in enumerate(chunks):
                conn.execute(
                    "INSERT OR REPLACE INTO ingestion_chunks(ingestion_id, chunk_index, pages_json, text, llm_result_json, created_at) VALUES (?,?,?,?,?,?)",
                    (ingestion_id, idx, _json_dump(ch.get("pages", [])), ch.get("text", ""), None, now),
                )

    return RedirectResponse(url=f"/import/pdf/run?ingestion_id={ingestion_id}&max_tasks={max_tasks}&max_chunks={max_chunks}&lmstudio_base_url={probe.get('base_url')}", status_code=303)


@app.get("/import/pdf/run", response_class=HTMLResponse)
def import_pdf_run(request: Request, ingestion_id: str, max_tasks: int = 10, max_chunks: int = 8, lmstudio_base_url: str = ""):
    require(request.state.role, "import:pdf")
    actor = request.state.user

    probe = _lmstudio_probe(lmstudio_base_url)
    if not probe.get("ok"):
        # Render a friendly error page (instead of raw 502)
        return templates.TemplateResponse(
            request,
            "import_pdf_preview.html",
            {
                "ingestion": {"id": ingestion_id, "cursor_chunk": "?", "filename": "(unknown)", "lmstudio_base_url": str(probe.get("base_url"))},
                "candidates": [],
                "workflows": [],
                "error": f"LM Studio is not reachable at {probe.get('base_url')} ({probe.get('detail')})",
                "done": False,
            },
        )

    max_tasks = max(1, min(int(max_tasks), 10))
    max_chunks = max(1, min(int(max_chunks), 20))

    with db() as conn:
        ing = conn.execute(
            "SELECT * FROM ingestions WHERE id=? AND created_by=?",
            (ingestion_id, actor),
        ).fetchone()
        if not ing:
            raise HTTPException(404)

        cursor = int(ing["cursor_chunk"])
        chunk_rows = conn.execute(
            "SELECT chunk_index, pages_json, text, llm_result_json FROM ingestion_chunks WHERE ingestion_id=? AND chunk_index>=? ORDER BY chunk_index ASC LIMIT ?",
            (ingestion_id, cursor, max_chunks),
        ).fetchall()

        if not chunk_rows:
            return templates.TemplateResponse(
                request,
                "import_pdf_preview.html",
                {
                    "ingestion": dict(ing),
                    "candidates": [],
                    "workflows": [],
                    "error": None,
                    "done": True,
                },
            )

        # Build existing task signatures for dedupe
        latest_rows = conn.execute(
            "SELECT record_id, MAX(version) AS v FROM tasks GROUP BY record_id"
        ).fetchall()
        existing_tasks: list[dict[str, Any]] = []
        for r in latest_rows:
            row = conn.execute(
                "SELECT title, outcome, steps_json FROM tasks WHERE record_id=? AND version=?",
                (r["record_id"], int(r["v"])),
            ).fetchone()
            if not row:
                continue
            existing_tasks.append(
                {
                    "record_id": r["record_id"],
                    "title": row["title"],
                    "outcome": row["outcome"],
                    "steps": _json_load(row["steps_json"]) or [],
                }
            )

    system = {
        "role": "system",
        "content": (
            "You are extracting governed learning Tasks from technical documentation. "
            "You MUST follow the schema strictly. Do not invent steps that are not supported by the provided source. "
            "If uncertain, omit. Every step MUST include a completion check. "
            "Concepts are best-effort and should be minimal."
        ),
    }

    user_prompt_tpl = (
        "From the following SOURCE TEXT (with page markers), propose up to {per_chunk} Task records.\n\n"
        "Rules:\n"
        "- A Task is one atomic outcome.\n"
        "- Provide: title, outcome, facts[], concepts[], dependencies[], procedure_name.\n"
        "- Provide steps[] where each step has: text, completion, and optional actions[].\n"
        "- Steps and completion MUST be concrete and verifiable.\n"
        "- Do NOT include troubleshooting.\n"
        "- Return JSON ONLY: {\"tasks\": [ ... ]} (no markdown, no commentary).\n\n"
        "SOURCE TEXT:\n{source}\n"
    )

    per_chunk = 3

    candidates: list[dict[str, Any]] = []

    # Fail whole run: any chunk failure aborts without advancing cursor.
    try:
        for cr in chunk_rows:
            chunk_index = int(cr["chunk_index"])
            cached = cr["llm_result_json"]

            if cached:
                try:
                    data = json.loads(cached)
                except Exception:
                    data = None
            else:
                user_prompt = user_prompt_tpl.replace("{per_chunk}", str(per_chunk)).replace("{source}", cr["text"])
                raw = _lmstudio_chat(
                    [system, {"role": "user", "content": user_prompt}],
                    temperature=0.2,
                    max_tokens=2000,
                    base_url=str(probe.get("base_url")),
                )
                try:
                    data = json.loads(raw)
                except Exception:
                    raise ValueError(f"Model returned non-JSON for chunk {chunk_index}")

                with db() as conn:
                    conn.execute(
                        "UPDATE ingestion_chunks SET llm_result_json=? WHERE ingestion_id=? AND chunk_index=?",
                        (_json_dump(data), ingestion_id, chunk_index),
                    )

            tasks = data.get("tasks") if isinstance(data, dict) else None
            if not isinstance(tasks, list):
                raise ValueError(f"Model returned invalid schema for chunk {chunk_index}")

            for t in tasks:
                if not isinstance(t, dict):
                    continue
                title = str(t.get("title", "")).strip()
                if not title:
                    continue
                # Keep candidates light for UI: store only what we need now.
                cand = {
                    "chunk_index": chunk_index,
                    "pages": _json_load(cr["pages_json"]) or [],
                    "task": t,
                }
                candidates.append(cand)
    except Exception as e:
        # Render friendly error; do not advance cursor.
        return templates.TemplateResponse(
            request,
            "import_pdf_preview.html",
            {
                "ingestion": {"id": ingestion_id, "cursor_chunk": cursor, "filename": ing["filename"], "lmstudio_base_url": str(probe.get("base_url"))},
                "candidates": [],
                "workflows": [],
                "error": str(e),
                "done": False,
            },
        )

    # Merge + cap to max_tasks
    # De-dupe within candidate list by fingerprint
    out: list[dict[str, Any]] = []
    seen_fp: set[str] = set()
    for c in candidates:
        fp = _task_fingerprint(c["task"])
        if fp in seen_fp:
            continue
        seen_fp.add(fp)
        out.append(c)
        if len(out) >= max_tasks:
            break

    # Attach dup flags
    flagged: list[dict[str, Any]] = []
    for c in out:
        t = c["task"]
        fp = _task_fingerprint(t)
        near_matches: list[dict[str, Any]] = []
        for ex in existing_tasks:
            ex_fp = _task_fingerprint(ex)
            if ex_fp == fp:
                near_matches.append({"record_id": ex["record_id"], "kind": "exact", "score": 1.0})
                continue
            score = _near_duplicate_score(t, ex)
            if score >= 0.72:
                near_matches.append({"record_id": ex["record_id"], "kind": "near", "score": round(score, 3)})
        near_matches = sorted(near_matches, key=lambda x: x["score"], reverse=True)[:3]

        flagged.append(
            {
                "id": _sha256_bytes((fp + str(c["chunk_index"])).encode("utf-8"))[:16],
                "title": str(t.get("title", "")).strip(),
                "chunk_index": c["chunk_index"],
                "pages": c["pages"],
                "dup_matches": near_matches,
            }
        )

    # Propose workflows from candidate titles (optional)
    wf_candidates: list[dict[str, Any]] = []
    if flagged:
        titles = [x["title"] for x in flagged]
        wf_system = {"role": "system", "content": "You propose small Workflows from a list of Task titles. Return JSON only."}
        wf_user = (
            "Given these Task titles, propose up to 3 Workflow candidates.\n"
            "Return JSON ONLY: {\"workflows\": [{\"title\":...,\"objective\":...,\"task_titles\":[...] }]}\n"
            "Rules: a workflow must reference 2-6 tasks by exact title; do not invent titles.\n\n"
            + _json_dump({"task_titles": titles})
        )
        raw = _lmstudio_chat([wf_system, {"role": "user", "content": wf_user}], temperature=0.2, max_tokens=800, base_url=str(probe.get("base_url")))
        data = json.loads(raw)
        wfs = data.get("workflows") if isinstance(data, dict) else None
        if isinstance(wfs, list):
            for wf in wfs[:3]:
                if not isinstance(wf, dict):
                    continue
                wt = str(wf.get("title", "")).strip()
                obj = str(wf.get("objective", "")).strip()
                tts = wf.get("task_titles") or []
                if not wt or not obj or not isinstance(tts, list):
                    continue
                wf_candidates.append(
                    {
                        "id": _sha256_bytes((wt + obj).encode("utf-8"))[:16],
                        "title": wt,
                        "objective": obj,
                        "task_titles": [str(x) for x in tts if str(x).strip() in titles],
                    }
                )

    return templates.TemplateResponse(
        request,
        "import_pdf_preview.html",
        {
            "ingestion": {"id": ingestion_id, "cursor_chunk": int(ing["cursor_chunk"]), "filename": ing["filename"], "lmstudio_base_url": str(probe.get("base_url"))},
            "candidates": flagged,
            "workflows": wf_candidates,
            "error": None,
            "done": False,
        },
    )


@app.post("/import/pdf/commit")
def import_pdf_commit(
    request: Request,
    ingestion_id: str = Form(...),
    candidate_id: list[str] = Form([]),
    workflow_id: list[str] = Form([]),
):
    require(request.state.role, "import:pdf")
    actor = request.state.user

    # Load last run candidates from cached llm results for current cursor window.
    with db() as conn:
        ing = conn.execute("SELECT * FROM ingestions WHERE id=? AND created_by=?", (ingestion_id, actor)).fetchone()
        if not ing:
            raise HTTPException(404)

        cursor = int(ing["cursor_chunk"])
        max_chunks = 8
        chunk_rows = conn.execute(
            "SELECT chunk_index, pages_json, text, llm_result_json FROM ingestion_chunks WHERE ingestion_id=? AND chunk_index>=? ORDER BY chunk_index ASC LIMIT ?",
            (ingestion_id, cursor, max_chunks),
        ).fetchall()

        if not chunk_rows:
            return RedirectResponse(url="/tasks?status=draft", status_code=303)

        # Reconstruct candidates deterministically
        reconstructed: list[dict[str, Any]] = []
        for cr in chunk_rows:
            if not cr["llm_result_json"]:
                continue
            data = json.loads(cr["llm_result_json"])
            tasks = data.get("tasks") if isinstance(data, dict) else []
            for t in tasks:
                if not isinstance(t, dict):
                    continue
                fp = _task_fingerprint(t)
                cid = _sha256_bytes((fp + str(int(cr["chunk_index"]))).encode("utf-8"))[:16]
                if cid not in candidate_id:
                    continue
                reconstructed.append({"task": t, "pages": _json_load(cr["pages_json"]) or []})

        now = utc_now_iso()
        created_tasks: dict[str, tuple[str, int]] = {}  # title -> (record_id, version)

        # Insert selected tasks
        for item in reconstructed:
            t = item["task"]
            title = str(t.get("title", "")).strip()
            outcome = str(t.get("outcome", "")).strip()
            procedure_name = str(t.get("procedure_name", "")).strip() or title
            facts = t.get("facts") or []
            concepts = t.get("concepts") or []
            deps = t.get("dependencies") or []
            steps = t.get("steps") or []

            steps_norm = _normalize_steps(steps)
            _validate_steps_required(steps_norm)

            record_id = str(uuid.uuid4())
            version = 1

            assets = [
                {
                    "url": f"ingestion:{ingestion_id}",
                    "type": "link",
                    "label": f"source_pdf:{ing['filename']} pages:{item['pages']}",
                }
            ]

            conn.execute(
                """
                INSERT INTO tasks(
                  record_id, version, status,
                  title, outcome, facts_json, concepts_json, procedure_name, steps_json, dependencies_json,
                  irreversible_flag, task_assets_json,
                  domain,
                  created_at, updated_at, created_by, updated_by,
                  reviewed_at, reviewed_by, change_note,
                  needs_review_flag, needs_review_note
                ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                """,
                (
                    record_id,
                    version,
                    "draft",
                    title,
                    outcome,
                    _json_dump([str(x) for x in facts]),
                    _json_dump([str(x) for x in concepts]),
                    procedure_name,
                    _json_dump(steps_norm),
                    _json_dump([str(x) for x in deps]),
                    0,
                    _json_dump(assets),
                    "",
                    now,
                    now,
                    actor,
                    actor,
                    None,
                    None,
                    f"import:pdf ingestion={ingestion_id}",
                    1,
                    "AI-imported: check for duplicates and correctness",
                ),
            )
            audit("task", record_id, version, "create", actor, note="import:pdf")
            created_tasks[title] = (record_id, version)

        # Insert workflows selected
        # Recompute workflow candidates from selected titles (best-effort)
        if workflow_id and created_tasks:
            titles = list(created_tasks.keys())
            wf_system = {"role": "system", "content": "You propose small Workflows from a list of Task titles. Return JSON only."}
            wf_user = (
                "Given these Task titles, propose up to 3 Workflow candidates.\n"
                "Return JSON ONLY: {\"workflows\": [{\"id\":...,\"title\":...,\"objective\":...,\"task_titles\":[...] }]}\n"
                "Rules: a workflow must reference 2-6 tasks by exact title; do not invent titles.\n\n"
                + _json_dump({"task_titles": titles})
            )
            raw = _lmstudio_chat([wf_system, {"role": "user", "content": wf_user}], temperature=0.2, max_tokens=900, base_url=str(probe.get("base_url")))
            data = json.loads(raw)
            wfs = data.get("workflows") if isinstance(data, dict) else None
            if isinstance(wfs, list):
                for wf in wfs:
                    if not isinstance(wf, dict):
                        continue
                    wid = str(wf.get("id", "")).strip() or _sha256_bytes((str(wf.get("title",""))+str(wf.get("objective",""))).encode("utf-8"))[:16]
                    if wid not in workflow_id:
                        continue
                    title = str(wf.get("title", "")).strip()
                    objective = str(wf.get("objective", "")).strip()
                    tts = wf.get("task_titles") or []
                    if not title or not objective or not isinstance(tts, list):
                        continue

                    wf_rid = str(uuid.uuid4())
                    wf_ver = 1
                    conn.execute(
                        "INSERT INTO workflows(record_id, version, status, title, objective, domains_json, tags_json, meta_json, created_at, updated_at, created_by, updated_by, reviewed_at, reviewed_by, change_note, needs_review_flag, needs_review_note) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                        (wf_rid, wf_ver, "draft", title, objective, "[]", "[]", "{}", now, now, actor, actor, None, None, f"import:pdf ingestion={ingestion_id}", 1, "AI-imported: check composition"),
                    )

                    order = 1
                    for tt in [str(x) for x in tts if str(x) in created_tasks]:
                        tr, tv = created_tasks[tt]
                        conn.execute(
                            "INSERT INTO workflow_task_refs(workflow_record_id, workflow_version, order_index, task_record_id, task_version) VALUES (?,?,?,?,?)",
                            (wf_rid, wf_ver, order, tr, tv),
                        )
                        order += 1

                    audit("workflow", wf_rid, wf_ver, "create", actor, note="import:pdf")

        # Advance cursor if commit happened (clean, deterministic)
        if reconstructed:
            conn.execute(
                "UPDATE ingestions SET cursor_chunk=cursor_chunk+? , status='in_progress' WHERE id=?",
                (len(chunk_rows), ingestion_id),
            )

    return RedirectResponse(url="/tasks?status=draft", status_code=303)


@app.get("/import/json", response_class=HTMLResponse)
def import_json_form(request: Request):
    require(request.state.role, "import:json")
    return templates.TemplateResponse(request, "import_json.html", {})


def _parse_task_json(obj: dict[str, Any]) -> dict[str, Any]:
    title = str(obj.get("title", "")).strip()
    outcome = str(obj.get("outcome", "")).strip()
    procedure_name = str(obj.get("procedure_name", "")).strip() or title
    if not title:
        raise HTTPException(status_code=400, detail="Task import: title is required")
    if not outcome:
        raise HTTPException(status_code=400, detail=f"Task import '{title}': outcome is required")

    facts = obj.get("facts") or []
    concepts = obj.get("concepts") or []
    deps = obj.get("dependencies") or []
    steps = obj.get("steps") or []

    if not isinstance(facts, list) or not isinstance(concepts, list) or not isinstance(deps, list):
        raise HTTPException(status_code=400, detail=f"Task import '{title}': facts/concepts/dependencies must be lists")

    steps_norm = _normalize_steps(steps)
    _validate_steps_required(steps_norm)

    irreversible_flag = 1 if bool(obj.get("irreversible_flag")) else 0
    assets = obj.get("task_assets") or obj.get("assets") or []
    if not isinstance(assets, list):
        raise HTTPException(status_code=400, detail=f"Task import '{title}': task_assets must be a list")

    return {
        "record_id": str(obj.get("record_id") or "").strip() or str(uuid.uuid4()),
        "version": int(obj.get("version") or 1),
        # Import is ingress: always draft. Trust boundary is human review.
        "status": "draft",
        "title": title,
        "outcome": outcome,
        "procedure_name": procedure_name,
        "facts": [str(x) for x in facts],
        "concepts": [str(x) for x in concepts],
        "dependencies": [str(x) for x in deps],
        "steps": steps_norm,
        "irreversible_flag": irreversible_flag,
        "task_assets": assets,
        "needs_review_flag": 1 if bool(obj.get("needs_review_flag")) else 0,
        "needs_review_note": (str(obj.get("needs_review_note")) if obj.get("needs_review_note") is not None else None),
    }


def _parse_workflow_json(obj: dict[str, Any]) -> dict[str, Any]:
    title = str(obj.get("title", "")).strip()
    objective = str(obj.get("objective", "")).strip()
    if not title:
        raise HTTPException(status_code=400, detail="Workflow import: title is required")
    if not objective:
        raise HTTPException(status_code=400, detail=f"Workflow import '{title}': objective is required")

    raw_refs = obj.get("task_refs") or obj.get("tasks") or []
    refs: list[tuple[str, int]] = []

    if isinstance(raw_refs, list):
        for item in raw_refs:
            if isinstance(item, str):
                if "@" not in item:
                    raise HTTPException(status_code=400, detail=f"Workflow import '{title}': invalid task ref '{item}'")
                rid, ver = item.split("@", 1)
                refs.append((rid.strip(), int(ver.strip())))
            elif isinstance(item, dict):
                rid = str(item.get("record_id") or item.get("task_record_id") or "").strip()
                ver = item.get("version") or item.get("task_version")
                if not rid or ver is None:
                    raise HTTPException(status_code=400, detail=f"Workflow import '{title}': task_refs items require record_id + version")
                refs.append((rid, int(ver)))
            else:
                raise HTTPException(status_code=400, detail=f"Workflow import '{title}': task_refs must contain strings or objects")
    else:
        raise HTTPException(status_code=400, detail=f"Workflow import '{title}': task_refs must be a list")

    if not refs:
        raise HTTPException(status_code=400, detail=f"Workflow import '{title}': at least one task_ref is required")

    return {
        "record_id": str(obj.get("record_id") or "").strip() or str(uuid.uuid4()),
        "version": int(obj.get("version") or 1),
        # Import is ingress: always draft. Trust boundary is human review.
        "status": "draft",
        "title": title,
        "objective": objective,
        "refs": refs,
        "needs_review_flag": 1 if bool(obj.get("needs_review_flag")) else 0,
        "needs_review_note": (str(obj.get("needs_review_note")) if obj.get("needs_review_note") is not None else None),
    }


@app.post("/import/json")
def import_json_run(
    request: Request,
    upload: UploadFile = File(...),
    actor_note: str = Form("Imported from JSON"),
):
    require(request.state.role, "import:json")
    actor = request.state.user

    raw = upload.file.read()
    try:
        payload = json.loads(raw)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON: {e}")

    tasks_in: list[dict[str, Any]] = []
    workflows_in: list[dict[str, Any]] = []

    if isinstance(payload, dict):
        if isinstance(payload.get("tasks"), list):
            tasks_in = [x for x in payload.get("tasks") if isinstance(x, dict)]
        if isinstance(payload.get("workflows"), list):
            workflows_in = [x for x in payload.get("workflows") if isinstance(x, dict)]
        # Allow single objects
        if payload.get("type") == "task":
            tasks_in = [payload]
        if payload.get("type") == "workflow":
            workflows_in = [payload]
    elif isinstance(payload, list):
        # list of heterogeneous objects
        for x in payload:
            if not isinstance(x, dict):
                continue
            if x.get("type") == "workflow":
                workflows_in.append(x)
            else:
                # default to task
                tasks_in.append(x)
    else:
        raise HTTPException(status_code=400, detail="Import JSON must be an object or a list")

    if not tasks_in and not workflows_in:
        raise HTTPException(status_code=400, detail="No tasks/workflows found in uploaded JSON")

    created_task_ids: list[str] = []
    created_workflow_ids: list[str] = []
    now = utc_now_iso()

    with db() as conn:
        # tasks first
        for t in tasks_in:
            item = _parse_task_json(t)
            # Import is ingress: always draft.
            # (Seeding/demo data should write directly to the DB via seed scripts, not via import.)
            item["status"] = "draft"

            # Prevent overwrite
            exists = conn.execute(
                "SELECT 1 FROM tasks WHERE record_id=? AND version=?",
                (item["record_id"], item["version"]),
            ).fetchone()
            if exists:
                raise HTTPException(
                    status_code=409,
                    detail=f"Task import conflict: {item['record_id']}@{item['version']} already exists",
                )

            conn.execute(
                """
                INSERT INTO tasks(
                  record_id, version, status,
                  title, outcome, facts_json, concepts_json, procedure_name, steps_json, dependencies_json,
                  irreversible_flag, task_assets_json,
                  domain,
                  created_at, updated_at, created_by, updated_by,
                  reviewed_at, reviewed_by, change_note,
                  needs_review_flag, needs_review_note
                ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                """,
                (
                    item["record_id"],
                    item["version"],
                    item["status"],
                    item["title"],
                    item["outcome"],
                    _json_dump(item["facts"]),
                    _json_dump(item["concepts"]),
                    item["procedure_name"],
                    _json_dump(item["steps"]),
                    _json_dump(item["dependencies"]),
                    item["irreversible_flag"],
                    _json_dump(item["task_assets"]),
                    "",
                    now,
                    now,
                    actor,
                    actor,
                    None,
                    None,
                    actor_note.strip() or "Imported from JSON",
                    item["needs_review_flag"],
                    item["needs_review_note"],
                ),
            )
            audit("task", item["record_id"], item["version"], "create", actor, note="import:json")
            created_task_ids.append(item["record_id"])

        # workflows
        for w in workflows_in:
            item = _parse_workflow_json(w)
            # Import is ingress: always draft.
            # (Seeding/demo data should write directly to the DB via seed scripts, not via import.)
            item["status"] = "draft"

            exists = conn.execute(
                "SELECT 1 FROM workflows WHERE record_id=? AND version=?",
                (item["record_id"], item["version"]),
            ).fetchone()
            if exists:
                raise HTTPException(
                    status_code=409,
                    detail=f"Workflow import conflict: {item['record_id']}@{item['version']} already exists",
                )

            enforce_workflow_ref_rules(conn, item["refs"])
            # Imported workflows always arrive as draft; confirmation remains a human-only trust boundary.

            conn.execute(
                """
                INSERT INTO workflows(
                  record_id, version, status,
                  title, objective,
                  domains_json,
                  tags_json, meta_json,
                  created_at, updated_at, created_by, updated_by,
                  reviewed_at, reviewed_by, change_note,
                  needs_review_flag, needs_review_note
                ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                """,
                (
                    item["record_id"],
                    item["version"],
                    item["status"],
                    item["title"],
                    item["objective"],
                    _json_dump(_workflow_domains(conn, item["refs"])),
                    "[]",
                    "{}",
                    now,
                    now,
                    actor,
                    actor,
                    None,
                    None,
                    actor_note.strip() or "Imported from JSON",
                    item["needs_review_flag"],
                    item["needs_review_note"],
                ),
            )
            for idx, (rid, ver) in enumerate(item["refs"], start=1):
                conn.execute(
                    """
                    INSERT INTO workflow_task_refs(workflow_record_id, workflow_version, order_index, task_record_id, task_version)
                    VALUES (?,?,?,?,?)
                    """,
                    (item["record_id"], item["version"], idx, rid, ver),
                )

            audit("workflow", item["record_id"], item["version"], "create", actor, note="import:json")
            created_workflow_ids.append(item["record_id"])

    # Redirect to something sensible
    if created_workflow_ids and not created_task_ids:
        return RedirectResponse(url="/workflows", status_code=303)
    return RedirectResponse(url="/tasks?status=draft", status_code=303)


# ---- Tasks ----


@app.get("/tasks", response_class=HTMLResponse)
def tasks_list(request: Request, status: str | None = None, q: str | None = None, domain: str | None = None, tag: str | None = None):
    q_norm = (q or "").strip().lower()
    domain_norm = (domain or "").strip().lower() or None
    tag_norm = (tag or "").strip().lower() or None

    with db() as conn:
        sql = "SELECT record_id, MAX(version) AS latest_version FROM tasks GROUP BY record_id ORDER BY record_id"
        rows = conn.execute(sql).fetchall()

        # Option lists (for dropdown filters)
        domains = _active_domains(conn)
        all_tags: set[str] = set()

        items = []
        for r in rows:
            rid = r["record_id"]
            latest_v = int(r["latest_version"])
            latest = conn.execute(
                "SELECT * FROM tasks WHERE record_id=? AND version=?", (rid, latest_v)
            ).fetchone()
            if not latest:
                continue
            if status and latest["status"] != status:
                continue
            if q_norm and q_norm not in (latest["title"] or "").lower():
                continue

            # derived: update_pending_confirmation
            confirmed_v = conn.execute(
                "SELECT MAX(version) AS v FROM tasks WHERE record_id=? AND status='confirmed'",
                (rid,),
            ).fetchone()["v"]
            update_pending = False
            if confirmed_v is not None and latest_v > int(confirmed_v) and latest["status"] in ("draft", "submitted"):
                update_pending = True

            tags = [str(x).strip().lower() for x in (_json_load(latest["tags_json"]) if "tags_json" in latest.keys() else [])]
            domain_val = (latest["domain"] if "domain" in latest.keys() else "")

            for t in tags:
                if t:
                    all_tags.add(t)

            # Apply filters
            if domain_norm and (domain_val or "").strip().lower() != domain_norm:
                continue
            if tag_norm and tag_norm not in set(tags):
                continue

            # Returned-for-changes signal: does this version have a return note?
            has_return_note = False
            if latest["status"] == "returned":
                rn = conn.execute(
                    "SELECT 1 FROM audit_log WHERE entity_type='task' AND record_id=? AND version=? AND action='return_for_changes' LIMIT 1",
                    (rid, latest_v),
                ).fetchone()
                has_return_note = bool(rn)

            items.append(
                {
                    "record_id": rid,
                    "latest_version": latest_v,
                    "title": latest["title"],
                    "status": latest["status"],
                    "needs_review_flag": bool(latest["needs_review_flag"]),
                    "update_pending_confirmation": update_pending,
                    "tags": tags,
                    "domain": domain_val,
                    "has_return_note": has_return_note,
                }
            )

    return templates.TemplateResponse(
        request,
        "tasks_list.html",
        {
            "items": items,
            "status": status,
            "q": q,
            "domain": domain_norm or "",
            "tag": tag_norm or "",
            "domains": domains,
            "tags": sorted(all_tags),
        },
    )


@app.get("/tasks/new", response_class=HTMLResponse)
def task_new_form(request: Request):
    require(request.state.role, "task:create")
    with db() as conn:
        domains = _active_domains(conn)
    return templates.TemplateResponse(
        request,
        "task_edit.html",
        {"mode": "new", "task": None, "warnings": [], "domains": domains},
    )


@app.post("/tasks/new")
def task_create(
    request: Request,
    title: str = Form(...),
    outcome: str = Form(...),
    procedure_name: str = Form(...),
    domain: str = Form(""),
    tags: str = Form(""),
    meta: str = Form(""),
    facts: str = Form(""),
    concepts: str = Form(""),
    dependencies: str = Form(""),
    step_text: list[str] = Form([]),
    step_completion: list[str] = Form([]),
    step_actions: list[str] = Form([]),
    irreversible_flag: bool = Form(False),
):
    require(request.state.role, "task:create")
    actor = request.state.user
    record_id = str(uuid.uuid4())
    version = 1

    facts_list = parse_lines(facts)
    concepts_list = parse_lines(concepts)
    deps_list = parse_lines(dependencies)
    tags_list = parse_tags(tags)
    meta_obj = parse_meta(meta)
    steps_list = _zip_steps(step_text, step_completion, step_actions)
    _validate_steps_required(steps_list)

    warnings = lint_steps(steps_list)

    now = utc_now_iso()
    with db() as conn:
        domains = _active_domains(conn)
        domain_norm = (domain or "").strip().lower()
        if domain_norm and domain_norm not in domains:
            raise HTTPException(status_code=400, detail=f"Invalid domain '{domain_norm}'")

        conn.execute(
            """
            INSERT INTO tasks(
              record_id, version, status,
              title, outcome, facts_json, concepts_json, procedure_name, steps_json, dependencies_json,
              irreversible_flag, task_assets_json,
              domain,
              tags_json, meta_json,
              created_at, updated_at, created_by, updated_by,
              reviewed_at, reviewed_by, change_note,
              needs_review_flag, needs_review_note
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (
                record_id,
                version,
                "draft",
                title.strip(),
                outcome.strip(),
                _json_dump(facts_list),
                _json_dump(concepts_list),
                procedure_name.strip(),
                _json_dump(steps_list),
                _json_dump(deps_list),
                1 if irreversible_flag else 0,
                _json_dump([]),
                domain_norm,
                _json_dump(tags_list),
                _json_dump(meta_obj),
                now,
                now,
                actor,
                actor,
                None,
                None,
                None,
                0,
                None,
            ),
        )
    audit("task", record_id, version, "create", actor)

    # Show edit page with warnings banner
    return RedirectResponse(url=f"/tasks/{record_id}/{version}/edit?created=1", status_code=303)


@app.get("/tasks/{record_id}/{version}", response_class=HTMLResponse)
def task_view(request: Request, record_id: str, version: int):
    with db() as conn:
        row = conn.execute(
            "SELECT * FROM tasks WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
    if not row:
        raise HTTPException(404)

    task = dict(row)
    task["facts"] = _json_load(task["facts_json"])
    task["concepts"] = _json_load(task["concepts_json"])
    task["dependencies"] = _json_load(task["dependencies_json"])

    raw_steps = _json_load(task["steps_json"])
    task["steps"] = _normalize_steps(raw_steps)
    for st in task["steps"]:
        if "actions" not in st or st["actions"] is None:
            st["actions"] = []

    warnings = lint_steps(task["steps"])

    # If returned, surface the most recent return note (if any)
    return_note = None
    if task.get("status") == "returned":
        with db() as conn:
            rn = conn.execute(
                "SELECT note, at, actor FROM audit_log WHERE entity_type='task' AND record_id=? AND version=? AND action='return_for_changes' ORDER BY at DESC LIMIT 1",
                (record_id, version),
            ).fetchone()
            if rn and rn["note"]:
                return_note = {"note": rn["note"], "at": rn["at"], "actor": rn["actor"]}

    return templates.TemplateResponse(
        request,
        "task_view.html",
        {"task": task, "warnings": warnings, "return_note": return_note},
    )


@app.get("/tasks/{record_id}/{version}/edit", response_class=HTMLResponse)
def task_edit_form(request: Request, record_id: str, version: int):
    require(request.state.role, "task:revise")
    with db() as conn:
        row = conn.execute(
            "SELECT * FROM tasks WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
    if not row:
        raise HTTPException(404)

    task = dict(row)
    task["facts"] = "\n".join(_json_load(task["facts_json"]) or [])
    task["concepts"] = "\n".join(_json_load(task["concepts_json"]) or [])
    task["dependencies"] = "\n".join(_json_load(task["dependencies_json"]) or [])

    raw_steps = _json_load(task["steps_json"])
    task["steps"] = _normalize_steps(raw_steps)
    for st in task["steps"]:
        if "actions" not in st or st["actions"] is None:
            st["actions"] = []

    warnings = lint_steps(task["steps"])

    with db() as conn:
        domains = _active_domains(conn)

    return templates.TemplateResponse(
        request,
        "task_edit.html",
        {"mode": "edit", "task": task, "warnings": warnings, "domains": domains},
    )


@app.post("/tasks/{record_id}/{version}/save")
def task_save(
    request: Request,
    record_id: str,
    version: int,
    title: str = Form(...),
    outcome: str = Form(...),
    procedure_name: str = Form(...),
    domain: str = Form(""),
    tags: str = Form(""),
    meta: str = Form(""),
    facts: str = Form(""),
    concepts: str = Form(""),
    dependencies: str = Form(""),
    step_text: list[str] = Form([]),
    step_completion: list[str] = Form([]),
    step_actions: list[str] = Form([]),
    irreversible_flag: bool = Form(False),
    change_note: str = Form(""),
):
    """Records are immutable.

    Saving changes always creates a NEW VERSION (draft) with a required change_note.
    """
    require(request.state.role, "task:revise")
    actor = request.state.user

    facts_list = parse_lines(facts)
    concepts_list = parse_lines(concepts)
    deps_list = parse_lines(dependencies)
    tags_list = parse_tags(tags)
    meta_obj = parse_meta(meta)
    steps_list = _zip_steps(step_text, step_completion, step_actions)
    _validate_steps_required(steps_list)

    note = change_note.strip()
    if not note:
        raise HTTPException(status_code=400, detail="change_note is required when creating a new version")

    with db() as conn:
        src = conn.execute(
            "SELECT * FROM tasks WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not src:
            raise HTTPException(404)

        # If the source version was returned for changes, force the change_note to reference the return note.
        if src["status"] == "returned":
            rn = conn.execute(
                "SELECT note, at, actor FROM audit_log WHERE entity_type='task' AND record_id=? AND version=? AND action='return_for_changes' ORDER BY at DESC LIMIT 1",
                (record_id, version),
            ).fetchone()
            if rn and rn["note"]:
                prefix = f"Response to return note by {rn['actor']} at {rn['at']}: {rn['note']} | "
                if prefix not in note:
                    note = prefix + note

        # New version number is latest + 1
        latest_v = get_latest_version(conn, "tasks", record_id) or version
        new_v = latest_v + 1
        now = utc_now_iso()

        conn.execute(
            """
            INSERT INTO tasks(
              record_id, version, status,
              title, outcome, facts_json, concepts_json, procedure_name, steps_json, dependencies_json,
              irreversible_flag, task_assets_json,
              domain,
              tags_json, meta_json,
              created_at, updated_at, created_by, updated_by,
              reviewed_at, reviewed_by, change_note,
              needs_review_flag, needs_review_note
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (
                record_id,
                new_v,
                "draft",
                title.strip(),
                outcome.strip(),
                _json_dump(facts_list),
                _json_dump(concepts_list),
                procedure_name.strip(),
                _json_dump(steps_list),
                _json_dump(deps_list),
                1 if irreversible_flag else 0,
                src["task_assets_json"],
                (domain or "").strip().lower(),
                _json_dump(tags_list),
                _json_dump(meta_obj),
                now,
                now,
                actor,
                actor,
                None,
                None,
                note,
                int(src["needs_review_flag"]),
                src["needs_review_note"],
            ),
        )

    audit("task", record_id, new_v, "new_version", actor, note=f"from v{version}: {note}")
    return RedirectResponse(url=f"/tasks/{record_id}/{new_v}", status_code=303)


@app.post("/tasks/{record_id}/{version}/new-version")
def task_new_version(request: Request, record_id: str, version: int):
    require(request.state.role, "task:revise")
    actor = request.state.user
    with db() as conn:
        src = conn.execute(
            "SELECT * FROM tasks WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not src:
            raise HTTPException(404)
        latest_v = get_latest_version(conn, "tasks", record_id) or version
        new_v = latest_v + 1
        now = utc_now_iso()

        conn.execute(
            """
            INSERT INTO tasks(
              record_id, version, status,
              title, outcome, facts_json, concepts_json, procedure_name, steps_json, dependencies_json,
              irreversible_flag, task_assets_json,
              domain,
              tags_json, meta_json,
              created_at, updated_at, created_by, updated_by,
              reviewed_at, reviewed_by, change_note,
              needs_review_flag, needs_review_note
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (
                record_id,
                new_v,
                "draft",
                src["title"],
                src["outcome"],
                src["facts_json"],
                src["concepts_json"],
                src["procedure_name"],
                src["steps_json"],
                src["dependencies_json"],
                src["irreversible_flag"],
                src["task_assets_json"],
                (src["domain"] if "domain" in src.keys() else ""),
                (src["tags_json"] if "tags_json" in src.keys() else "[]"),
                (src["meta_json"] if "meta_json" in src.keys() else "{}"),
                now,
                now,
                actor,
                actor,
                None,
                None,
                f"Created new version from v{version}",
                int(src["needs_review_flag"]),
                src["needs_review_note"],
            ),
        )

    audit("task", record_id, new_v, "new_version", actor, note=f"from v{version}")
    return RedirectResponse(url=f"/tasks/{record_id}/{new_v}/edit", status_code=303)


@app.post("/tasks/{record_id}/{version}/submit")
def task_submit(request: Request, record_id: str, version: int):
    require(request.state.role, "task:submit")
    actor = request.state.user
    with db() as conn:
        row = conn.execute(
            "SELECT status, domain FROM tasks WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not row:
            raise HTTPException(404)
        if row["status"] != "draft":
            raise HTTPException(409, detail="Only draft tasks can be submitted")

        domain = (row["domain"] or "").strip()
        if not domain:
            raise HTTPException(status_code=409, detail="Cannot submit task: domain is required")
        if not _user_has_domain(conn, actor, domain):
            raise HTTPException(status_code=403, detail=f"Forbidden: you are not authorized for domain '{domain}'")

        conn.execute(
            "UPDATE tasks SET status='submitted', updated_at=?, updated_by=? WHERE record_id=? AND version=?",
            (utc_now_iso(), actor, record_id, version),
        )
    audit("task", record_id, version, "submit", actor)
    return RedirectResponse(url=f"/tasks/{record_id}/{version}", status_code=303)


@app.post("/tasks/{record_id}/{version}/force-submit")
def task_force_submit(request: Request, record_id: str, version: int):
    require(request.state.role, "task:force_submit")
    actor = request.state.user
    with db() as conn:
        row = conn.execute(
            "SELECT status FROM tasks WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not row:
            raise HTTPException(404)
        if row["status"] in ("deprecated", "confirmed"):
            raise HTTPException(409, detail=f"Cannot force-submit a {row['status']} task")
        conn.execute(
            "UPDATE tasks SET status='submitted', updated_at=?, updated_by=? WHERE record_id=? AND version=?",
            (utc_now_iso(), actor, record_id, version),
        )
    audit("task", record_id, version, "force_submit", actor, note="admin forced submission")
    return RedirectResponse(url=f"/tasks/{record_id}/{version}", status_code=303)


@app.post("/tasks/{record_id}/{version}/return")
def task_return_for_changes(request: Request, record_id: str, version: int, note: str = Form("")):
    require(request.state.role, "task:confirm")
    actor = request.state.user
    msg = (note or "").strip()
    if not msg:
        raise HTTPException(status_code=400, detail="Return note is required")

    with db() as conn:
        row = conn.execute(
            "SELECT status, domain FROM tasks WHERE record_id=? AND version=?",
            (record_id, version),
        ).fetchone()
        if not row:
            raise HTTPException(404)
        if row["status"] != "submitted":
            raise HTTPException(status_code=409, detail="Only submitted tasks can be returned")

        # Domain gate (admin implicitly authorized via _user_has_domain)
        domain = (row["domain"] or "").strip()
        if domain and not _user_has_domain(conn, actor, domain):
            raise HTTPException(status_code=403, detail=f"Forbidden: you are not authorized for domain '{domain}'")

        conn.execute(
            "UPDATE tasks SET status='returned', updated_at=?, updated_by=? WHERE record_id=? AND version=?",
            (utc_now_iso(), actor, record_id, version),
        )

    audit("task", record_id, version, "return_for_changes", actor, note=msg)
    return RedirectResponse(url=f"/tasks/{record_id}/{version}", status_code=303)


@app.post("/tasks/{record_id}/{version}/confirm")
def task_confirm(request: Request, record_id: str, version: int):
    require(request.state.role, "task:confirm")
    actor = request.state.user
    with db() as conn:
        row = conn.execute(
            "SELECT status, domain FROM tasks WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not row:
            raise HTTPException(404)
        if row["status"] != "submitted":
            raise HTTPException(409, detail="Only submitted tasks can be confirmed")

        domain = (row["domain"] or "").strip()
        if not domain:
            raise HTTPException(status_code=409, detail="Cannot confirm task: domain is required")
        if not _user_has_domain(conn, actor, domain):
            raise HTTPException(status_code=403, detail=f"Forbidden: you are not authorized to confirm domain '{domain}'")

        # Deprecate any previously confirmed version
        conn.execute(
            "UPDATE tasks SET status='deprecated', updated_at=?, updated_by=? WHERE record_id=? AND status='confirmed'",
            (utc_now_iso(), actor, record_id),
        )

        conn.execute(
            """
            UPDATE tasks
            SET status='confirmed', reviewed_at=?, reviewed_by=?, updated_at=?, updated_by=?
            WHERE record_id=? AND version=?
            """,
            (utc_now_iso(), actor, utc_now_iso(), actor, record_id, version),
        )

    audit("task", record_id, version, "confirm", actor)
    return RedirectResponse(url=f"/tasks/{record_id}/{version}", status_code=303)


@app.post("/tasks/{record_id}/{version}/force-confirm")
def task_force_confirm(request: Request, record_id: str, version: int):
    require(request.state.role, "task:force_confirm")
    actor = request.state.user
    with db() as conn:
        row = conn.execute(
            "SELECT status FROM tasks WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not row:
            raise HTTPException(404)
        if row["status"] == "deprecated":
            raise HTTPException(409, detail="Cannot force-confirm a deprecated task")

        # Still enforce: you can't confirm an empty/bad record (structure checks are enforced earlier).
        # Admin override is for lifecycle, not semantics.

        # Deprecate any previously confirmed version
        conn.execute(
            "UPDATE tasks SET status='deprecated', updated_at=?, updated_by=? WHERE record_id=? AND status='confirmed'",
            (utc_now_iso(), actor, record_id),
        )

        conn.execute(
            """
            UPDATE tasks
            SET status='confirmed', reviewed_at=?, reviewed_by=?, updated_at=?, updated_by=?
            WHERE record_id=? AND version=?
            """,
            (utc_now_iso(), actor, utc_now_iso(), actor, record_id, version),
        )

    audit("task", record_id, version, "force_confirm", actor, note="admin forced confirmation")
    return RedirectResponse(url=f"/tasks/{record_id}/{version}", status_code=303)


# ---- Workflows ----


@app.get("/workflows", response_class=HTMLResponse)
def workflows_list(request: Request, status: str | None = None, q: str | None = None, domain: str | None = None, tag: str | None = None):
    q_norm = (q or "").strip().lower()
    domain_norm = (domain or "").strip().lower() or None
    tag_norm = (tag or "").strip().lower() or None

    with db() as conn:
        sql = "SELECT record_id, MAX(version) AS latest_version FROM workflows GROUP BY record_id ORDER BY record_id"
        rows = conn.execute(sql).fetchall()

        domains = _active_domains(conn)
        all_tags: set[str] = set()

        items = []
        for r in rows:
            rid = r["record_id"]
            latest_v = int(r["latest_version"])
            latest = conn.execute(
                "SELECT * FROM workflows WHERE record_id=? AND version=?", (rid, latest_v)
            ).fetchone()
            if not latest:
                continue
            if status and latest["status"] != status:
                continue
            if q_norm and q_norm not in (latest["title"] or "").lower():
                continue

            # Derived readiness + domains
            refs = conn.execute(
                "SELECT task_record_id, task_version FROM workflow_task_refs WHERE workflow_record_id=? AND workflow_version=? ORDER BY order_index",
                (rid, latest_v),
            ).fetchall()
            pairs = [(x["task_record_id"], int(x["task_version"])) for x in refs]
            readiness = workflow_readiness(conn, pairs)
            doms = _workflow_domains(conn, pairs)

            tags = [str(x).strip().lower() for x in (_json_load(latest["tags_json"]) if "tags_json" in latest.keys() else [])]
            for t in tags:
                if t:
                    all_tags.add(t)

            # Apply filters
            if domain_norm and domain_norm not in set([d.lower() for d in doms]):
                continue
            if tag_norm and tag_norm not in set(tags):
                continue

            # Store domains_json opportunistically (keeps DB queryable and consistent)
            if "domains_json" in latest.keys():
                conn.execute(
                    "UPDATE workflows SET domains_json=? WHERE record_id=? AND version=?",
                    (_json_dump(doms), rid, latest_v),
                )

            items.append(
                {
                    "record_id": rid,
                    "latest_version": latest_v,
                    "title": latest["title"],
                    "status": latest["status"],
                    "readiness": readiness,
                    "domains": doms,
                    "tags": tags,
                }
            )

    return templates.TemplateResponse(
        request,
        "workflows_list.html",
        {"items": items, "status": status, "q": q, "domain": domain_norm or "", "tag": tag_norm or "", "domains": domains, "tags": sorted(all_tags)},
    )


@app.get("/workflows/new", response_class=HTMLResponse)
def workflow_new_form(request: Request):
    require(request.state.role, "workflow:create")
    with db() as conn:
        confirmed_tasks = conn.execute(
            "SELECT record_id, version, title FROM tasks WHERE status='confirmed' ORDER BY title"
        ).fetchall()
    return templates.TemplateResponse(
        request,
        "workflow_edit.html",
        {"mode": "new", "workflow": None, "confirmed_tasks": confirmed_tasks, "refs_text": ""},
    )


def _parse_task_refs(task_refs: str) -> list[tuple[str, int]]:
    """Parse newline-separated: record_id@version."""
    refs: list[tuple[str, int]] = []
    for ln in parse_lines(task_refs):
        if "@" not in ln:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid task reference '{ln}'. Use the format record_id@version, one per line.",
            )
        rid, ver = ln.split("@", 1)
        try:
            refs.append((rid.strip(), int(ver.strip())))
        except ValueError:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid task version in '{ln}'. Use the format record_id@version.",
            )
    if not refs:
        raise HTTPException(status_code=400, detail="Workflow must include at least one confirmed Task reference")
    return refs


def workflow_readiness(conn: sqlite3.Connection, refs: list[tuple[str, int]]) -> Literal[
    "ready", "awaiting_task_confirmation", "invalid"
]:
    # Backward-compatible: keep readiness as a simple derived label.
    info = workflow_readiness_detail(conn, refs)
    return info["readiness"]


def workflow_readiness_detail(conn: sqlite3.Connection, refs: list[tuple[str, int]]) -> dict[str, Any]:
    """Derived readiness + human-readable reasons.

    Rules:
      - If any reference is missing => invalid
      - If any reference is deprecated => invalid
      - Else if any reference is not confirmed (draft/submitted) => awaiting_task_confirmation
      - Else => ready

    Returns:
      { readiness: str, reasons: [str], blocking_task_refs: [(rid,ver,status)] }
    """
    reasons: list[str] = []
    blocking: list[tuple[str, int, str]] = []

    if not refs:
        return {
            "readiness": "invalid",
            "reasons": ["Workflow has no Task references."],
            "blocking_task_refs": [],
        }

    awaiting = False

    for rid, ver in refs:
        row = conn.execute(
            "SELECT status, title FROM tasks WHERE record_id=? AND version=?", (rid, ver)
        ).fetchone()
        if not row:
            reasons.append(f"Missing Task reference: {rid}@{ver} does not exist")
            return {"readiness": "invalid", "reasons": reasons, "blocking_task_refs": blocking}

        st = row["status"]
        if st == "deprecated":
            reasons.append(f"Deprecated Task reference: {rid}@{ver} is deprecated")
            return {"readiness": "invalid", "reasons": reasons, "blocking_task_refs": blocking}

        if st != "confirmed":
            awaiting = True
            blocking.append((rid, int(ver), str(st)))

    if awaiting:
        reasons.append("One or more referenced Task versions are not confirmed.")
        return {"readiness": "awaiting_task_confirmation", "reasons": reasons, "blocking_task_refs": blocking}

    return {"readiness": "ready", "reasons": [], "blocking_task_refs": []}


def enforce_workflow_ref_rules(conn: sqlite3.Connection, refs: list[tuple[str, int]]) -> None:
    """Hard constraints for workflow composition.

    Draft/submitted workflows may reference draft/submitted/confirmed Task versions.
    Confirmed workflows must reference confirmed Task versions only (enforced at confirm-time).

    Hard constraints here:
      - at least one task reference must exist
      - referenced task versions must exist
      - referenced task versions must not be deprecated
    """
    if not refs:
        raise HTTPException(status_code=400, detail="Workflow must include at least one Task reference")

    derived = workflow_readiness(conn, refs)
    if derived == "invalid":
        raise HTTPException(
            status_code=409,
            detail="Workflow contains invalid Task references (missing or deprecated task versions)",
        )


@app.post("/workflows/new")
def workflow_create(
    request: Request,
    title: str = Form(...),
    objective: str = Form(...),
    tags: str = Form(""),
    meta: str = Form(""),
    task_refs: str = Form(""),
):
    require(request.state.role, "workflow:create")
    actor = request.state.user
    record_id = str(uuid.uuid4())
    version = 1
    now = utc_now_iso()

    refs = _parse_task_refs(task_refs)
    tags_list = parse_tags(tags)
    meta_obj = parse_meta(meta)

    with db() as conn:
        enforce_workflow_ref_rules(conn, refs)

        conn.execute(
            """
            INSERT INTO workflows(
              record_id, version, status,
              title, objective,
              domains_json,
              tags_json, meta_json,
              created_at, updated_at, created_by, updated_by,
              reviewed_at, reviewed_by, change_note,
              needs_review_flag, needs_review_note
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (
                record_id,
                version,
                "draft",
                title.strip(),
                objective.strip(),
                _json_dump(_workflow_domains(conn, refs)),
                _json_dump(tags_list),
                _json_dump(meta_obj),
                now,
                now,
                actor,
                actor,
                None,
                None,
                None,
                0,
                None,
            ),
        )
        for idx, (rid, ver) in enumerate(refs, start=1):
            conn.execute(
                """
                INSERT INTO workflow_task_refs(workflow_record_id, workflow_version, order_index, task_record_id, task_version)
                VALUES (?,?,?,?,?)
                """,
                (record_id, version, idx, rid, ver),
            )

    audit("workflow", record_id, version, "create", actor)
    return RedirectResponse(url=f"/workflows/{record_id}/{version}", status_code=303)


@app.get("/workflows/{record_id}/{version}", response_class=HTMLResponse)
def workflow_view(request: Request, record_id: str, version: int):
    with db() as conn:
        wf = conn.execute(
            "SELECT * FROM workflows WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not wf:
            raise HTTPException(404)
        refs = conn.execute(
            """
            SELECT r.order_index, t.record_id, t.version, t.title, t.status as task_status
            FROM workflow_task_refs r
            JOIN tasks t ON t.record_id=r.task_record_id AND t.version=r.task_version
            WHERE r.workflow_record_id=? AND r.workflow_version=?
            ORDER BY r.order_index
            """,
            (record_id, version),
        ).fetchall()

        refs_pairs = [(r["record_id"], int(r["version"])) for r in refs]
        readiness_info = workflow_readiness_detail(conn, refs_pairs)
        doms = _workflow_domains(conn, refs_pairs)

    return templates.TemplateResponse(
        request,
        "workflow_view.html",
        {
            "workflow": dict(wf),
            "refs": refs,
            "readiness": readiness_info["readiness"],
            "readiness_reasons": readiness_info["reasons"],
            "blocking_task_refs": readiness_info["blocking_task_refs"],
            "domains": doms,
        },
    )


@app.get("/workflows/{record_id}/{version}/revise", response_class=HTMLResponse)
def workflow_revise_form(request: Request, record_id: str, version: int):
    require(request.state.role, "workflow:revise")
    with db() as conn:
        wf = conn.execute(
            "SELECT * FROM workflows WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not wf:
            raise HTTPException(404)
        refs = conn.execute(
            """
            SELECT r.order_index, r.task_record_id, r.task_version
            FROM workflow_task_refs r
            WHERE r.workflow_record_id=? AND r.workflow_version=?
            ORDER BY r.order_index
            """,
            (record_id, version),
        ).fetchall()
        confirmed_tasks = conn.execute(
            "SELECT record_id, version, title FROM tasks WHERE status='confirmed' ORDER BY title"
        ).fetchall()

    refs_text = "\n".join([f"{r['task_record_id']}@{r['task_version']}" for r in refs])
    return templates.TemplateResponse(
        request,
        "workflow_edit.html",
        {
            "mode": "revise",
            "workflow": dict(wf),
            "confirmed_tasks": confirmed_tasks,
            "refs_text": refs_text,
        },
    )


@app.post("/workflows/{record_id}/{version}/revise")
def workflow_revise(
    request: Request,
    record_id: str,
    version: int,
    title: str = Form(...),
    objective: str = Form(...),
    tags: str = Form(""),
    meta: str = Form(""),
    task_refs: str = Form(""),
    change_note: str = Form(""),
):
    require(request.state.role, "workflow:revise")
    actor = request.state.user

    note = change_note.strip()
    if not note:
        raise HTTPException(status_code=400, detail="change_note is required when creating a new version")

    refs = _parse_task_refs(task_refs)
    tags_list = parse_tags(tags)
    meta_obj = parse_meta(meta)

    with db() as conn:
        src = conn.execute(
            "SELECT * FROM workflows WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not src:
            raise HTTPException(404)

        enforce_workflow_ref_rules(conn, refs)

        latest_v = get_latest_version(conn, "workflows", record_id) or version
        new_v = latest_v + 1
        now = utc_now_iso()

        conn.execute(
            """
            INSERT INTO workflows(
              record_id, version, status,
              title, objective,
              domains_json,
              tags_json, meta_json,
              created_at, updated_at, created_by, updated_by,
              reviewed_at, reviewed_by, change_note,
              needs_review_flag, needs_review_note
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (
                record_id,
                new_v,
                "draft",
                title.strip(),
                objective.strip(),
                _json_dump(_workflow_domains(conn, refs)),
                _json_dump(tags_list),
                _json_dump(meta_obj),
                now,
                now,
                actor,
                actor,
                None,
                None,
                note,
                int(src["needs_review_flag"]),
                src["needs_review_note"],
            ),
        )

        for idx, (rid, ver) in enumerate(refs, start=1):
            conn.execute(
                """
                INSERT INTO workflow_task_refs(workflow_record_id, workflow_version, order_index, task_record_id, task_version)
                VALUES (?,?,?,?,?)
                """,
                (record_id, new_v, idx, rid, ver),
            )

    audit("workflow", record_id, new_v, "new_version", actor, note=f"from v{version}: {note}")
    return RedirectResponse(url=f"/workflows/{record_id}/{new_v}", status_code=303)


@app.post("/workflows/{record_id}/{version}/submit")
def workflow_submit(request: Request, record_id: str, version: int):
    require(request.state.role, "workflow:submit")
    actor = request.state.user
    with db() as conn:
        row = conn.execute(
            "SELECT status FROM workflows WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not row:
            raise HTTPException(404)
        if row["status"] != "draft":
            raise HTTPException(409, detail="Only draft workflows can be submitted")

        # Author domain gate: you can't submit workflows containing domains you don't hold.
        refs = conn.execute(
            "SELECT task_record_id, task_version FROM workflow_task_refs WHERE workflow_record_id=? AND workflow_version=? ORDER BY order_index",
            (record_id, version),
        ).fetchall()
        doms = _workflow_domains(conn, [(r["task_record_id"], int(r["task_version"])) for r in refs])
        missing = [d for d in doms if not _user_has_domain(conn, actor, d)]
        if missing:
            raise HTTPException(status_code=403, detail=f"Forbidden: you are not authorized for workflow domain(s): {', '.join(missing)}")

        conn.execute(
            "UPDATE workflows SET status='submitted', updated_at=?, updated_by=? WHERE record_id=? AND version=?",
            (utc_now_iso(), actor, record_id, version),
        )
    audit("workflow", record_id, version, "submit", actor)
    return RedirectResponse(url=f"/workflows/{record_id}/{version}", status_code=303)


@app.post("/workflows/{record_id}/{version}/force-submit")
def workflow_force_submit(request: Request, record_id: str, version: int):
    require(request.state.role, "workflow:force_submit")
    actor = request.state.user
    with db() as conn:
        row = conn.execute(
            "SELECT status FROM workflows WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not row:
            raise HTTPException(404)
        if row["status"] in ("deprecated", "confirmed"):
            raise HTTPException(409, detail=f"Cannot force-submit a {row['status']} workflow")
        conn.execute(
            "UPDATE workflows SET status='submitted', updated_at=?, updated_by=? WHERE record_id=? AND version=?",
            (utc_now_iso(), actor, record_id, version),
        )
    audit("workflow", record_id, version, "force_submit", actor, note="admin forced submission")
    return RedirectResponse(url=f"/workflows/{record_id}/{version}", status_code=303)


@app.post("/workflows/{record_id}/{version}/confirm")
def workflow_confirm(request: Request, record_id: str, version: int):
    require(request.state.role, "workflow:confirm")
    actor = request.state.user
    with db() as conn:
        row = conn.execute(
            "SELECT status FROM workflows WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not row:
            raise HTTPException(404)
        if row["status"] != "submitted":
            raise HTTPException(409, detail="Only submitted workflows can be confirmed")

        refs = conn.execute(
            "SELECT task_record_id, task_version FROM workflow_task_refs WHERE workflow_record_id=? AND workflow_version=? ORDER BY order_index",
            (record_id, version),
        ).fetchall()
        readiness = workflow_readiness(conn, [(r["task_record_id"], int(r["task_version"])) for r in refs])
        if readiness != "ready":
            raise HTTPException(
                status_code=409,
                detail="Cannot confirm workflow: all referenced Task versions must be confirmed.",
            )

        doms = _workflow_domains(conn, [(r["task_record_id"], int(r["task_version"])) for r in refs])
        missing = [d for d in doms if not _user_has_domain(conn, actor, d)]
        if missing:
            raise HTTPException(status_code=403, detail=f"Forbidden: you are not authorized to confirm workflow domain(s): {', '.join(missing)}")

        # Deprecate any previously confirmed version
        conn.execute(
            "UPDATE workflows SET status='deprecated', updated_at=?, updated_by=? WHERE record_id=? AND status='confirmed'",
            (utc_now_iso(), actor, record_id),
        )

        conn.execute(
            """
            UPDATE workflows
            SET status='confirmed', reviewed_at=?, reviewed_by=?, updated_at=?, updated_by=?
            WHERE record_id=? AND version=?
            """,
            (utc_now_iso(), actor, utc_now_iso(), actor, record_id, version),
        )

    audit("workflow", record_id, version, "confirm", actor)
    return RedirectResponse(url=f"/workflows/{record_id}/{version}", status_code=303)


@app.post("/workflows/{record_id}/{version}/force-confirm")
def workflow_force_confirm(request: Request, record_id: str, version: int):
    require(request.state.role, "workflow:force_confirm")
    actor = request.state.user
    with db() as conn:
        row = conn.execute(
            "SELECT status FROM workflows WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not row:
            raise HTTPException(404)
        if row["status"] == "deprecated":
            raise HTTPException(409, detail="Cannot force-confirm a deprecated workflow")

        refs = conn.execute(
            "SELECT task_record_id, task_version FROM workflow_task_refs WHERE workflow_record_id=? AND workflow_version=? ORDER BY order_index",
            (record_id, version),
        ).fetchall()
        readiness = workflow_readiness(conn, [(r["task_record_id"], int(r["task_version"])) for r in refs])
        if readiness != "ready":
            raise HTTPException(
                status_code=409,
                detail="Cannot force-confirm workflow: referenced Task versions must still be confirmed.",
            )

        # Deprecate any previously confirmed version
        conn.execute(
            "UPDATE workflows SET status='deprecated', updated_at=?, updated_by=? WHERE record_id=? AND status='confirmed'",
            (utc_now_iso(), actor, record_id),
        )

        conn.execute(
            """
            UPDATE workflows
            SET status='confirmed', reviewed_at=?, reviewed_by=?, updated_at=?, updated_by=?
            WHERE record_id=? AND version=?
            """,
            (utc_now_iso(), actor, utc_now_iso(), actor, record_id, version),
        )

    audit("workflow", record_id, version, "force_confirm", actor, note="admin forced confirmation")
    return RedirectResponse(url=f"/workflows/{record_id}/{version}", status_code=303)


# ---- Assessments (MVP) ----

AssessmentClaim = Literal["auto", "fact_probe", "concept_probe", "procedure_proxy"]


def _assessment_domains(conn: sqlite3.Connection, refs: list[dict[str, Any]]) -> list[str]:
    """Derive domains from attached task/workflow refs."""
    doms: set[str] = set()
    for r in refs or []:
        rt = (r.get("ref_type") or "").strip().lower()
        rid = str(r.get("ref_record_id") or "").strip()
        ver = int(r.get("ref_version") or 0)
        if not rt or not rid or ver <= 0:
            continue
        if rt == "task":
            row = conn.execute("SELECT domain FROM tasks WHERE record_id=? AND version=?", (rid, ver)).fetchone()
            d = (str(row["domain"]) if row else "").strip().lower()
            if d:
                doms.add(d)
        elif rt == "workflow":
            row = conn.execute("SELECT domains_json FROM workflows WHERE record_id=? AND version=?", (rid, ver)).fetchone()
            if row and row["domains_json"]:
                for d in (_json_load(row["domains_json"]) or []):
                    dn = (str(d) or "").strip().lower()
                    if dn:
                        doms.add(dn)
    return sorted(doms)


def _assessment_lint(stem: str, options: list[dict[str, str]], correct_key: str, claim: str) -> list[dict[str, Any]]:
    """Return lint findings [{level, code, msg}]. level in (error|warn)."""
    findings: list[dict[str, Any]] = []

    stem_norm = (stem or "").strip()
    if not stem_norm:
        findings.append({"level": "error", "code": "stem.empty", "msg": "Stem is required"})

    if "which of the following" in stem_norm.lower():
        findings.append({"level": "warn", "code": "stem.which_of_following", "msg": "Avoid 'which of the following' phrasing"})

    # Options
    if len(options) != 4:
        findings.append({"level": "error", "code": "options.count", "msg": "Exactly 4 options are required"})

    keys = [str(o.get("key") or "").strip().upper() for o in options]
    texts = [str(o.get("text") or "").strip() for o in options]

    if any(not k or k not in ("A", "B", "C", "D") for k in keys):
        findings.append({"level": "error", "code": "options.keys", "msg": "Option keys must be A, B, C, D"})

    if any(not t for t in texts):
        findings.append({"level": "error", "code": "options.empty", "msg": "All option texts are required"})

    ck = (correct_key or "").strip().upper()
    if ck not in keys:
        findings.append({"level": "error", "code": "correct.missing", "msg": "Correct answer key must match one of the options"})

    # Duplicate texts
    seen: set[str] = set()
    for t in texts:
        tn = re.sub(r"\s+", " ", (t or "").strip().lower())
        if not tn:
            continue
        if tn in seen:
            findings.append({"level": "error", "code": "options.duplicate", "msg": "Duplicate option text detected"})
            break
        seen.add(tn)

    # Absolute terms heuristic
    abs_terms = (" always ", " never ", " only ")
    for idx, t in enumerate(texts):
        tl = f" {t.lower()} "
        if any(a in tl for a in abs_terms):
            findings.append({"level": "warn", "code": "options.absolute_terms", "msg": f"Option {keys[idx] or '?'} contains absolute terms (always/never/only)"})
            break

    # Length band (warn)
    wcounts = [len((t or "").split()) for t in texts if t]
    if wcounts:
        if max(wcounts) - min(wcounts) > 6:
            findings.append({"level": "warn", "code": "options.length_band", "msg": "Options vary widely in length; this can create visual clues"})

    claim_norm = (claim or "").strip()
    if claim_norm not in ("auto", "fact_probe", "concept_probe", "procedure_proxy"):
        findings.append({"level": "error", "code": "claim.invalid", "msg": "Invalid claim (auto|fact_probe|concept_probe|procedure_proxy)"})

    # Auto is allowed and intentionally does not apply claim-specific linting.
    if claim_norm == "auto":
        return findings

    # Only apply scenario hinting when explicitly in procedure_proxy.
    if claim_norm == "procedure_proxy":

        # Soft check: require some scenario signal.
        if not any(x in stem_norm.lower() for x in ("scenario", "you", "environment", "given", "after")):
            findings.append({"level": "warn", "code": "procedure_proxy.weak_scenario", "msg": "Procedure proxy items should usually be scenario-framed"})

    return findings


def _assessment_export_dict(row: sqlite3.Row | dict[str, Any]) -> dict[str, Any]:
    r = dict(row)
    return {
        "type": "assessment",
        "record_id": r["record_id"],
        "version": int(r["version"]),
        "status": r["status"],
        "stem": r["stem"],
        "options": _json_load(r["options_json"]) or [],
        "correct_key": r["correct_key"],
        "rationale": r.get("rationale") or "",
        "claim": r.get("claim") or "fact_probe",
        "domains": _json_load(r.get("domains_json") or "[]") or [],
        "refs": _json_load(r.get("refs_json") or "[]") or [],
        "lint": _json_load(r.get("lint_json") or "[]") or [],
        "needs_review_flag": bool(r.get("needs_review_flag")),
        "needs_review_note": r.get("needs_review_note"),
        "meta": {
            "created_at": r.get("created_at"),
            "updated_at": r.get("updated_at"),
            "created_by": r.get("created_by"),
            "updated_by": r.get("updated_by"),
            "reviewed_at": r.get("reviewed_at"),
            "reviewed_by": r.get("reviewed_by"),
            "change_note": r.get("change_note"),
        },
    }


@app.get("/_refs/search")
def refs_search(request: Request, kind: str = "task", q: str = "", limit: int = 20):
    """Lightweight ref search for the assessment ref selector UI."""
    require(request.state.role, "assessment:create")

    kind = (kind or "task").strip().lower()
    q_norm = (q or "").strip().lower()
    limit = max(1, min(int(limit or 20), 50))

    with db() as conn:
        if kind == "workflow":
            rows = conn.execute("SELECT record_id, MAX(version) AS latest_version FROM workflows GROUP BY record_id").fetchall()
            items: list[dict[str, Any]] = []
            for r in rows:
                rid = r["record_id"]
                v = int(r["latest_version"])
                w = conn.execute(
                    "SELECT record_id, version, title, status, domains_json FROM workflows WHERE record_id=? AND version=?",
                    (rid, v),
                ).fetchone()
                if not w:
                    continue
                if q_norm and q_norm not in (w["title"] or "").lower():
                    continue
                doms = [str(x).strip().lower() for x in (_json_load(w["domains_json"]) or []) if str(x).strip()]
                items.append({"ref_type": "workflow", "record_id": w["record_id"], "version": int(w["version"]), "title": w["title"], "status": w["status"], "domains": doms})
            items.sort(key=lambda it: (it.get("title") or ""))
            return {"items": items[:limit]}

        # default tasks
        rows = conn.execute("SELECT record_id, MAX(version) AS latest_version FROM tasks GROUP BY record_id").fetchall()
        items = []
        for r in rows:
            rid = r["record_id"]
            v = int(r["latest_version"])
            t = conn.execute(
                "SELECT record_id, version, title, status, domain FROM tasks WHERE record_id=? AND version=?",
                (rid, v),
            ).fetchone()
            if not t:
                continue
            if q_norm and q_norm not in (t["title"] or "").lower():
                continue
            items.append({"ref_type": "task", "record_id": t["record_id"], "version": int(t["version"]), "title": t["title"], "status": t["status"], "domain": t["domain"]})
        items.sort(key=lambda it: (it.get("title") or ""))
        return {"items": items[:limit]}


@app.get("/_refs/peek", response_class=HTMLResponse)
def refs_peek(request: Request, ref_type: str, record_id: str, version: int, component: str = "facts"):
    """Open a small window for assessment authors to view underlying task/workflow data.

    Intentional constraint: we do not surface dependencies as an assessable target.
    Dependencies tend to produce low-quality "prereq chain" questions.

    This is intentionally a separate page (not inline) to avoid spamming the assessment form.
    """
    require(request.state.role, "assessment:create")

    ref_type = (ref_type or "").strip().lower()
    component = (component or "facts").strip().lower()
    if component not in ("facts", "concepts", "procedure"):
        component = "facts"

    with db() as conn:
        if ref_type == "task":
            row = conn.execute(
                "SELECT record_id, version, title, status, domain, outcome, procedure_name, facts_json, concepts_json, steps_json FROM tasks WHERE record_id=? AND version=?",
                (record_id, int(version)),
            ).fetchone()
            if not row:
                raise HTTPException(404)
            return templates.TemplateResponse(
                request,
                "ref_peek.html",
                {
                    "kind": "task",
                    "ref_type": "task",
                    "record_id": row["record_id"],
                    "version": int(row["version"]),
                    "title": row["title"],
                    "status": row["status"],
                    "domain": row["domain"],
                    "outcome": row["outcome"],
                    "procedure_name": row["procedure_name"],
                    "component": component,
                    "facts": _json_load(row["facts_json"]) or [],
                    "concepts": _json_load(row["concepts_json"]) or [],
                    "steps": _normalize_steps(_json_load(row["steps_json"]) or []),
                    "tasks": [],
                },
            )

        if ref_type == "workflow":
            row = conn.execute(
                "SELECT record_id, version, title, status, objective FROM workflows WHERE record_id=? AND version=?",
                (record_id, int(version)),
            ).fetchone()
            if not row:
                raise HTTPException(404)

            refs = conn.execute(
                "SELECT task_record_id, task_version FROM workflow_task_refs WHERE workflow_record_id=? AND workflow_version=? ORDER BY order_index",
                (record_id, int(version)),
            ).fetchall()

            tasks: list[dict[str, Any]] = []
            agg: list[str] = []
            seen: dict[str, str] = {}

            def _norm_key(s: str) -> str:
                return re.sub(r"\s+", " ", (s or "").strip().lower())

            for r in refs:
                t = conn.execute(
                    "SELECT record_id, version, title, domain, outcome, procedure_name, facts_json, concepts_json, steps_json FROM tasks WHERE record_id=? AND version=?",
                    (r["task_record_id"], int(r["task_version"])),
                ).fetchone()
                if not t:
                    continue

                tasks.append(
                    {
                        "record_id": t["record_id"],
                        "version": int(t["version"]),
                        "title": t["title"],
                        "domain": t["domain"],
                        "outcome": t["outcome"],
                        "procedure_name": t["procedure_name"],
                        "steps": _normalize_steps(_json_load(t["steps_json"]) or []),
                    }
                )

                if component == "facts":
                    for x in (_json_load(t["facts_json"]) or []):
                        k = _norm_key(str(x))
                        if k and k not in seen:
                            seen[k] = str(x)
                elif component == "concepts":
                    for x in (_json_load(t["concepts_json"]) or []):
                        k = _norm_key(str(x))
                        if k and k not in seen:
                            seen[k] = str(x)

            if component in ("facts", "concepts"):
                agg = sorted(seen.values(), key=lambda s: _norm_key(s))

            return templates.TemplateResponse(
                request,
                "ref_peek.html",
                {
                    "kind": "workflow",
                    "ref_type": "workflow",
                    "record_id": row["record_id"],
                    "version": int(row["version"]),
                    "title": row["title"],
                    "status": row["status"],
                    "objective": row["objective"],
                    "component": component,
                    "agg": agg,
                    "tasks": tasks,
                    # unused for workflow branch but template expects keys
                    "domain": "",
                    "outcome": "",
                    "procedure_name": "",
                    "facts": [],
                    "concepts": [],
                    "steps": [],
                },
            )

    raise HTTPException(status_code=400, detail="Invalid ref_type")


@app.get("/assessments", response_class=HTMLResponse)
def assessments_list(request: Request, status: str | None = None, q: str | None = None, domain: str | None = None, claim: str | None = None):
    q_norm = (q or "").strip().lower()
    domain_norm = (domain or "").strip().lower() or None
    claim_norm = (claim or "").strip().lower() or None

    with db() as conn:
        sql = "SELECT record_id, MAX(version) AS latest_version FROM assessment_items GROUP BY record_id ORDER BY record_id"
        rows = conn.execute(sql).fetchall()

        domains = _active_domains(conn)
        items: list[dict[str, Any]] = []
        for r in rows:
            rid = r["record_id"]
            latest_v = int(r["latest_version"])
            latest = conn.execute(
                "SELECT * FROM assessment_items WHERE record_id=? AND version=?", (rid, latest_v)
            ).fetchone()
            if not latest:
                continue
            if status and latest["status"] != status:
                continue
            if q_norm and q_norm not in (latest["stem"] or "").lower():
                continue
            # sqlite3.Row doesn't implement .get
            if claim_norm and (str(latest["claim"] or "").strip().lower() != claim_norm):
                continue

            doms = [str(d).strip().lower() for d in (_json_load(latest["domains_json"] or "[]") or [])]
            if domain_norm and domain_norm not in set(doms):
                continue

            items.append(
                {
                    "record_id": rid,
                    "latest_version": latest_v,
                    "stem": latest["stem"],
                    "status": latest["status"],
                    "claim": (latest["claim"] or "auto"),
                    "domains": doms,
                    "needs_review_flag": bool(latest["needs_review_flag"]),
                }
            )

    return templates.TemplateResponse(
        request,
        "assessments_list.html",
        {"items": items, "status": status, "q": q, "domain": domain_norm or "", "domains": domains, "claim": claim_norm or ""},
    )


@app.get("/delivery", response_class=HTMLResponse)
def delivery_page(request: Request, q: str | None = None, domain: str | None = None, tag: str | None = None):
    require(request.state.role, "delivery:view")

    q_norm = (q or "").strip().lower()
    domain_norm = (domain or "").strip().lower() or None
    tag_norm = (tag or "").strip().lower() or None

    with db() as conn:
        domains = _active_domains(conn)

        # Latest confirmed version per workflow
        rows = conn.execute(
            "SELECT record_id, MAX(version) AS v FROM workflows WHERE status='confirmed' GROUP BY record_id ORDER BY record_id"
        ).fetchall()

        workflows: list[dict[str, Any]] = []
        all_tags: set[str] = set()

        for r in rows:
            rid = r["record_id"]
            v = int(r["v"])
            wf = conn.execute(
                "SELECT record_id, version, title, tags_json, domains_json FROM workflows WHERE record_id=? AND version=?",
                (rid, v),
            ).fetchone()
            if not wf:
                continue

            title = str(wf["title"] or "")
            if q_norm and q_norm not in title.lower():
                continue

            wf_tags = [str(x).strip().lower() for x in (_json_load(wf["tags_json"]) or []) if str(x).strip()]
            for t in wf_tags:
                all_tags.add(t)

            if tag_norm and tag_norm not in set(wf_tags):
                continue

            wf_domains = [str(x).strip().lower() for x in (_json_load(wf["domains_json"]) or []) if str(x).strip()]
            if domain_norm and domain_norm not in set(wf_domains):
                continue

            workflows.append({"record_id": wf["record_id"], "version": v, "title": title})

    workflows.sort(key=lambda w: (w.get("title") or ""))

    return templates.TemplateResponse(
        request,
        "delivery.html",
        {"workflows": workflows, "q": q, "domain": domain_norm or "", "tag": tag_norm or "", "domains": domains, "tags": sorted(all_tags)},
    )


@app.post("/delivery/export")
def delivery_export(request: Request, workflow_key: str = Form(""), modality: str = Form("docx")):
    require(request.state.role, "delivery:export")

    wk = (workflow_key or "").strip()
    if "@" not in wk:
        raise HTTPException(status_code=400, detail="workflow_key is required")
    rid, ver_s = wk.split("@", 1)
    try:
        ver = int(ver_s)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid workflow version")

    mod = (modality or "docx").strip().lower()
    if mod == "docx":
        return RedirectResponse(url=f"/workflows/{rid}/{ver}/export.docx", status_code=303)
    if mod == "md":
        return RedirectResponse(url=f"/workflows/{rid}/{ver}/export.md", status_code=303)

    raise HTTPException(status_code=409, detail=f"Modality '{mod}' is not operational yet")


@app.get("/assessments/new", response_class=HTMLResponse)
def assessment_new_form(
    request: Request,
    q: str | None = None,
    task_record_id: str | None = None,
    task_version: int | None = None,
    ref_type: str | None = None,
    ref_record_id: str | None = None,
    ref_version: int | None = None,
):
    require(request.state.role, "assessment:create")

    # Support older task_* params, and newer generic ref_* params.
    if task_record_id and task_version and not (ref_type and ref_record_id and ref_version):
        ref_type = "task"
        ref_record_id = task_record_id
        ref_version = int(task_version)

    if not (ref_type and ref_record_id and ref_version):
        # Picker UI
        q_norm = (q or "").strip().lower()
        with db() as conn:
            # Tasks (latest versions)
            t_rows = conn.execute("SELECT record_id, MAX(version) AS latest_version FROM tasks GROUP BY record_id ORDER BY record_id").fetchall()
            tasks: list[dict[str, Any]] = []
            for r in t_rows:
                rid = r["record_id"]
                v = int(r["latest_version"])
                t = conn.execute("SELECT record_id, version, title, status, domain FROM tasks WHERE record_id=? AND version=?", (rid, v)).fetchone()
                if not t:
                    continue
                if q_norm and q_norm not in (t["title"] or "").lower():
                    continue
                tasks.append(dict(t))
            tasks = tasks[:30]

            # Workflows (latest versions)
            w_rows = conn.execute("SELECT record_id, MAX(version) AS latest_version FROM workflows GROUP BY record_id ORDER BY record_id").fetchall()
            workflows: list[dict[str, Any]] = []
            for r in w_rows:
                rid = r["record_id"]
                v = int(r["latest_version"])
                w = conn.execute("SELECT record_id, version, title, status, domains_json FROM workflows WHERE record_id=? AND version=?", (rid, v)).fetchone()
                if not w:
                    continue
                if q_norm and q_norm not in (w["title"] or "").lower():
                    continue
                wd = dict(w)
                wd["domains"] = [str(x).strip().lower() for x in (_json_load(w["domains_json"]) or []) if str(x).strip()]
                workflows.append(wd)
            workflows = workflows[:30]

        return templates.TemplateResponse(request, "assessment_pick_ref.html", {"q": q, "tasks": tasks, "workflows": workflows})

    ref = {"ref_type": (ref_type or "").strip().lower(), "ref_record_id": (ref_record_id or "").strip(), "ref_version": int(ref_version)}

    return templates.TemplateResponse(
        request,
        "assessment_edit.html",
        {
            "mode": "new",
            "item": None,
            "refs": [ref],
            "lint": [],
        },
    )


@app.post("/assessments/new")
def assessment_create(
    request: Request,
    stem: str = Form(""),
    claim: str = Form("auto"),
    correct_key: str = Form("A"),
    option_a: str = Form(""),
    option_b: str = Form(""),
    option_c: str = Form(""),
    option_d: str = Form(""),
    rationale: str = Form(""),
    change_note: str = Form(""),
    # authoring helpers (stored in meta_json)
    target_fact: str = Form(""),
    relation_verb: str = Form(""),
    scenario_truth: str = Form(""),
    ref_type: list[str] = Form([]),
    ref_record_id: list[str] = Form([]),
    ref_version: list[int] = Form([]),
):
    require(request.state.role, "assessment:create")
    actor = request.state.user

    record_id = str(uuid.uuid4())
    version = 1
    now = utc_now_iso()

    options = [
        {"key": "A", "text": (option_a or "").strip()},
        {"key": "B", "text": (option_b or "").strip()},
        {"key": "C", "text": (option_c or "").strip()},
        {"key": "D", "text": (option_d or "").strip()},
    ]

    refs: list[dict[str, Any]] = []
    for rt, rid, ver in zip(ref_type or [], ref_record_id or [], ref_version or []):
        rt_n = (rt or "").strip().lower()
        rid_n = (rid or "").strip()
        try:
            ver_i = int(ver)
        except Exception:
            ver_i = 0
        if rt_n and rid_n and ver_i > 0:
            refs.append({"ref_type": rt_n, "ref_record_id": rid_n, "ref_version": ver_i})

    with db() as conn:
        domains = _assessment_domains(conn, refs)
        lint = _assessment_lint(stem, options, correct_key, claim)

        meta_obj = {
            "target_fact": (target_fact or "").strip(),
            "relation_verb": (relation_verb or "").strip(),
            "scenario_truth": (scenario_truth or "").strip(),
        }

        conn.execute(
            """
            INSERT INTO assessment_items(
              record_id, version, status,
              stem, options_json, correct_key, rationale,
              claim, domains_json, lint_json, refs_json,
              tags_json, meta_json,
              created_at, updated_at, created_by, updated_by,
              reviewed_at, reviewed_by, change_note,
              needs_review_flag, needs_review_note
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (
                record_id,
                version,
                "draft",
                (stem or "").strip(),
                _json_dump(options),
                (correct_key or "A").strip().upper(),
                (rationale or "").strip(),
                (claim or "auto").strip().lower(),
                _json_dump(domains),
                _json_dump(lint),
                _json_dump(refs),
                _json_dump([]),
                _json_dump(meta_obj),
                now,
                now,
                actor,
                actor,
                None,
                None,
                (change_note or "").strip() or None,
                0,
                None,
            ),
        )

        # refs table
        for idx, r in enumerate(refs, start=1):
            conn.execute(
                "INSERT INTO assessment_refs(assessment_record_id, assessment_version, order_index, ref_type, ref_record_id, ref_version) VALUES (?,?,?,?,?,?)",
                (record_id, version, idx, r["ref_type"], r["ref_record_id"], int(r["ref_version"])),
            )

        audit("assessment", record_id, version, "create", actor, conn=conn)

    return RedirectResponse(url=f"/assessments/{record_id}/{version}/edit?created=1", status_code=303)


@app.get("/assessments/{record_id}/{version}", response_class=HTMLResponse)
def assessment_view(request: Request, record_id: str, version: int):
    with db() as conn:
        row = conn.execute(
            "SELECT * FROM assessment_items WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not row:
            raise HTTPException(404)

        ref_rows = conn.execute(
            "SELECT ref_type, ref_record_id, ref_version FROM assessment_refs WHERE assessment_record_id=? AND assessment_version=? ORDER BY order_index",
            (record_id, version),
        ).fetchall()

    item = _assessment_export_dict(row)
    item["refs"] = [dict(r) for r in ref_rows]

    return_note = None
    if item.get("status") == "returned":
        with db() as conn:
            rn = conn.execute(
                "SELECT note, at, actor FROM audit_log WHERE entity_type='assessment' AND record_id=? AND version=? AND action='return_for_changes' ORDER BY at DESC LIMIT 1",
                (record_id, version),
            ).fetchone()
            if rn and rn["note"]:
                return_note = {"note": rn["note"], "at": rn["at"], "actor": rn["actor"]}

    return templates.TemplateResponse(request, "assessment_view.html", {"item": item, "return_note": return_note})


@app.get("/assessments/{record_id}/{version}/edit", response_class=HTMLResponse)
def assessment_edit_form(request: Request, record_id: str, version: int):
    require(request.state.role, "assessment:revise")
    with db() as conn:
        row = conn.execute(
            "SELECT * FROM assessment_items WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not row:
            raise HTTPException(404)

        ref_rows = conn.execute(
            "SELECT ref_type, ref_record_id, ref_version FROM assessment_refs WHERE assessment_record_id=? AND assessment_version=? ORDER BY order_index",
            (record_id, version),
        ).fetchall()

    item = dict(row)
    lint = _json_load(item.get("lint_json") or "[]") or []
    options = _json_load(item.get("options_json") or "[]") or []
    item["options"] = options

    return templates.TemplateResponse(
        request,
        "assessment_edit.html",
        {"mode": "edit", "item": item, "refs": [dict(r) for r in ref_rows], "lint": lint},
    )


@app.post("/assessments/{record_id}/{version}/save")
def assessment_save(
    request: Request,
    record_id: str,
    version: int,
    stem: str = Form(""),
    claim: str = Form("auto"),
    correct_key: str = Form("A"),
    option_a: str = Form(""),
    option_b: str = Form(""),
    option_c: str = Form(""),
    option_d: str = Form(""),
    rationale: str = Form(""),
    change_note: str = Form(""),
    # authoring helpers (stored in meta_json)
    target_fact: str = Form(""),
    relation_verb: str = Form(""),
    scenario_truth: str = Form(""),
    ref_type: list[str] = Form([]),
    ref_record_id: list[str] = Form([]),
    ref_version: list[int] = Form([]),
):
    """Immutable records: saving creates a NEW VERSION (draft) with required change_note."""
    require(request.state.role, "assessment:revise")
    actor = request.state.user

    note = (change_note or "").strip()
    if not note:
        raise HTTPException(status_code=400, detail="change_note is required when creating a new version")

    options = [
        {"key": "A", "text": (option_a or "").strip()},
        {"key": "B", "text": (option_b or "").strip()},
        {"key": "C", "text": (option_c or "").strip()},
        {"key": "D", "text": (option_d or "").strip()},
    ]

    refs: list[dict[str, Any]] = []
    for rt, rid, ver in zip(ref_type or [], ref_record_id or [], ref_version or []):
        rt_n = (rt or "").strip().lower()
        rid_n = (rid or "").strip()
        try:
            ver_i = int(ver)
        except Exception:
            ver_i = 0
        if rt_n and rid_n and ver_i > 0:
            refs.append({"ref_type": rt_n, "ref_record_id": rid_n, "ref_version": ver_i})

    with db() as conn:
        src = conn.execute(
            "SELECT * FROM assessment_items WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not src:
            raise HTTPException(404)

        if src["status"] == "returned":
            rn = conn.execute(
                "SELECT note, at, actor FROM audit_log WHERE entity_type='assessment' AND record_id=? AND version=? AND action='return_for_changes' ORDER BY at DESC LIMIT 1",
                (record_id, version),
            ).fetchone()
            if rn and rn["note"]:
                prefix = f"Response to return note by {rn['actor']} at {rn['at']}: {rn['note']} | "
                if prefix not in note:
                    note = prefix + note

        latest_v = get_latest_version(conn, "assessment_items", record_id) or version
        new_v = latest_v + 1
        now = utc_now_iso()

        domains = _assessment_domains(conn, refs)
        lint = _assessment_lint(stem, options, correct_key, claim)

        meta_prev = _json_load((src["meta_json"] if "meta_json" in src.keys() else "{}") or "{}") or {}
        meta_prev.update(
            {
                "target_fact": (target_fact or "").strip(),
                "relation_verb": (relation_verb or "").strip(),
                "scenario_truth": (scenario_truth or "").strip(),
            }
        )

        conn.execute(
            """
            INSERT INTO assessment_items(
              record_id, version, status,
              stem, options_json, correct_key, rationale,
              claim, domains_json, lint_json, refs_json,
              tags_json, meta_json,
              created_at, updated_at, created_by, updated_by,
              reviewed_at, reviewed_by, change_note,
              needs_review_flag, needs_review_note
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (
                record_id,
                new_v,
                "draft",
                (stem or "").strip(),
                _json_dump(options),
                (correct_key or "A").strip().upper(),
                (rationale or "").strip(),
                (claim or "auto").strip().lower(),
                _json_dump(domains),
                _json_dump(lint),
                _json_dump(refs),
                (src["tags_json"] if "tags_json" in src.keys() else "[]"),
                _json_dump(meta_prev),
                now,
                now,
                actor,
                actor,
                None,
                None,
                note,
                int(src["needs_review_flag"]),
                src["needs_review_note"],
            ),
        )

        # refs table for new version
        for idx, r in enumerate(refs, start=1):
            conn.execute(
                "INSERT INTO assessment_refs(assessment_record_id, assessment_version, order_index, ref_type, ref_record_id, ref_version) VALUES (?,?,?,?,?,?)",
                (record_id, new_v, idx, r["ref_type"], r["ref_record_id"], int(r["ref_version"])),
            )

        audit("assessment", record_id, new_v, "new_version", actor, note=f"from v{version}: {note}", conn=conn)

    return RedirectResponse(url=f"/assessments/{record_id}/{new_v}", status_code=303)


@app.post("/assessments/{record_id}/{version}/submit")
def assessment_submit(request: Request, record_id: str, version: int):
    require(request.state.role, "assessment:submit")
    actor = request.state.user

    with db() as conn:
        row = conn.execute(
            "SELECT status, domains_json FROM assessment_items WHERE record_id=? AND version=?",
            (record_id, version),
        ).fetchone()
        if not row:
            raise HTTPException(404)
        if row["status"] != "draft":
            raise HTTPException(409, detail="Only draft assessments can be submitted")

        # Require at least one domain via refs
        doms = [str(d).strip().lower() for d in (_json_load(row["domains_json"]) or [])]
        if not doms:
            raise HTTPException(status_code=409, detail="Cannot submit assessment: attach it to at least one task/workflow")
        for d in doms:
            if not _user_has_domain(conn, actor, d):
                raise HTTPException(status_code=403, detail=f"Forbidden: you are not authorized for domain '{d}'")

        # Must pass lint errors
        lint = _json_load(
            conn.execute("SELECT lint_json FROM assessment_items WHERE record_id=? AND version=?", (record_id, version)).fetchone()["lint_json"]
        )
        for f in (lint or []):
            if (f.get("level") or "").lower() == "error":
                raise HTTPException(status_code=409, detail="Cannot submit assessment: fix lint errors first")

        conn.execute(
            "UPDATE assessment_items SET status='submitted', updated_at=?, updated_by=? WHERE record_id=? AND version=?",
            (utc_now_iso(), actor, record_id, version),
        )
        audit("assessment", record_id, version, "submit", actor, conn=conn)

    return RedirectResponse(url=f"/assessments/{record_id}/{version}", status_code=303)


@app.post("/assessments/{record_id}/{version}/return")
def assessment_return_for_changes(request: Request, record_id: str, version: int, note: str = Form("")):
    require(request.state.role, "assessment:confirm")
    actor = request.state.user
    msg = (note or "").strip()
    if not msg:
        raise HTTPException(status_code=400, detail="Return note is required")

    with db() as conn:
        row = conn.execute(
            "SELECT status FROM assessment_items WHERE record_id=? AND version=?",
            (record_id, version),
        ).fetchone()
        if not row:
            raise HTTPException(404)
        if row["status"] != "submitted":
            raise HTTPException(status_code=409, detail="Only submitted assessments can be returned")

        conn.execute(
            "UPDATE assessment_items SET status='returned', updated_at=?, updated_by=? WHERE record_id=? AND version=?",
            (utc_now_iso(), actor, record_id, version),
        )
        audit("assessment", record_id, version, "return_for_changes", actor, note=msg, conn=conn)

    return RedirectResponse(url=f"/assessments/{record_id}/{version}", status_code=303)


@app.post("/assessments/{record_id}/{version}/confirm")
def assessment_confirm(request: Request, record_id: str, version: int):
    require(request.state.role, "assessment:confirm")
    actor = request.state.user

    with db() as conn:
        row = conn.execute(
            "SELECT status FROM assessment_items WHERE record_id=? AND version=?",
            (record_id, version),
        ).fetchone()
        if not row:
            raise HTTPException(404)
        if row["status"] != "submitted":
            raise HTTPException(status_code=409, detail="Only submitted assessments can be confirmed")

        conn.execute(
            "UPDATE assessment_items SET status='deprecated', updated_at=?, updated_by=? WHERE record_id=? AND status='confirmed'",
            (utc_now_iso(), actor, record_id),
        )
        conn.execute(
            "UPDATE assessment_items SET status='confirmed', reviewed_at=?, reviewed_by=?, updated_at=?, updated_by=? WHERE record_id=? AND version=?",
            (utc_now_iso(), actor, utc_now_iso(), actor, record_id, version),
        )
        audit("assessment", record_id, version, "confirm", actor, conn=conn)

    return RedirectResponse(url=f"/assessments/{record_id}/{version}", status_code=303)


# ---- Export helpers ----

def _task_export_dict(row: sqlite3.Row | dict[str, Any]) -> dict[str, Any]:
    r = dict(row)
    return {
        "type": "task",
        "record_id": r["record_id"],
        "version": int(r["version"]),
        "status": r["status"],
        "title": r["title"],
        "outcome": r["outcome"],
        "facts": _json_load(r["facts_json"]) or [],
        "concepts": _json_load(r["concepts_json"]) or [],
        "procedure_name": r["procedure_name"],
        "steps": _normalize_steps(_json_load(r["steps_json"]) or []),
        "dependencies": _json_load(r["dependencies_json"]) or [],
        "irreversible_flag": bool(r["irreversible_flag"]),
        "task_assets": _json_load(r["task_assets_json"]) or [],
        "needs_review_flag": bool(r["needs_review_flag"]),
        "needs_review_note": r["needs_review_note"],
        "meta": {
            "created_at": r["created_at"],
            "updated_at": r["updated_at"],
            "created_by": r["created_by"],
            "updated_by": r["updated_by"],
            "reviewed_at": r["reviewed_at"],
            "reviewed_by": r["reviewed_by"],
            "change_note": r["change_note"],
        },
    }


def _workflow_export_dict(wf_row: sqlite3.Row, refs_rows: list[sqlite3.Row]) -> dict[str, Any]:
    wf = dict(wf_row)
    return {
        "type": "workflow",
        "record_id": wf["record_id"],
        "version": int(wf["version"]),
        "status": wf["status"],
        "title": wf["title"],
        "objective": wf["objective"],
        "task_refs": [
            {
                "order_index": int(r["order_index"]),
                "record_id": r["task_record_id"],
                "version": int(r["task_version"]),
            }
            for r in refs_rows
        ],
        "needs_review_flag": bool(wf["needs_review_flag"]),
        "needs_review_note": wf["needs_review_note"],
        "meta": {
            "created_at": wf["created_at"],
            "updated_at": wf["updated_at"],
            "created_by": wf["created_by"],
            "updated_by": wf["updated_by"],
            "reviewed_at": wf["reviewed_at"],
            "reviewed_by": wf["reviewed_by"],
            "change_note": wf["change_note"],
        },
    }


@app.get("/export/task/{record_id}/{version}.json")
def export_task_json(record_id: str, version: int):
    with db() as conn:
        row = conn.execute(
            "SELECT * FROM tasks WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
    if not row:
        raise HTTPException(404)

    if row["status"] != "confirmed":
        raise HTTPException(status_code=409, detail="Export is allowed for confirmed tasks only")

    payload = _task_export_dict(row)
    filename = f"task__{record_id}__v{version}.json"
    return JSONResponse(
        content=payload,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@app.get("/export/workflow/{record_id}/{version}.json")
def export_workflow_json(record_id: str, version: int):
    with db() as conn:
        wf = conn.execute(
            "SELECT * FROM workflows WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not wf:
            raise HTTPException(404)

        if wf["status"] != "confirmed":
            raise HTTPException(status_code=409, detail="Export is allowed for confirmed workflows only")

        refs = conn.execute(
            """
            SELECT order_index, task_record_id, task_version
            FROM workflow_task_refs
            WHERE workflow_record_id=? AND workflow_version=?
            ORDER BY order_index
            """,
            (record_id, version),
        ).fetchall()

        readiness = workflow_readiness(conn, [(r["task_record_id"], int(r["task_version"])) for r in refs])
        if readiness != "ready":
            raise HTTPException(status_code=409, detail="Export is allowed only when all referenced Task versions are confirmed")

    payload = _workflow_export_dict(wf, refs)
    filename = f"workflow__{record_id}__v{version}.json"
    return JSONResponse(
        content=payload,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@app.get("/workflows/{record_id}/{version}/export.docx")
def workflow_export_docx(request: Request, record_id: str, version: int):
    """Export a confirmed workflow to DOCX (v0 ILT handout).

    Governance rule: exports are allowed for confirmed workflows only, and only when
    all referenced Task versions are confirmed.
    """
    actor = request.state.user

    with db() as conn:
        wf = conn.execute(
            "SELECT * FROM workflows WHERE record_id=? AND version=?",
            (record_id, version),
        ).fetchone()
        if not wf:
            raise HTTPException(404)

        if wf["status"] != "confirmed":
            raise HTTPException(status_code=409, detail="Export is allowed for confirmed workflows only")

        ref_rows = conn.execute(
            """
            SELECT order_index, task_record_id, task_version
            FROM workflow_task_refs
            WHERE workflow_record_id=? AND workflow_version=?
            ORDER BY order_index
            """,
            (record_id, version),
        ).fetchall()

        readiness = workflow_readiness(conn, [(r["task_record_id"], int(r["task_version"])) for r in ref_rows])
        if readiness != "ready":
            raise HTTPException(status_code=409, detail="Export is allowed only when all referenced Task versions are confirmed")

        tasks: list[dict[str, Any]] = []
        for r in ref_rows:
            t = conn.execute(
                "SELECT * FROM tasks WHERE record_id=? AND version=?",
                (r["task_record_id"], int(r["task_version"])),
            ).fetchone()
            if not t:
                raise HTTPException(status_code=409, detail="Export failed: referenced task not found")
            if t["status"] != "confirmed":
                raise HTTPException(status_code=409, detail="Export is allowed only when all referenced Task versions are confirmed")
            tasks.append({"order_index": int(r["order_index"]), "row": dict(t)})

        # Build doc
        doc = Document()

        trace = f"Trace: {_short_code('WF', record_id)} v{version} · {utc_now_iso().split('T')[0]}"

        # Footer trace on each section
        for s in doc.sections:
            fp = s.footer.paragraphs[0] if s.footer.paragraphs else s.footer.add_paragraph()
            fp.text = trace

        doc.add_heading(str(wf["title"]), level=1)
        doc.add_paragraph(str(wf["objective"])).style = doc.styles["Normal"]

        doc.add_heading("Tasks", level=2)
        for t in tasks:
            r = t["row"]
            doc.add_paragraph(f"{t['order_index']}. {r['title']}")

        for t in tasks:
            r = t["row"]
            doc.add_page_break()
            doc.add_heading(f"Task {t['order_index']}: {r['title']}", level=2)
            doc.add_paragraph(f"Outcome: {r['outcome']}")

            steps = _normalize_steps(_json_load(r["steps_json"]) or [])
            doc.add_paragraph(f"Procedure: {r['procedure_name']}")

            table = doc.add_table(rows=1, cols=3)
            hdr = table.rows[0].cells
            hdr[0].text = "Step"
            hdr[1].text = "Actions"
            hdr[2].text = "Completion"

            for st in steps:
                row = table.add_row().cells
                row[0].text = str(st.get("text", "") or "")
                actions = st.get("actions") or []
                row[1].text = "\n".join([str(a) for a in actions if str(a).strip()]) if actions else ""
                row[2].text = str(st.get("completion", "") or "")

        # Provenance (full UUIDs)
        doc.add_page_break()
        doc.add_heading("Provenance (internal)", level=2)
        doc.add_paragraph(f"Workflow: {record_id} v{version}")
        for r in ref_rows:
            doc.add_paragraph(f"Task: {r['task_record_id']} v{int(r['task_version'])}")

        # Write artifact
        export_id = str(uuid.uuid4())
        ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
        filename = f"workflow__{_short_code('WF', record_id)}__v{version}__{ts}.docx"
        out_path = os.path.join(EXPORTS_DIR, filename)

        doc.save(out_path)
        file_bytes = Path(out_path).read_bytes()
        sha = _sha256_bytes(file_bytes)

        conn.execute(
            "INSERT INTO export_artifacts(id, kind, filename, path, sha256, workflow_record_id, workflow_version, task_refs_json, exported_at, exported_by, retention_days) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
            (
                export_id,
                "docx",
                filename,
                out_path,
                sha,
                record_id,
                int(version),
                _json_dump([{ "record_id": r["task_record_id"], "version": int(r["task_version"])} for r in ref_rows]),
                utc_now_iso(),
                actor,
                30,
            ),
        )
        audit("export", export_id, 1, "create", actor, note=f"kind=docx workflow={record_id}@{version}", conn=conn)

    return FileResponse(
        out_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )


@app.get("/workflows/{record_id}/{version}/export.md")
def workflow_export_md(record_id: str, version: int):
    with db() as conn:
        wf = conn.execute(
            "SELECT * FROM workflows WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not wf:
            raise HTTPException(404)

        if wf["status"] != "confirmed":
            raise HTTPException(status_code=409, detail="Export is allowed for confirmed workflows only")

        refs = conn.execute(
            """
            SELECT r.order_index, t.*
            FROM workflow_task_refs r
            JOIN tasks t ON t.record_id=r.task_record_id AND t.version=r.task_version
            WHERE r.workflow_record_id=? AND r.workflow_version=?
            ORDER BY r.order_index
            """,
            (record_id, version),
        ).fetchall()

        readiness = workflow_readiness(
            conn,
            [(r["record_id"], int(r["version"])) for r in refs],
        )
        if readiness != "ready":
            raise HTTPException(status_code=409, detail="Export is allowed only when all referenced Task versions are confirmed")

    lines: list[str] = []
    lines.append(f"# {wf['title']}")
    lines.append("")

    lines.append(f"**Objective:** {wf['objective']}")
    lines.append("")

    for r in refs:
        steps = _normalize_steps(_json_load(r["steps_json"]))
        facts = _json_load(r["facts_json"]) or []
        concepts = _json_load(r["concepts_json"]) or []
        deps = _json_load(r["dependencies_json"]) or []

        lines.append(f"## Task {r['order_index']}: {r['title']} ({r['record_id']}@{r['version']})")
        if r["status"] != "confirmed":
            lines.append(f"**Task status:** {r['status']} (unconfirmed)")
            lines.append("")
        lines.append("")
        lines.append(f"**Outcome:** {r['outcome']}")
        lines.append("")

        if facts:
            lines.append("**Facts:**")
            for f in facts:
                lines.append(f"- {f}")
            lines.append("")

        if concepts:
            lines.append("**Concepts:**")
            for c in concepts:
                lines.append(f"- {c}")
            lines.append("")

        if deps:
            lines.append("**Dependencies:**")
            for d in deps:
                lines.append(f"- {d}")
            lines.append("")

        lines.append(f"**Procedure:** {r['procedure_name']}")
        lines.append("")
        def _md_cell(s: str) -> str:
            s = (s or "").replace("\n", "<br>")
            s = s.replace("|", "\\|")
            return s

        lines.append("| Step | Actions | Completion |")
        lines.append("| --- | --- | --- |")
        for st in steps:
            txt = _md_cell(str(st.get("text", "") or ""))
            actions = st.get("actions") or []
            actions_txt = _md_cell("<br>".join([str(a) for a in actions if str(a).strip()])) if actions else "—"
            comp = _md_cell(str(st.get("completion", "") or "")) or "—"
            lines.append(f"| {txt} | {actions_txt} | {comp} |")
        lines.append("")

    md = "\n".join(lines)
    filename = f"workflow__{record_id}__v{version}.md"
    return HTMLResponse(
        content=md,
        media_type="text/markdown",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )
