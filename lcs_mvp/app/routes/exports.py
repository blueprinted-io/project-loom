from __future__ import annotations

import io
import json
import logging
import os
import sqlite3
import uuid
import zipfile
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Any

from docx import Document
from fastapi import APIRouter, Form, HTTPException, Request
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, RedirectResponse

from ..config import templates, EXPORTS_DIR, TASK_IMAGES_DIR, LOGO_TEMP_DIR
from ..database import db, utc_now_iso, workflow_readiness, _user_id
from ..audit import audit, _normalize_domains
from ..linting import _normalize_steps
from ..auth import require
from ..ingestion import _sha256_bytes, _short_code
from ..utils import _json_dump, _json_load

logger = logging.getLogger(__name__)

router = APIRouter()


def _parse_iso_dt(s: str) -> datetime | None:
    s = (s or "").strip()
    if not s:
        return None
    if s.endswith("Z"):
        s = s[:-1] + "+00:00"
    try:
        dt = datetime.fromisoformat(s)
    except ValueError:
        return None
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(timezone.utc)


def _cleanup_export_artifacts(conn: sqlite3.Connection) -> dict[str, int]:
    """Delete expired export artifacts (files + DB rows).

    Safety: only deletes files inside EXPORTS_DIR.
    """
    now = datetime.now(timezone.utc)
    exports_root = Path(EXPORTS_DIR).resolve()

    scanned = 0
    expired = 0
    deleted_files = 0
    missing_files = 0
    deleted_rows = 0

    rows = conn.execute(
        "SELECT id, path, exported_at, retention_days FROM export_artifacts ORDER BY exported_at ASC"
    ).fetchall()

    for r in rows:
        scanned += 1
        exported_at = _parse_iso_dt(str(r["exported_at"] or ""))
        if not exported_at:
            continue

        retention_days = int(r["retention_days"]) if r["retention_days"] is not None else 30
        if now <= (exported_at + timedelta(days=retention_days)):
            continue

        expired += 1

        p = Path(str(r["path"] or ""))
        try:
            p_abs = p.resolve()
        except (ValueError, OSError) as e:
            logger.warning("Export cleanup: could not resolve path %r: %s", r["path"], e)
            continue

        if exports_root not in p_abs.parents and p_abs != exports_root:
            # refuse to delete outside exports dir
            continue

        if p_abs.exists():
            try:
                p_abs.unlink()
                deleted_files += 1
            except OSError as e:
                logger.warning("Export cleanup: could not delete %s: %s", p_abs, e)
                continue
        else:
            missing_files += 1

        conn.execute("DELETE FROM export_artifacts WHERE id=?", (str(r["id"]),))
        deleted_rows += 1

    return {
        "scanned": scanned,
        "expired": expired,
        "deleted_files": deleted_files,
        "missing_files": missing_files,
        "deleted_rows": deleted_rows,
    }


def _task_export_dict(row: sqlite3.Row | dict[str, Any]) -> dict[str, Any]:
    r = dict(row)
    return {
        "type": "task",
        "record_id": r["record_id"],
        "version": int(r["version"]),
        "status": r["status"],
        "title": r["title"],
        "outcome": r["outcome"],
        "facts": _json_load(r["facts_json"]) or [],
        "concepts": _json_load(r["concepts_json"]) or [],
        "procedure_name": r["procedure_name"],
        "steps": _normalize_steps(_json_load(r["steps_json"]) or []),
        "dependencies": _json_load(r["dependencies_json"]) or [],
        "irreversible_flag": bool(r["irreversible_flag"]),
        "task_assets": _json_load(r["task_assets_json"]) or [],
        "needs_review_flag": bool(r["needs_review_flag"]),
        "needs_review_note": r["needs_review_note"],
        "meta": {
            "created_at": r["created_at"],
            "updated_at": r["updated_at"],
            "created_by": r["created_by"],
            "updated_by": r["updated_by"],
            "reviewed_at": r["reviewed_at"],
            "reviewed_by": r["reviewed_by"],
            "change_note": r["change_note"],
        },
    }


_PRIMER_ATTACH_QUERY = """
    SELECT p.record_id, p.title, p.summary, p.explanation, p.analogies, p.levels_json, p.version
    FROM workflow_primer_refs wpr
    JOIN primers p ON p.record_id = wpr.primer_record_id
    WHERE wpr.workflow_record_id = ?
      AND p.version = (
        SELECT MAX(p2.version) FROM primers p2
        WHERE p2.record_id = p.record_id AND p2.status = 'confirmed'
      )
    ORDER BY p.title
"""


def _workflow_export_dict(wf_row: sqlite3.Row, refs_rows: list[sqlite3.Row], primer_rows: list[sqlite3.Row] | None = None) -> dict[str, Any]:
    wf = dict(wf_row)
    return {
        "type": "workflow",
        "record_id": wf["record_id"],
        "version": int(wf["version"]),
        "status": wf["status"],
        "title": wf["title"],
        "objective": wf["objective"],
        "task_refs": [
            {
                "order_index": int(r["order_index"]),
                "record_id": r["task_record_id"],
                "version": int(r["task_version"]),
            }
            for r in refs_rows
        ],
        "primers": [
            {
                "record_id": p["record_id"],
                "version": int(p["version"]),
                "title": p["title"],
                "summary": p["summary"],
                "explanation": p["explanation"],
                "analogies": p["analogies"],
            }
            for p in (primer_rows or [])
        ],
        "needs_review_flag": bool(wf["needs_review_flag"]),
        "needs_review_note": wf["needs_review_note"],
        "meta": {
            "created_at": wf["created_at"],
            "updated_at": wf["updated_at"],
            "created_by": wf["created_by"],
            "updated_by": wf["updated_by"],
            "reviewed_at": wf["reviewed_at"],
            "reviewed_by": wf["reviewed_by"],
            "change_note": wf["change_note"],
        },
    }


@router.get("/exports", response_class=HTMLResponse)
def exports_library(request: Request, workflow: str = "", kind: str = "", by: str = "", msg: str | None = None):
    require(request.state.role, "export:library")

    wf = (workflow or "").strip()
    kind = (kind or "").strip().lower()
    by = (by or "").strip()

    where: list[str] = []
    params: list[object] = []

    if wf:
        where.append("workflow_record_id=?")
        params.append(wf)
    if kind:
        where.append("kind=?")
        params.append(kind)
    if by:
        where.append("exported_by=?")
        params.append(by)

    where_sql = ("WHERE " + " AND ".join(where)) if where else ""

    with db() as conn:
        artifacts = conn.execute(
            f"SELECT id, kind, filename, path, workflow_record_id, workflow_version, exported_at, exported_by, retention_days "
            f"FROM export_artifacts {where_sql} ORDER BY exported_at DESC LIMIT 200",
            params,
        ).fetchall()

    now = datetime.now(timezone.utc)
    out: list[dict[str, Any]] = []
    for a in artifacts:
        d = dict(a)
        exported_at = _parse_iso_dt(str(d.get("exported_at") or ""))
        retention_days = int(d.get("retention_days") or 30)
        if exported_at:
            expires_at = exported_at + timedelta(days=retention_days)
            d["expires_at"] = expires_at.replace(microsecond=0).isoformat()
            d["is_expired"] = now > expires_at
        else:
            d["expires_at"] = None
            d["is_expired"] = False
        out.append(d)

    return templates.TemplateResponse(
        request,
        "admin/exports.html",
        {
            "artifacts": out,
            "msg": msg,
            "filters": {"workflow": wf, "kind": kind, "by": by},
        },
    )


@router.get("/admin/exports", response_class=HTMLResponse)
def admin_exports_redirect(request: Request):
    # Backward-compatible URL for admins.
    require(request.state.role, "export:cleanup")
    return RedirectResponse(url="/exports", status_code=303)


@router.post("/admin/exports/cleanup")
def admin_exports_cleanup(request: Request):
    require(request.state.role, "export:cleanup")
    actor = request.state.user
    with db() as conn:
        stats = _cleanup_export_artifacts(conn)
        conn.commit()
        audit(
            "export_artifacts",
            "retention",
            1,
            "cleanup",
            actor,
            note=f"expired={stats['expired']} deleted_rows={stats['deleted_rows']} deleted_files={stats['deleted_files']} missing_files={stats['missing_files']}",
            conn=conn,
        )

    msg = (
        f"scanned={stats['scanned']} expired={stats['expired']} deleted_rows={stats['deleted_rows']} "
        f"deleted_files={stats['deleted_files']} missing_files={stats['missing_files']}"
    )
    return RedirectResponse(url=f"/exports?msg={msg}", status_code=303)


@router.get("/exports/{export_id}/download")
def export_download(request: Request, export_id: str):
    require(request.state.role, "export:library")
    with db() as conn:
        row = conn.execute(
            "SELECT filename, path FROM export_artifacts WHERE id=?",
            (export_id,),
        ).fetchone()
        if not row:
            raise HTTPException(404)

    p = Path(str(row["path"] or ""))
    try:
        p_abs = p.resolve()
        exports_root = Path(EXPORTS_DIR).resolve()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid artifact path")

    if exports_root not in p_abs.parents and p_abs != exports_root:
        raise HTTPException(status_code=400, detail="Invalid artifact path")

    if not p_abs.exists():
        raise HTTPException(status_code=404, detail="Artifact file missing")

    return FileResponse(
        str(p_abs),
        filename=str(row["filename"]),
        media_type="application/octet-stream",
    )


def _primers_html(primer_rows: list, esc_fn: Any) -> str:
    if not primer_rows:
        return ""
    parts = ["<h2>Pre-reading</h2>"]
    for p in primer_rows:
        parts.append(f"<h3>{esc_fn(str(p['title']))}</h3>")
        if p["summary"]:
            parts.append(f"<p><em>{esc_fn(str(p['summary']))}</em></p>")
        if p["explanation"]:
            parts.append(f"<p>{esc_fn(str(p['explanation']))}</p>")
        if p["analogies"]:
            parts.append(f"<p><strong>Analogy:</strong> {esc_fn(str(p['analogies']))}</p>")
    return "\n".join(parts)


@router.get("/workflows/{record_id}/{version}/export.html")
def workflow_export_html(record_id: str, version: int):
    """Export a confirmed workflow as a standalone, print-friendly HTML file."""
    with db() as conn:
        wf = conn.execute(
            "SELECT * FROM workflows WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not wf:
            raise HTTPException(404)

        if wf["status"] != "confirmed":
            raise HTTPException(status_code=409, detail="Export is allowed for confirmed workflows only")

        refs = conn.execute(
            """
            SELECT r.order_index, t.*
            FROM workflow_task_refs r
            JOIN tasks t ON t.record_id=r.task_record_id AND t.version=r.task_version
            WHERE r.workflow_record_id=? AND r.workflow_version=?
            ORDER BY r.order_index
            """,
            (record_id, version),
        ).fetchall()

        readiness = workflow_readiness(
            conn,
            [(r["record_id"], int(r["version"])) for r in refs],
        )
        if readiness != "ready":
            raise HTTPException(status_code=409, detail="Export is allowed only when all referenced Task versions are confirmed")

        primer_rows = conn.execute(_PRIMER_ATTACH_QUERY, (record_id,)).fetchall()

    def _esc(s: str) -> str:
        return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def _steps_rows(steps_data: list[dict[str, Any]]) -> str:
        rows = []
        for i, st in enumerate(steps_data, 1):
            actions = st.get("actions") or []
            actions_html = "<br>".join(_esc(str(a)) for a in actions if str(a).strip()) if actions else "—"
            rows.append(
                f"<tr><td>{i}</td><td>{_esc(str(st.get('text','') or ''))}</td>"
                f"<td>{actions_html}</td>"
                f"<td>{_esc(str(st.get('notes','') or '')) or '—'}</td>"
                f"<td>{_esc(str(st.get('completion','') or '')) or '—'}</td></tr>"
            )
        return "\n".join(rows)

    tasks_html = []
    for r in refs:
        steps = _normalize_steps(_json_load(r["steps_json"]) or [])
        facts = _json_load(r["facts_json"]) or []
        concepts = _json_load(r["concepts_json"]) or []
        deps = _json_load(r["dependencies_json"]) or []

        facts_html = "".join(f"<li>{_esc(str(f))}</li>" for f in facts) if facts else "<li><em>None</em></li>"
        concepts_html = "".join(f"<li>{_esc(str(c))}</li>" for c in concepts) if concepts else "<li><em>None</em></li>"
        deps_html = "".join(f"<li>{_esc(str(d))}</li>" for d in deps) if deps else "<li><em>None</em></li>"
        irr_note = "<p><strong>\u26a0 Irreversible:</strong> This task cannot be undone.</p>" if r["irreversible_flag"] else ""

        tasks_html.append(f"""
<section class="task">
  <h2>Task {r['order_index']}: {_esc(str(r['title']))}</h2>
  {irr_note}
  <p><strong>Outcome:</strong> {_esc(str(r['outcome']))}</p>
  <h3>Facts</h3><ul>{facts_html}</ul>
  <h3>Concepts</h3><ul>{concepts_html}</ul>
  <h3>Dependencies</h3><ul>{deps_html}</ul>
  <h3>Procedure: {_esc(str(r['procedure_name']))}</h3>
  <table>
    <thead><tr><th>#</th><th>Step</th><th>Actions</th><th>Notes</th><th>Completion</th></tr></thead>
    <tbody>{_steps_rows(steps)}</tbody>
  </table>
</section>""")

    trace = f"{record_id}@{version} \u00b7 exported {utc_now_iso().split('T')[0]}"
    html = f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{_esc(str(wf['title']))}</title>
<style>
  body{{font-family:Georgia,serif;max-width:900px;margin:40px auto;padding:0 24px;color:#111;line-height:1.6}}
  h1{{font-size:2em;border-bottom:2px solid #111;padding-bottom:8px}}
  h2{{font-size:1.35em;margin-top:2em;border-bottom:1px solid #ccc;padding-bottom:4px}}
  h3{{font-size:1em;text-transform:uppercase;letter-spacing:.05em;color:#555;margin-top:1.2em}}
  table{{border-collapse:collapse;width:100%;margin-top:8px;font-size:.9em}}
  th,td{{border:1px solid #ccc;padding:6px 10px;vertical-align:top;text-align:left}}
  th{{background:#f5f5f5;font-weight:600}}
  ul{{margin:4px 0;padding-left:20px}}
  .task{{page-break-before:always}}
  .task:first-child{{page-break-before:avoid}}
  .provenance{{font-size:.75em;color:#888;margin-top:3em;border-top:1px solid #eee;padding-top:8px}}
  @media print{{body{{margin:0;padding:16px}}.task{{page-break-before:always}}}}
</style>
</head>
<body>
<h1>{_esc(str(wf['title']))}</h1>
<p><strong>Objective:</strong> {_esc(str(wf['objective']))}</p>
{_primers_html(primer_rows, _esc)}
<h2>Task Overview</h2>
<ol>{"".join(f"<li>{_esc(str(r['title']))}</li>" for r in refs)}</ol>
{"".join(tasks_html)}
<div class="provenance">Trace: {_esc(trace)}</div>
</body>
</html>"""

    filename = f"workflow__{_short_code('WF', record_id)}__v{version}.html"
    return HTMLResponse(
        content=html,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.get("/review", response_class=HTMLResponse)
def review_queue(request: Request, item_type: str = ""):
    # Contributors and admins only. (Admin implicitly has all domains.)
    if request.state.role not in ("contributor", "admin"):
        raise HTTPException(status_code=403, detail="Forbidden: contributor/admin only")

    with db() as conn:
        # Determine authorized domains
        if request.state.role == "admin":
            from ..database import _active_domains
            doms = _active_domains(conn)
        else:
            uid = _user_id(conn, request.state.user)
            dom_rows = conn.execute("SELECT domain FROM user_domains WHERE user_id=?", (uid,)).fetchall() if uid else []
            doms = [str(r["domain"]) for r in dom_rows]

        filter_type = (item_type or "").strip().lower()
        if filter_type not in ("task", "workflow", "assessment"):
            filter_type = ""

        items: list[dict[str, Any]] = []
        if doms:
            dset = {d.strip().lower() for d in doms if d}
            qmarks = ",".join(["?"] * len(doms))

            # Priority bucket for tasks that block workflow progression
            # (derived from workflow-side review flag/linkage).
            blocking_task_refs: set[tuple[str, int]] = set()
            block_rows = conn.execute(
                """
                SELECT wr.task_record_id, wr.task_version
                FROM workflow_task_refs wr
                JOIN workflows w
                  ON w.record_id = wr.workflow_record_id
                 AND w.version = wr.workflow_version
                WHERE w.status='submitted' AND w.needs_review_flag=1
                """
            ).fetchall()
            for br in block_rows:
                blocking_task_refs.add((str(br["task_record_id"]), int(br["task_version"])))

            # Tasks
            t_rows = conn.execute(
                f"SELECT record_id, version, title, status, domain, created_at FROM tasks WHERE status='submitted' AND domain IN ({qmarks})",
                doms,
            ).fetchall()
            for r in t_rows:
                rid = str(r["record_id"])
                ver = int(r["version"])
                items.append(
                    {
                        "type": "task",
                        "record_id": rid,
                        "version": ver,
                        "title": r["title"],
                        "status": r["status"],
                        "domains": [str(r["domain"])],
                        "created_at": str(r["created_at"] or ""),
                        "priority_bucket": 0 if (rid, ver) in blocking_task_refs else 2,
                    }
                )

            # Workflows (domain derived)
            w_rows = conn.execute(
                "SELECT record_id, version, title, status, domains_json, created_at FROM workflows WHERE status='submitted'"
            ).fetchall()
            for r in w_rows:
                wdoms = _normalize_domains(r["domains_json"])
                if dset.intersection(wdoms):
                    items.append(
                        {
                            "type": "workflow",
                            "record_id": r["record_id"],
                            "version": int(r["version"]),
                            "title": r["title"],
                            "status": r["status"],
                            "domains": wdoms,
                            "created_at": str(r["created_at"] or ""),
                            "priority_bucket": 2,
                        }
                    )

            # Assessments
            a_rows = conn.execute(
                "SELECT record_id, version, stem, status, domains_json, created_at FROM assessment_items WHERE status='submitted'"
            ).fetchall()
            for r in a_rows:
                adoms = _normalize_domains(r["domains_json"])
                if dset.intersection(adoms):
                    title = str(r["stem"])[:80]
                    items.append(
                        {
                            "type": "assessment",
                            "record_id": r["record_id"],
                            "version": int(r["version"]),
                            "title": title,
                            "status": r["status"],
                            "domains": adoms,
                            "created_at": str(r["created_at"] or ""),
                            "priority_bucket": 2,
                        }
                    )

            if filter_type:
                items = [it for it in items if str(it.get("type")) == filter_type]

            # Sort: blockers first, then oldest, then type/title.
            items.sort(
                key=lambda it: (
                    int(it.get("priority_bucket", 2)),
                    str(it.get("created_at") or ""),
                    str(it.get("type") or ""),
                    str(it.get("title") or "").lower(),
                )
            )

    return templates.TemplateResponse(request, "review_queue.html", {"items": items, "domains": doms, "item_type": filter_type})


@router.get("/audit", response_class=HTMLResponse)
def audit_list(
    request: Request,
    entity_type: str | None = None,
    record_id: str | None = None,
    limit: int = 200,
):
    require(request.state.role, "audit:view")

    entity_type = (entity_type or "").strip() or None
    record_id = (record_id or "").strip() or None
    limit = max(1, min(int(limit or 200), 1000))

    where: list[str] = []
    params: list[Any] = []
    if entity_type:
        where.append("entity_type=?")
        params.append(entity_type)
    if record_id:
        where.append("record_id=?")
        params.append(record_id)

    sql = "SELECT id, entity_type, record_id, version, action, actor, at, note FROM audit_log"
    if where:
        sql += " WHERE " + " AND ".join(where)
    sql += " ORDER BY at DESC, id DESC LIMIT ?"
    params.append(limit)

    with db() as conn:
        rows = conn.execute(sql, params).fetchall()

    items = [dict(r) for r in rows]

    return templates.TemplateResponse(
        request,
        "audit_list.html",
        {"items": items, "entity_type": entity_type, "record_id": record_id, "limit": limit},
    )


@router.get("/export/task/{record_id}/{version}.json")
def export_task_json(record_id: str, version: int):
    with db() as conn:
        row = conn.execute(
            "SELECT * FROM tasks WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
    if not row:
        raise HTTPException(404)

    if row["status"] != "confirmed":
        raise HTTPException(status_code=409, detail="Export is allowed for confirmed tasks only")

    payload = _task_export_dict(row)
    filename = f"task__{record_id}__v{version}.json"
    return JSONResponse(
        content=payload,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.get("/export/workflow/{record_id}/{version}.json")
def export_workflow_json(record_id: str, version: int):
    with db() as conn:
        wf = conn.execute(
            "SELECT * FROM workflows WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not wf:
            raise HTTPException(404)

        if wf["status"] != "confirmed":
            raise HTTPException(status_code=409, detail="Export is allowed for confirmed workflows only")

        refs = conn.execute(
            """
            SELECT order_index, task_record_id, task_version
            FROM workflow_task_refs
            WHERE workflow_record_id=? AND workflow_version=?
            ORDER BY order_index
            """,
            (record_id, version),
        ).fetchall()

        readiness = workflow_readiness(conn, [(r["task_record_id"], int(r["task_version"])) for r in refs])
        if readiness != "ready":
            raise HTTPException(status_code=409, detail="Export is allowed only when all referenced Task versions are confirmed")

        primers = conn.execute(_PRIMER_ATTACH_QUERY, (record_id,)).fetchall()

    payload = _workflow_export_dict(wf, refs, primers)
    filename = f"workflow__{record_id}__v{version}.json"
    return JSONResponse(
        content=payload,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.get("/workflows/{record_id}/{version}/export.docx")
def workflow_export_docx(request: Request, record_id: str, version: int):
    """Export a confirmed workflow to DOCX (v0 ILT handout).

    Governance rule: exports are allowed for confirmed workflows only, and only when
    all referenced Task versions are confirmed.
    """
    actor = request.state.user

    with db() as conn:
        wf = conn.execute(
            "SELECT * FROM workflows WHERE record_id=? AND version=?",
            (record_id, version),
        ).fetchone()
        if not wf:
            raise HTTPException(404)

        if wf["status"] != "confirmed":
            raise HTTPException(status_code=409, detail="Export is allowed for confirmed workflows only")

        ref_rows = conn.execute(
            """
            SELECT order_index, task_record_id, task_version
            FROM workflow_task_refs
            WHERE workflow_record_id=? AND workflow_version=?
            ORDER BY order_index
            """,
            (record_id, version),
        ).fetchall()

        readiness = workflow_readiness(conn, [(r["task_record_id"], int(r["task_version"])) for r in ref_rows])
        if readiness != "ready":
            raise HTTPException(status_code=409, detail="Export is allowed only when all referenced Task versions are confirmed")

        primer_rows = conn.execute(_PRIMER_ATTACH_QUERY, (record_id,)).fetchall()

        tasks: list[dict[str, Any]] = []
        for r in ref_rows:
            t = conn.execute(
                "SELECT * FROM tasks WHERE record_id=? AND version=?",
                (r["task_record_id"], int(r["task_version"])),
            ).fetchone()
            if not t:
                raise HTTPException(status_code=409, detail="Export failed: referenced task not found")
            if t["status"] != "confirmed":
                raise HTTPException(status_code=409, detail="Export is allowed only when all referenced Task versions are confirmed")
            tasks.append({"order_index": int(r["order_index"]), "row": dict(t)})

        # Build doc
        doc = Document()

        trace = f"Trace: {_short_code('WF', record_id)} v{version} · {utc_now_iso().split('T')[0]}"

        # Footer trace on each section
        for s in doc.sections:
            fp = s.footer.paragraphs[0] if s.footer.paragraphs else s.footer.add_paragraph()
            fp.text = trace

        doc.add_heading(str(wf["title"]), level=1)
        doc.add_paragraph(str(wf["objective"])).style = doc.styles["Normal"]

        if primer_rows:
            doc.add_heading("Pre-reading", level=2)
            for p in primer_rows:
                doc.add_heading(str(p["title"]), level=3)
                if p["summary"]:
                    doc.add_paragraph(str(p["summary"])).style = doc.styles["Normal"]
                if p["explanation"]:
                    doc.add_paragraph(str(p["explanation"])).style = doc.styles["Normal"]
                if p["analogies"]:
                    doc.add_paragraph(f"Analogy: {p['analogies']}").style = doc.styles["Normal"]

        doc.add_heading("Tasks", level=2)
        for t in tasks:
            r = t["row"]
            doc.add_paragraph(f"{t['order_index']}. {r['title']}")

        for t in tasks:
            r = t["row"]
            doc.add_page_break()
            doc.add_heading(f"Task {t['order_index']}: {r['title']}", level=2)
            doc.add_paragraph(f"Outcome: {r['outcome']}")

            steps = _normalize_steps(_json_load(r["steps_json"]) or [])
            doc.add_paragraph(f"Procedure: {r['procedure_name']}")

            table = doc.add_table(rows=1, cols=4)
            hdr = table.rows[0].cells
            hdr[0].text = "Step"
            hdr[1].text = "Actions"
            hdr[2].text = "Notes"
            hdr[3].text = "Completion"

            for st in steps:
                row = table.add_row().cells
                row[0].text = str(st.get("text", "") or "")
                actions = st.get("actions") or []
                row[1].text = "\n".join([str(a) for a in actions if str(a).strip()]) if actions else ""
                row[2].text = str(st.get("notes", "") or "")
                row[3].text = str(st.get("completion", "") or "")

        # Provenance (full UUIDs)
        doc.add_page_break()
        doc.add_heading("Provenance (internal)", level=2)
        doc.add_paragraph(f"Workflow: {record_id} v{version}")
        for r in ref_rows:
            doc.add_paragraph(f"Task: {r['task_record_id']} v{int(r['task_version'])}")

        # Write artifact
        export_id = str(uuid.uuid4())
        ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
        filename = f"workflow__{_short_code('WF', record_id)}__v{version}__{ts}.docx"
        out_path = os.path.join(EXPORTS_DIR, filename)

        doc.save(out_path)
        file_bytes = Path(out_path).read_bytes()
        sha = _sha256_bytes(file_bytes)

        conn.execute(
            "INSERT INTO export_artifacts(id, kind, filename, path, sha256, workflow_record_id, workflow_version, task_refs_json, exported_at, exported_by, retention_days) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
            (
                export_id,
                "docx",
                filename,
                out_path,
                sha,
                record_id,
                int(version),
                _json_dump([{"record_id": r["task_record_id"], "version": int(r["task_version"])} for r in ref_rows]),
                utc_now_iso(),
                actor,
                30,
            ),
        )
        audit("export", export_id, 1, "create", actor, note=f"kind=docx workflow={record_id}@{version}", conn=conn)

    return FileResponse(
        out_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )


@router.get("/workflows/{record_id}/{version}/export.md")
def workflow_export_md(record_id: str, version: int):
    with db() as conn:
        wf = conn.execute(
            "SELECT * FROM workflows WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not wf:
            raise HTTPException(404)

        if wf["status"] != "confirmed":
            raise HTTPException(status_code=409, detail="Export is allowed for confirmed workflows only")

        refs = conn.execute(
            """
            SELECT r.order_index, t.*
            FROM workflow_task_refs r
            JOIN tasks t ON t.record_id=r.task_record_id AND t.version=r.task_version
            WHERE r.workflow_record_id=? AND r.workflow_version=?
            ORDER BY r.order_index
            """,
            (record_id, version),
        ).fetchall()

        readiness = workflow_readiness(
            conn,
            [(r["record_id"], int(r["version"])) for r in refs],
        )
        if readiness != "ready":
            raise HTTPException(status_code=409, detail="Export is allowed only when all referenced Task versions are confirmed")

        primer_rows = conn.execute(_PRIMER_ATTACH_QUERY, (record_id,)).fetchall()

    lines: list[str] = []
    lines.append(f"# {wf['title']}")
    lines.append("")

    lines.append(f"**Objective:** {wf['objective']}")
    lines.append("")

    if primer_rows:
        lines.append("## Pre-reading")
        lines.append("")
        for p in primer_rows:
            lines.append(f"### {p['title']}")
            lines.append("")
            if p["summary"]:
                lines.append(f"_{p['summary']}_")
                lines.append("")
            if p["explanation"]:
                lines.append(str(p["explanation"]))
                lines.append("")
            if p["analogies"]:
                lines.append(f"**Analogy:** {p['analogies']}")
                lines.append("")

    for r in refs:
        steps = _normalize_steps(_json_load(r["steps_json"]))
        facts = _json_load(r["facts_json"]) or []
        concepts = _json_load(r["concepts_json"]) or []
        deps = _json_load(r["dependencies_json"]) or []

        lines.append(f"## Task {r['order_index']}: {r['title']} ({r['record_id']}@{r['version']})")
        if r["status"] != "confirmed":
            lines.append(f"**Task status:** {r['status']} (unconfirmed)")
            lines.append("")
        lines.append("")
        lines.append(f"**Outcome:** {r['outcome']}")
        lines.append("")

        if facts:
            lines.append("**Facts:**")
            for f in facts:
                lines.append(f"- {f}")
            lines.append("")

        if concepts:
            lines.append("**Concepts:**")
            for c in concepts:
                lines.append(f"- {c}")
            lines.append("")

        if deps:
            lines.append("**Dependencies:**")
            for d in deps:
                lines.append(f"- {d}")
            lines.append("")

        lines.append(f"**Procedure:** {r['procedure_name']}")
        lines.append("")

        def _md_cell(s: str) -> str:
            s = (s or "").replace("\n", "<br>")
            s = s.replace("|", "\\|")
            return s

        lines.append("| Step | Actions | Notes | Completion |")
        lines.append("| --- | --- | --- | --- |")
        for st in steps:
            txt = _md_cell(str(st.get("text", "") or ""))
            notes = _md_cell(str(st.get("notes", "") or "").strip()) or "—"
            actions = st.get("actions") or []
            actions_txt = _md_cell("<br>".join([str(a) for a in actions if str(a).strip()])) if actions else "—"
            comp = _md_cell(str(st.get("completion", "") or "")) or "—"
            lines.append(f"| {txt} | {actions_txt} | {notes} | {comp} |")
        lines.append("")

    md = "\n".join(lines)
    filename = f"workflow__{record_id}__v{version}.md"
    return HTMLResponse(
        content=md,
        media_type="text/markdown",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


def _build_presentation_payload(conn, workflow_record_id: str, workflow_version: int) -> dict:
    """Fetch full workflow content for presentation output. Used by both the prompt page and the fetch endpoint."""
    wf = conn.execute(
        "SELECT * FROM workflows WHERE record_id=? AND version=?",
        (workflow_record_id, workflow_version),
    ).fetchone()
    if not wf:
        raise HTTPException(status_code=404, detail="Workflow not found")

    refs = conn.execute(
        "SELECT task_record_id, task_version, order_index "
        "FROM workflow_task_refs WHERE workflow_record_id=? AND workflow_version=? "
        "ORDER BY order_index",
        (workflow_record_id, workflow_version),
    ).fetchall()
    tasks = []
    for ref in refs:
        t = conn.execute(
            "SELECT record_id, version, title, outcome, steps_json, domain "
            "FROM tasks WHERE record_id=? AND version=?",
            (ref["task_record_id"], ref["task_version"]),
        ).fetchone()
        if t:
            tasks.append({
                "order_index": ref["order_index"],
                "record_id": t["record_id"],
                "version": int(t["version"]),
                "title": t["title"],
                "outcome": t["outcome"],
                "domain": t["domain"],
                "steps": _json_load(t["steps_json"] or "[]") or [],
            })

    primer_rows = conn.execute(_PRIMER_ATTACH_QUERY, (workflow_record_id,)).fetchall()
    primers = [
        {
            "record_id": p["record_id"],
            "version": int(p["version"]),
            "title": p["title"],
            "summary": p["summary"],
            "explanation": p["explanation"],
            "analogies": p["analogies"],
        }
        for p in primer_rows
    ]

    return {
        "workflow": {
            "record_id": wf["record_id"],
            "version": int(wf["version"]),
            "title": wf["title"],
            "objective": wf["objective"],
            "domains": _normalize_domains(wf["domains_json"]),
            "tags": _json_load(wf["tags_json"] or "[]") or [],
        },
        "tasks": tasks,
        "primers": primers,
    }


# ---------------------------------------------------------------------------
# Export package (ZIP): format-specific LLM prompt bundles
# ---------------------------------------------------------------------------

_EXPORT_FORMATS: dict[str, str] = {
    "ilt_slides": "ILT Slides",
    "ilt_facilitators_guide": "ILT Facilitators Guide",
    "self_paced_html": "Self-Paced HTML",
    "helpsheet": "Helpsheet",
}

_FORMAT_PROMPTS: dict[str, str] = {
    "ilt_slides": """\
# Export Package: ILT Slides

You are generating a slide deck for an instructor-led training (ILT) session from a Blueprinted workflow export package.

This ZIP contains:
- `data.json`: full structured workflow and task content
- `images/`: screenshots and diagrams (referenced by task_record_id)
- `logos/`: client logo file (e.g. `client-logo.png`); use it if present
- `styles.md`: visual design specification (colours, typography, spacing, buttons)
- `editorial.md`: writing rules, product name spelling, and tone guidance

Read `styles.md` and `editorial.md` before generating output. They are authoritative. Where any instruction in this prompt conflicts with those files, the guidance in those files takes precedence.

If a logo file is present in `logos/`, place it on the title slide (top-left or top-right corner, maintaining native aspect ratio) and in the footer area of every slide alongside the workflow record_id and version. If no logo is present, leave those positions empty.

---

## Understanding the data model

Before generating output, understand what each entity is.

### Workflow

A Workflow is a composite outcome made of ordered Tasks. It has:

- `title`: learner-facing name
- `objective`: the organisation-defined outcome the workflow produces
- `record_id`: unique identifier
- `version`: content version number
- `tasks[]`: ordered list of Task references

Task order is strict and reflects capability dependencies only.

### Task

A Task is an atomic, self-contained unit of performance. It produces exactly one observable outcome. A Task has:

- `title`: verb-driven learner-facing name
- `outcome`: the observable result when the task is complete
- `facts[]`: literal information the learner must know before executing the task (discrete, verifiable statements)
- `concepts[]`: the mental models required to understand WHY the steps work
- `procedure_name`: the name of the step sequence
- `steps[]`: ordered atomic instructions
- `dependencies[]`: conditions or access requirements that must be true before this task can be executed
- `task_assets[]`: optional media objects (see Media below)
- `irreversible`: boolean; true means the task cannot be undone

Facts without Concepts produce rote behaviour. Concepts without Facts produce abstraction. Neither produces capability without the Procedure. Present all three.

### Step

Each Step inside a Task has:

- `text`: (required) the primary action in imperative form
- `actions[]`: (optional) sub-steps describing HOW to execute the step
- `completion`: (required) the observable signal confirming the step is complete
- `notes`: (optional) additional context or warnings
- `screenshots[]`: (optional) image filenames in `images/{task_record_id}/`

### Primer

A Primer is pre-reading material attached to the workflow. Primers are not tasks and do not have steps. Each primer has:

- `title`, `summary`, `explanation`

Primers establish orientation and context before the learner encounters task procedures.

### Media (task_assets)

Each task asset object has:

- `url`: fully-qualified external URL
- `type`: one of: `video`, `demo`, `image`, `audio`, `module`, `link`
- `label`: short descriptive label

Reference media assets as clearly labelled links with the `label` text. Do not attempt to embed external URLs. For `type: demo` (Storylane), note it as an interactive demo link. For `type: video`, note it as a video resource. Include media references on the context or procedure slide for the relevant task.

---

## Visual-first design principle

Slides must lead with visuals, not text. Every time you can replace a bullet list with a diagram, icon grid, visual grouping, or table layout: do it. Learners in a live session respond to visual stimuli, not walls of text. Apply this principle slide by slide:

- **Facts**: do not list them as plain bullets. Represent them as a grid of labelled cards or a two-column icon-plus-text layout. If there are 3 or more facts, use a card grid (3-up or 4-up). Each card: short label at top, fact statement below.
- **Concepts**: use a visual metaphor where possible: a simple diagram, a before/after comparison, or a cause-effect layout. If the concept is abstract, a labelled diagram with arrows is better than a paragraph.
- **Primers** (pre-reading slide): render each primer as a distinct visual card showing title and one-line summary. Never list them as plain text.
- **Procedures**: use a numbered step flow (horizontal or vertical) rather than a plain numbered list. Show step number, action, and completion signal as three discrete visual zones per step. Split across slides if more than 4 steps.
- **Summary slide**: use a visual timeline or numbered card row showing each task title: not a plain bullet list.
- Slide text should act as a label for a visual, not as the primary content. If a slide has more than 25 words of body text, redesign it.

Express all visual layouts in Markdown using tables, block structures, or clearly labelled ASCII layout diagrams so a human or tool can recreate them faithfully in PowerPoint.

---

## Your output

Produce a complete slide deck in Markdown format (one `---` separator per slide).
Use the following structure:

1. **Title slide**: workflow title; client logo if present in `logos/` (top corner, native aspect ratio); and a visually prominent **social contract statement**: "By the end of this session you will have [workflow.objective]": this must be large, standalone, and impossible to miss; include a simple visual motif (e.g. a relevant icon or abstract shape description)
2. **Learning objectives slide**: task outcomes as a visual card row, not a plain bullet list; one card per task showing the outcome in one line
3. **Pre-reading slide** (if primers present): one visual card per primer showing title and one-line summary
4. For each task in order:
   a. **Section divider slide**: task title only, large centred heading; include a brief visual context cue (e.g. icon, colour block description, or thematic illustration note)
   b. **Context slide**: facts as a card grid or icon-list layout; one line of concept framing below; list any dependencies as a short visual checklist; include media asset references if present
   c. **Procedure slide(s)**: steps as a numbered visual flow; split across slides if more than 4 steps; note irreversible steps with a caution marker; include `actions` as indented sub-items
   d. **Screenshot slide(s)**: one slide per screenshot image if present; caption with the step it relates to; reference path: `images/{task_record_id}/{filename}`; always insert images at their native aspect ratio: never stretch, crop, or scale to fill the slide; use "fit" sizing with centred alignment and empty space around the image if needed
   e. **Completion slide**: "How do you know you're done?": render completion signals as a visual checklist, not a plain bullet list
5. **Summary slide**: task titles in a visual numbered card row or timeline layout
6. **Closing slide**: a visually prominent **closing statement**: "You have just [workflow.objective]": large, standalone, and distinct from the summary slide; this is the closing of the social contract opened on the title slide; follow with "Questions?" beneath it

---

## Style rules
- Tone: professional, direct, second-person ("you will", "select", "click")
- Maximum 25 words of body text per slide; prefer visual layouts over text
- No plain bullet lists: always use a card grid, icon list, flow diagram, or table instead
- Speaker notes: add below each slide (under a `> Notes:` blockquote) with talking points, timing cues, and facilitator tips
- Do not use em dashes (—) anywhere. Hard rule, no exceptions. Use commas or colons instead
- Include the workflow record_id and version in the footer of the title slide
- Ensure content is clear, concise, and visually engaging for live training; adhere to learning best practices; ensure sections are properly linked to create a throughline across the whole product

---

## Data reference

    workflow.title
    workflow.objective
    workflow.record_id
    workflow.version
    workflow.tasks[]
    workflow.primers[]          // optional

    tasks[].title
    tasks[].outcome
    tasks[].facts[]
    tasks[].concepts[]
    tasks[].procedure_name
    tasks[].steps[]
    tasks[].dependencies[]
    tasks[].task_assets[]       // optional; objects with url, type, label
    tasks[].irreversible        // boolean, optional

    steps[].text
    steps[].actions[]           // optional
    steps[].completion
    steps[].notes               // optional
    steps[].screenshots[]       // optional; filenames only
""",

    "ilt_facilitators_guide": """\
# Export Package: ILT Facilitators Guide

You are generating a facilitator guide for an instructor-led training (ILT) session from a Blueprinted workflow export package.

This ZIP contains:
- `data.json`: full structured workflow and task content
- `images/`: screenshots and diagrams (reference these to direct facilitators' attention)
- `logos/`: client logo file (e.g. `client-logo.png`); use it if present
- `styles.md`: visual design specification (colours, typography, spacing, buttons)
- `editorial.md`: writing rules, product name spelling, and tone guidance

Read `styles.md` and `editorial.md` before generating output. They are authoritative. Where any instruction in this prompt conflicts with those files, the guidance in those files takes precedence.

If a logo file is present in `logos/`, place it on the document cover page (top-left or top-right, maintaining native aspect ratio) and in the footer of every page alongside the workflow record_id and version. If no logo is present, leave those positions empty.

---

## Understanding the data model

Before generating output, understand what each entity is.

### Workflow

A Workflow is a composite outcome made of ordered Tasks. It has:

- `title`: learner-facing name
- `objective`: the organisation-defined outcome the workflow produces
- `record_id`: unique identifier
- `version`: content version number
- `tasks[]`: ordered list of Task references

Task order is strict and reflects capability dependencies only.

### Task

A Task is an atomic, self-contained unit of performance. It produces exactly one observable outcome. A Task has:

- `title`: verb-driven learner-facing name
- `outcome`: the observable result when the task is complete
- `facts[]`: literal information the learner must know before executing the task (discrete, verifiable statements)
- `concepts[]`: the mental models required to understand WHY the steps work
- `procedure_name`: the name of the step sequence
- `steps[]`: ordered atomic instructions
- `dependencies[]`: conditions or access requirements that must be true before this task can be executed
- `task_assets[]`: optional media objects (see Media below)
- `irreversible`: boolean; true means the task cannot be undone

Facts without Concepts produce rote behaviour. Concepts without Facts produce abstraction. Neither produces capability without the Procedure. Present all three.

### Step

Each Step inside a Task has:

- `text`: (required) the primary action in imperative form
- `actions[]`: (optional) sub-steps describing HOW to execute the step
- `completion`: (required) the observable signal confirming the step is complete
- `notes`: (optional) additional context or warnings
- `screenshots[]`: (optional) image filenames in `images/{task_record_id}/`

### Primer

A Primer is pre-reading material attached to the workflow. Primers are not tasks and do not have steps. Each primer has:

- `title`, `summary`, `explanation`

Primers establish orientation and context before the learner encounters task procedures.

### Media (task_assets)

Each task asset object has:

- `url`: fully-qualified external URL
- `type`: one of: `video`, `demo`, `image`, `audio`, `module`, `link`
- `label`: short descriptive label

Reference media assets in the facilitator notes for the relevant task. Include the label text and note the resource type (video, interactive demo, etc.) so the facilitator knows when to direct participants to it. Do not embed URLs in participant-facing content.

---

## Before you generate

Before drafting any content, search online for common mistakes, failure points, and misconceptions that people new to this process typically encounter. Use the workflow title, task titles, and domain as search terms. Incorporate what you find into the facilitator notes: specifically in the Talking points and Walkthrough steps for each task. Cite or paraphrase specific failure patterns; do not use generic advice.

---

## Your output

Produce a detailed facilitators guide as a Microsoft Word document (DOCX). Use heading styles (Heading 1, Heading 2, Heading 3), bulleted and numbered lists, and bold/italic formatting as appropriate: do not produce plain text or Markdown. Structure:

### Front matter
- Workflow title and objective
- **Social contract statement: opening**: render this prominently and in full as a standalone block: "By the end of this session participants will have [workflow.objective]": this is the facilitator's commitment to the learner and must appear verbatim at the top of the front matter, before any other content
- Total estimated duration (estimate 5 minutes per step across all tasks, plus 10 minutes intro and 5 minutes close)
- Materials checklist (projector, participant workbooks, access to the software, etc.)
- Learning objectives (one per task outcome)
- Pre-reading list (if primers are present): title and one-line summary for each primer

### For each task (in order)

Use this repeating structure:

**Task N: [Title]**
- *Duration estimate:* [N] minutes
- *Key message:* one sentence summarising why this task matters (from concepts)
- *Preparation:* what the facilitator must have ready before starting this task; note any media assets to cue up
- *Talking points:* expand on the facts and concepts: 3-5 bullet points with fuller explanation than the slides show
- *Walkthrough steps:* numbered list mirroring the task steps; add facilitator notes on what to watch for, when to pause, and: drawing on your pre-generation research: the specific failure patterns and misconceptions that newcomers to this process commonly encounter at each step; note irreversible steps with a caution marker
- *Discussion prompt:* one open question relating the task to participants' real work

### Close
- **Social contract statement: closing**: render this prominently as a standalone block before anything else in the close section: "Participants have just [workflow.objective]": this closes the commitment made at the start; it must be explicit and standalone, not buried in a summary paragraph
- Summary of the session (task titles and one-line outcomes)
- Next steps / recommended follow-up
- Evaluation prompt (what to ask participants to assess learning)

---

## Style rules
- Tone: direct, collegial: written for a subject-matter expert facilitating peers
- Use second-person for participant instructions, third-person for facilitator guidance
- Do not use em dashes (—) anywhere. Hard rule, no exceptions. Use commas or colons instead
- Include workflow record_id and version in the document footer
- Images: insert at native size and aspect ratio; never stretch or distort; if an image is wider than the text column, scale it down proportionally (maintain aspect ratio); do not set explicit height values independently of width
- Ensure that the content is clear, concise, and visually engaging for live training; adhere to learning best practices; ensure sections are properly linked to create a throughline across the whole product

---

## Data reference

    workflow.title
    workflow.objective
    workflow.record_id
    workflow.version
    workflow.tasks[]
    workflow.primers[]          // optional

    tasks[].title
    tasks[].outcome
    tasks[].facts[]
    tasks[].concepts[]
    tasks[].procedure_name
    tasks[].steps[]
    tasks[].dependencies[]
    tasks[].task_assets[]       // optional; objects with url, type, label
    tasks[].irreversible        // boolean, optional

    steps[].text
    steps[].actions[]           // optional
    steps[].completion
    steps[].notes               // optional
    steps[].screenshots[]       // optional; filenames only
""",

    "self_paced_html": """\
# Export Package: Self-Paced HTML Course

You are generating a self-paced e-learning course in HTML format from a Blueprinted workflow export package.

This ZIP contains:

- `data.json`: full structured workflow and task content
- `images/`: screenshots and diagrams (referenced by task_record_id)
- `logos/`: client logo file (e.g. `client-logo.png`); use it if present
- `styles.md`: visual design specification (colours, typography, spacing, buttons)
- `editorial.md`: writing rules, product name spelling, and tone guidance

Read both files before writing any HTML or CSS. They are authoritative. Where
any instruction in this prompt conflicts with `styles.md` or `editorial.md`,
the guidance in those files takes precedence.

---

## Important: two-phase generation: do not skip this step

**Do not generate any HTML yet.**

Before producing the course, you must first generate the "In practice" scenario text for each task and have it reviewed by the person who uploaded this package.

**Phase 1: scenario review (do this first):**

Before writing any scenario, invent a single character who will appear in every task's scenario throughout the course. Give them a first name, a job title, a named organisation, a department, and a location. This character must be consistent across all tasks: the learner follows the same person from task 1 to the final task. Output the character profile first, as a single short paragraph, so it can be reviewed alongside the scenarios.

When choosing the character's details, select randomly from the approved lists below. Do not use your own suggestions or training-data defaults — always pick from these lists.

**First names:** Tom, Ben, Oliver, Ethan, Liam, James, Nate, Leo, Finn, Glen, Craig, Aaron, Wei, Yusuf, Kenji, Mateo, Kofi, Emma, Sophie, Claire, Rachel, Laura, Hannah, Megan, Amy, Fiona, Naomi, Jade, Zoe, Beth, Aisha, Ingrid, Ananya, Alex, Sam, Jordan, Taylor, Casey, Morgan, Riley, Jamie, Quinn, Blake, Ryan, Dylan, Connor, Derek, Shane, Karl, Ivan, Tariq, Soren, Felipe, Riku, Caleb, Owen, Brett, Dean, Kate, Molly, Nina, Tara, Leah, Chloe, Helen, Petra, Yuki, Fatima, Siobhan, Renata, Lucia, Drew, Robin, Reece, Hayden, Finley, Sage, Avery, Cameron, Dana, Kendall

**Surnames:** Clarke, Murphy, Nguyen, Thomson, Okafor, Fernandez, Yamamoto, Hassan, Reid, Walsh, Kowalski, Johansson, Dube, Petrov, O'Brien, Chen, Svensson, Mbeki, Nakamura, Patel, Turner, Mitchell, Cooper, Bailey, Hughes, Watson, Morrison, Fischer, Andersen, Diallo, Santos, Kimura, Brennan, Osei, Larsson, Hoffmann, Chowdhury, Reyes, Flanagan, Varga

**Organisations:** Meridian Systems, Corepath Solutions, Nexvale Group, Trident Infrastructure, Bridgeway Technologies, Ironshore Networks, Crestwood Partners, Stonebridge Consulting, Verity Group, Halcyon Technologies, Thornfield Solutions, Archway Networks, Clearwater Systems, Fairmont Industries, Harborview Group, Pinnacle Group, Redstone Technologies, Helix Solutions, Cascade Partners, Apex Infrastructure, Summit Networks, Keystone Consulting, Horizon Group, Tidal Systems, Driftwood Partners, Pillar Technologies, Bedrock Systems, Lodestar Group, Vantage Point Consulting, Crossroads Solutions

**Locations:** Manchester, Glasgow, Bristol, Leeds, Liverpool, Edinburgh, Cardiff, Dublin, Birmingham, Sheffield, Newcastle, Nottingham, Auckland, Wellington, Christchurch, Dunedin, Brisbane, Perth, Adelaide, Melbourne, Toronto, Vancouver, Calgary, Ottawa, Montreal, Winnipeg, Singapore, Amsterdam, Johannesburg, Nairobi, Boston, Philadelphia, Baltimore, Pittsburgh, Atlanta, Charlotte, Nashville, Raleigh, Tampa, New Orleans, Chicago, Detroit, Minneapolis, Columbus, Indianapolis, Kansas City, Milwaukee, Cleveland, St. Louis, Dallas, Houston, Austin, Phoenix, Denver, Salt Lake City, Portland, Seattle, Las Vegas, Miami

Then, for each task in `data.json`, in order, output the proposed "In practice" scenario as plain text in the chat: not as HTML. Each scenario must feature the same character established above. Follow the scenario content rules in the "Real-world scenario" section below. After outputting all scenarios, stop and ask: "Please review the character and scenarios above for accuracy and suitability. Confirm they are appropriate for your audience, or provide any corrections or replacements, and I will proceed to generate the full course."

Wait for an explicit confirmation or correction before proceeding.

**Phase 2: full course generation:**

Once the scenarios have been confirmed (or corrected), generate the complete HTML course using the confirmed scenario text for each "In practice" card. Do not regenerate or alter confirmed scenarios during this phase.

If a logo file is present in `logos/`, display it on the cover page (top-left or top-right of the header, maintaining native aspect ratio, max-height 60px) and in the footer of every page alongside the workflow record_id and version. Reference it with a relative path: `logos/client-logo.{ext}`. If no logo is present, leave those positions empty.

---

## Understanding the data model

Before generating output, understand what each entity is.

### Workflow

A Workflow is a composite outcome made of ordered Tasks. It has:

- `title`: learner-facing name
- `objective`: the organisation-defined outcome the workflow produces
- `record_id`: unique identifier
- `version`: content version number
- `tasks[]`: ordered list of Task references

A workflow does not contain steps, prerequisites, or learning-sequence guidance. Task order is strict and reflects capability dependencies only.

### Task

A Task is an atomic, self-contained unit of performance. It produces exactly one observable outcome. A Task has:

- `title`: verb-driven learner-facing name
- `outcome`: the observable result when the task is complete
- `facts[]`: literal information the learner must know before executing the task (not opinions, not background; facts are discrete, verifiable statements)
- `concepts[]`: the mental models required to understand WHY the steps work (concepts answer "what is this" and "why does it behave this way")
- `procedure_name`: the name of the step sequence
- `steps[]`: ordered atomic instructions (see below)
- `dependencies[]`: conditions, completed tasks, or access requirements that must be true before this task can be executed
- `task_assets[]`: optional media objects (see Media below)
- `irreversible`: boolean flag; true means the task cannot be undone

Facts, Concepts, and Steps are owned entirely by their Task. They are not shared between Tasks.

Facts without Concepts produce rote behaviour. Concepts without Facts produce abstraction. Neither produces capability without the Procedure. The three elements work together: the learner must know the facts, understand the concepts, and follow the steps.

### Step

Each Step inside a Task has:

- `text`: (required) the primary action in imperative form
- `actions[]`: (optional) sub-steps describing HOW to execute the step in a specific tool or environment
- `completion`: (required) the observable signal confirming the step is complete
- `notes`: (optional) additional context or warnings
- `screenshots[]`: (optional) list of image filenames in `images/{task_record_id}/`

Each Step describes exactly one action. Completion signals are objective, not subjective ("The success banner appears" not "You feel confident").

### Primer

A Primer is pre-reading material attached to the workflow. Primers establish orientation and context before the learner encounters task procedures. They are not tasks and do not have steps. Each primer has:

- `title`: heading for the primer card
- `summary`: one or two sentences describing what the primer covers
- `explanation`: full text content, rendered as formatted paragraphs

Primers reduce cognitive load in the task pages by front-loading conceptual orientation. Render them before the task sequence begins.

### Media (task_assets)

Each task asset object has:

- `url`: fully-qualified external URL
- `type`: one of: `video`, `demo`, `image`, `audio`, `module`, `link`
- `label`: short descriptive label

For `type: video`, render an embedded video player or a clearly labelled link card using the label text.

For `type: demo`, render a Storylane demo link card with an icon and the label text. Do not embed as iframe (Storylane blocks this). Open in a new tab.

For other types, render an appropriate inline link card.

Task-level assets appear after the Concepts section and before the procedure table.

---

## Interactive-first design principle

This HTML output must be interactive and visually driven throughout. Avoid rendering content as static text where an interactive or visual treatment is possible. Apply this principle to every section:

- **Primers**: do not render primer explanation text as prose paragraphs. Before rendering each primer, read its explanation and classify what kind of content it contains, then choose the matching visual treatment:
  - **Grouped concepts or categories** (e.g. "there are three types of X", "the four components are"): render as an interactive card group: one card per item, laid out in a responsive grid. Each card shows the item name/label on the front. On click/tap the card expands or flips to reveal the full description.
  - **A process, sequence, or lifecycle** (e.g. steps, stages, phases, a workflow): render as an animated step-flow or timeline: numbered nodes connected by lines or arrows, each node expanding on click to show detail. Animate the progression left-to-right or top-to-bottom using CSS transitions triggered on page load or on scroll into view.
  - **A comparison or set of options** (e.g. "option A vs option B", trade-offs, when to use which): render as a side-by-side comparison card layout with a clear visual divider.
  - **A definition or single concept with depth**: render as a large visual card with a bold title, a one-line summary visible immediately, and the full explanation revealed via an expand toggle.
  - **Mixed content**: split the explanation into its logical parts and apply the appropriate visual treatment to each part independently.
  Use HTML5, CSS animations, and vanilla JS freely to express these layouts. Animated SVG paths, CSS keyframe animations, and interactive state transitions are all permitted and encouraged where they serve the content. No external libraries.
- **Configuration and option steps**: Before rendering any task's procedure section, inspect the task content for two special step patterns.

  A **configuration step** is one where the learner sets multiple parameter values (e.g. schedule type, retention count, destination path, job mode) that together produce a combined outcome. Signals: the step's `actions[]` lists configuration fields, and the `completion` describes a result that depends on multiple chosen values; notes contain "if X then Y" relationships describing how setting values interact.

  If a configuration step is detected, render a **Configuration Explorer** immediately above the procedure table for that task:
  - One selector per configurable parameter: use `<select>` dropdowns or a segmented button group
  - Populate each selector's options from the step's actions text and notes; derive concise human-readable labels for each option (not raw technical tokens)
  - Below all selectors, display a live outcome sentence in a visually distinct callout (accent border, light background tint). Derive a sentence template from the task content: e.g. "Your [job type] backup will run [schedule] and retain [retention] restore points." Update it instantly on any selector change via vanilla JS event listeners
  - The procedure table remains fully visible below the Configuration Explorer: the explorer is an orientation aid, not a replacement for the steps

  An **option-selection step** is one where the learner must choose between 2–5 discrete modes, types, or approaches before executing the procedure. Signals: a concept or fact states "there are N types of X", describes mutually exclusive options, or the procedure has conditional paths depending on which approach is chosen (e.g. different destination types, agent vs agentless modes, job types).

  If an option-selection pattern is detected, render **Option Selector Cards** above the Facts section for that task:
  - A row (or wrapping grid) of visually distinct card buttons: one per option: styled with the accent palette from `styles.md`
  - Clicking a card: highlights it (accent border + background tint), slides open a detail panel directly below showing the option's description, implications, and which procedure steps apply; all other panels collapse
  - Only one card selected at a time; default to the first card selected on page load
  - If procedure steps are conditional on the selected option, dim non-applicable rows (opacity: 0.3) and show applicable rows at full opacity; update dynamically on card change
  - Keyboard accessible: left/right arrow keys move between cards, Enter/Space selects

  **Content-type quick reference**: use this table to decide which pattern applies:

  | Signal in data.json | Pattern |
  |---|---|
  | Step actions list config fields + completion describes combined outcome | Configuration Explorer |
  | Concept/fact describes N mutually exclusive types, modes, or approaches | Option Selector Cards |
  | Primer: grouped concepts or categories | Interactive card group |
  | Primer: process, sequence, or lifecycle | Animated step-flow / timeline |
  | Primer: comparison or trade-off | Side-by-side comparison layout |
  | Primer: single concept with depth | Expandable definition card |
  | Facts: 3 or more | Flip cards |
  | Facts: fewer than 3 | Highlighted callout cards |

  **Icons**: use inline SVG for all icons (gear, calendar, checkmark, warning triangle, arrow). Never load an icon library or any resource from a CDN. Embed SVG markup directly in the HTML. Unicode characters (✓ ⚠ →) are acceptable as a fallback for simple inline indicators only.

- **Facts**: render as interactive flip cards. Front face shows a short label or key phrase (derived from the fact). Back face reveals the full fact statement on click/tap. Use vanilla JS: no external libraries. If there are fewer than 3 facts, use highlighted callout cards instead of flip cards.
- **Concepts**: render as an accordion: each concept is a collapsed row showing the first sentence as a preview; clicking expands the full explanation. Do not render concepts as static prose paragraphs.
- **Dependencies**: render as a visual checklist with checkbox-style indicators (purely decorative, not functional), not a plain bullet list.
- **Procedure steps**: the existing table format is acceptable, but enhance it: highlight the active/focused row on hover, and use a distinct visual treatment for the completion signal column (coloured badge or pill, not plain text).
- **Overall**: every page should feel designed, not word-processed. Use whitespace, colour, and interaction to guide the learner's attention. Never render a section as a plain paragraph list when a card, accordion, or interactive component is available.

All interactive components must be implemented in vanilla JS with embedded CSS. No external libraries or CDN dependencies of any kind.

---

## Output

Produce a ZIP file. The ZIP must contain:

```
index.html          ← the course (all CSS and JS embedded inline)
images/             ← copy the images/ folder from this ZIP as-is
logos/              ← copy the logos/ folder from this ZIP as-is
```

Do not inline images as base64. All `<img>` src attributes and logo references must use relative paths that resolve within the ZIP (e.g. `images/{task_record_id}/{filename}`, `logos/client-logo.png`). When a learner extracts the ZIP and opens `index.html` in a browser, all images and logos must load correctly. CSS and JavaScript remain embedded in `index.html`: no external CDN dependencies.

---

## Page model: critical constraint

This output is **not a single-page application**. Do not build a monolithic HTML document where all content exists in the DOM simultaneously and sections are shown or hidden by toggling `display`, `visibility`, `opacity`, `height`, `pointer-events`, or any other CSS or JS property. That approach produces a slow, fragile document that breaks screen readers, search, and print.

Instead, implement a **page-based navigation model**:

- The document is divided into discrete named page views, each a separate `<section>` or `<div>` with a unique `id`
- Only the active page is visible; all others have `display: none`
- Navigating to a page sets the target to `display: block` (or `flex`/`grid` as appropriate) and hides all others: a single JS function handles this for every transition
- Each page transition scrolls to the top of the content area and applies a short CSS opacity fade (150ms)
- The active page `id` is reflected in the URL hash so the browser back/forward buttons work and the page can be bookmarked or reloaded
- The sidebar highlights the currently active page

There is no scrolling between pages. A learner should never need to scroll to find a navigation button. Prev / Next controls are always visible within the viewport.

---

## Page structure

### 1. Cover page

This page forms the opening of the social contract with the learner. It must make the commitment unmistakable.

- Client logo if present in `logos/`: top of page, native aspect ratio, max-height 60px
- Workflow title (large heading)
- **Social contract statement: opening**: render this as a large, visually dominant, standalone element: not a paragraph, not a subtitle. The exact text: **"By the end of this course you will have [workflow.objective]"**: style it as the most prominent thing on the page after the title; use a distinct background block, large type, or a visually separated callout so it cannot be overlooked or confused with supporting text
- Task titles as a visual numbered card row: display only, not clickable links. Do not add any navigation links or anchors to this list.
- Progress bar showing 0 / total tasks
- A single "Start →" button: this is the only navigation control on this page. Do not add any other buttons, links, or shortcuts (no "read the primers first", no "skip to task N", no anchor links into the course).

### 2. Character introduction page

This page introduces the course character established in Phase 1. It appears after the cover and before the primers, giving the learner a human anchor for the scenarios they will encounter in each task.

- Heading: "Who you'll be following" (or a similarly inviting heading)
- Display the character's name, job title, organisation, department, and location as a visually distinct profile card: not a plain paragraph
- 2–3 sentences of narrative prose introducing the character's professional context and the circumstances that have brought them to this course
- "Continue →" button that navigates to the first primer (if primers exist) or the first task

### 3. Primers (one page per primer: render only if primers[] is non-empty)

Each primer is its own page. Do not combine primers onto a single page. If there are three primers, there are three primer pages.

Each primer page:

- Primer title as the page heading; primer summary as a subtitle
- The primer explanation rendered using the classified visual treatment defined in the Interactive-first design principle: never as plain prose
- "← Back" and "Next →" navigation buttons; the final primer's Next button navigates to the first task

### 4. Task pages

Each task is split across three consecutive pages. Do not combine these onto a single page. A learner navigates through all three pages for each task before moving to the next task.

---

#### Page A: Task introduction

- Task number and title (e.g. "Task 2 of 5: Configure backup schedule")
- Outcome: styled callout: "When you complete this task: [outcome text]"
- If `irreversible` is true: prominent warning banner: "This task includes irreversible changes. Review all steps before proceeding."

**"In practice" scenario** (content rules in the Real-world scenario section below):

In Phase 2, render the confirmed scenario text on this page as a visually distinct collapsible card:

- Heading: "In practice"
- A subtle icon (inline SVG: a building, person at desk, or similar; no CDN)
- The scenario text rendered as short paragraphs with comfortable line spacing
- A "Show / Hide" toggle; default state is expanded
- Visually distinct from the surrounding content: different background tint or left border in the accent colour

**Dependencies** (render only if dependencies[] is non-empty):

- Heading: "Before you start"
- Visual checklist with decorative checkbox indicators: not a plain bullet list

**Facts:**

- Heading: "What you need to know"
- Render as interactive flip cards (see Interactive-first design principle). Facts are discrete and literal: each card represents exactly one fact.

- "Next: Why this works →" button navigates to Page B

---

#### Page B: Concepts

- Task number and title repeated as a sub-heading so the learner knows where they are
- Heading: "Why this works"
- Render all concepts as an accordion (see Interactive-first design principle). Each row shows a short preview; clicking expands the full explanation.
- Task-level media (render only if task_assets[] is non-empty): heading "Supporting resources"; render each asset using the Media rules above

- "← Back" button returns to Page A; "Next: The procedure →" button navigates to Page C

---

#### Page C: Procedure

- Task number and title repeated as a sub-heading
- Configuration Explorer and/or Option Selector Cards if applicable (see Interactive-first design principle)
- The full procedure step-through (spec below)

**g. Procedure step-through**

Do not render the procedure as a table. Render it as a full-width interactive step-through where only one step is visible at a time and imagery is the primary element.

**Structure of the step-through container:**

At the top of the procedure section, display the heading ([procedure_name]) and a segmented progress track: a thin horizontal bar divided into N equal segments (one per step), with the current segment filled in the accent colour. Segment count updates as the learner advances.

Each individual step panel uses a **responsive two-column layout**:

- **Desktop (≥ 700px)**: CSS grid with two columns: text column on the left (`1fr`), image column on the right (`1fr`). Both columns are full height, vertically aligned to the top. The step navigation controls span both columns at the bottom.
- **Mobile (< 700px)**: single column, stacked: image first, then text content below. Use a `@media (max-width: 699px)` breakpoint to override the grid to a single column and reorder so the image pane appears before the text column (`order: -1` on the image pane).

**Left column (text)**: in this order:

1. **Step header**:
   - A large step-number badge (e.g. a circle or pill containing "3")
   - step.text rendered as a prominent heading: this is the primary instruction the learner acts on

2. **Actions** (render only when step.actions[] is non-empty):
   - A numbered sub-list of the action sub-steps, indented beneath the step heading

3. **Concept in context**: see spec below (item 4)

4. **Irreversible warning**, **Notes**, **Completion signal**: see specs below (items 5–7)

**Right column (image)**: render only when the step has screenshots:

- Apply to every `<img>`: `width: 100%; height: auto; display: block; object-fit: contain; border-radius: 6px`. Never set both width and height explicitly. Never use `object-fit: cover` or `object-fit: fill`. Images must always render at their native aspect ratio.
- Image `src`: `images/{task_record_id}/{filename}`. Image `alt`: step.text.
- Every screenshot opens the lightbox on click; apply `cursor: zoom-in`.
- When a step has no screenshots, the right column is absent: the text column expands to full width.

**One screenshot**: display it filling the right column width.

**Two or more screenshots**: render a carousel scoped to this step filling the right column:
- One image visible at a time at full column width
- Previous / Next arrow buttons positioned left and right of the image
- Position indicator below the image ("2 / 4" style, not dots)
- Implemented in vanilla JS with inline CSS: no external libraries
- Each carousel instance is fully independent (multiple steps on the same page operate independently)
- Arrow button click handlers must call `event.stopPropagation()` so they do not bubble to step-level or page-level navigation handlers
- Each image in the carousel also opens the lightbox on click

---

The remaining left-column elements, in order below the actions list:

4. **Concept in context** (render only when directly applicable: do not force one onto every step):
   Before rendering this element, ask: does executing this step correctly depend on understanding a specific concept from this task's `concepts[]`? A concept is directly applicable when misunderstanding it would cause the learner to execute the step incorrectly, choose the wrong value, or not understand what the completion signal means. If yes, render a small inline concept hook immediately below the actions list:
   - A subtle expandable chip: label "Why this works: [concept name]": styled in a lighter background with a left accent border, smaller than body type
   - Clicking/tapping expands to reveal the concept's explanation inline, without navigating away from the step. Collapse on second click.
   - Only the concept(s) that directly apply to this specific step: not every concept in the task
   - If no concept from `concepts[]` is directly load-bearing for this step, render nothing here

5. **Irreversible warning** (render only when the step is flagged irreversible):
   - A prominent caution banner in Ignis red or Suma orange tint, with an inline SVG warning icon: "This step cannot be undone. Review carefully before proceeding."

6. **Notes** (render only when step.notes is present):
   - A styled note/warning callout block (distinct background, left border accent)

7. **Completion signal**:
   - Always present; rendered last before the navigation controls
   - Label: "Done when:" in muted text
   - The completion text: displayed in a visually prominent pill or callout with a Viridis green tint background and an inline SVG checkmark icon
   - This must be clearly visible before the learner clicks Next: it is the gate check

8. **Step navigation controls** (always rendered at the bottom of each step panel):
   - Left button: "← Back": disabled and visually subdued on step 1
   - Centre label: "Step [n] of [N]"
   - Right button: "Next →": on all steps except the last; on the last step render "Mark task complete →" which triggers task completion and advances the page-level progress indicator
   - Transition between steps: a short CSS opacity fade (150ms ease)
   - Do not bind left/right arrow keys to step navigation: page-level ← → keys handle task page navigation and must not conflict with step navigation

**Lightbox**: implement a single shared lightbox overlay in vanilla JS: one instance for the entire page, reused for all step images. The lightbox must:

- Cover the full viewport with a dark semi-transparent backdrop (`rgba(0,0,0,0.85)`)
- Display the full-resolution image centred, scaled to fit the viewport maintaining native aspect ratio (`max-width: 90vw; max-height: 90vh; object-fit: contain`)
- Show a close button (×) in the top-right corner of the overlay
- Close on: close button click, backdrop click, or Escape key
- Trap focus while open; restore focus to the triggering image on close
- Animate open and close with a short CSS opacity/scale transition
- Require no external libraries

**h. Task footer (Page C only)**

- The final step's "Mark task complete →" button marks the task done, updates the progress indicator, and navigates to Page A of the next task. On the final task it advances to the closing page.

---

### 5. Closing page

This page closes the social contract opened on the cover. It must be as visually prominent and standalone as the cover page social contract statement.

- **Social contract statement: closing**: render as the dominant element of the page: **"You have just [workflow.objective]"**: same visual weight and treatment as the opening statement on the cover; a distinct background block or large-type callout; it must feel conclusive and affirming, not like a footnote
- Below it: a visual summary row of all task titles (numbered cards or a timeline): this is secondary to the closing statement, not equal to it
- "Return to start" link at the bottom
- Footer: workflow record_id | Version [version] | Generated by Blueprinted; include logo if present

---

## Real-world scenario content rules

These rules apply to the "In practice" scenario text generated for each task during Phase 1.

Before generating each scenario, perform a web search for recent real-world incidents, outages, data loss events, compliance requirements, or case studies relevant to this task's technology domain. Look for examples from the past 2–3 years that illustrate why a professional would need to perform exactly this kind of task. Specific, current details make scenarios credible: training data alone tends to produce generic examples. If web search is unavailable or returns nothing relevant, use training data but keep the scenario grounded and specific.

Write the scenario as narrative prose: 3 to 5 short paragraphs. This is storytelling, not a briefing document. It should read the way a case study or a well-written trade article opens: it puts a real person in a real place under real pressure. Do not write in the style of "This scenario involves a storage administrator who needs to...". Write it in third-person narrative: show the situation unfolding.

The tone is professional and grounded. Never describe how a character feels, their emotional state, or their internal reaction to events. Do not use literary or flowery language: no metaphors, no atmospheric description, no dramatic phrasing. Write what happened and what is at stake, not how it felt. If a sentence could belong in a novel, rewrite it as a sentence that could belong in a trade publication.

Each scenario must feature the same character established at the start of Phase 1. The scenario must:

- **Show how the character finds out**: the scenario must open from the character's point of view: specifically, the moment they become aware that this task needs doing. This is not a background summary of a situation; it is the character encountering the problem. Be specific about the trigger: an email arrives, a monitoring alert fires, a colleague stops by their desk, a ticket lands in the queue, they notice something anomalous themselves. The learner should understand not just what the problem is, but how it landed on this person. Do not open with context: open with the moment of discovery.
- **Build the pressure**: describe the event or circumstance that has made this specific task urgent right now: a recent incident, a compliance deadline, a failed audit, a growth event, a near-miss. Draw on web search findings to make this feel current and earned.
- **End with intent**: close with a sentence describing what the character is about to do and what success looks like: mapped directly to the task's `outcome` field. This should feel like the moment before action, not a summary.

---

## Navigation and interaction

**Page order** (every page is a distinct view: see Page model above):

1. Cover
2. Character introduction
3. Primer 1, Primer 2, … (one page each, if primers exist)
4. For each task: Page A (Introduction) → Page B (Concepts) → Page C (Procedure)
5. Closing page

**Sidebar**: shows the character intro, each primer by title, each task by title with a completion indicator (dot or checkmark), and the closing page at the bottom. The currently active page is highlighted. Task sub-pages (A/B/C) are not individually listed in the sidebar: the task entry in the sidebar links to Page A of that task.

**Linear navigation only: critical constraint**: every page has exactly one way to move forward: a single Next/Continue/Start button. Do not add any secondary navigation shortcuts, deep links, anchor tags, or buttons that jump to a non-adjacent page. This includes: but is not limited to: "read the primers first" links on the cover, clickable task titles that skip ahead, "jump to task N" shortcuts, and "skip intro" buttons. The learner moves through the course one page at a time in sequence. The sidebar shows progress and allows the learner to go back to previously visited pages, but must never provide a shortcut past unvisited content.

**Previous / Next buttons**: present on every page. Never absent. Always within the viewport without scrolling.

**Cover page**: one button only: "Start →": navigates to the character introduction page.

**Character introduction page**: one button only: "Continue →": navigates to Primer 1 (if primers exist) or Task 1 Page A.

**Keyboard navigation**: left arrow = previous page, right arrow = next page. Step-through navigation within a task uses button clicks only: do not bind arrow keys to step navigation.

**Progress indicator**: "Task N of N complete": updated when a learner completes Page C of a task (clicks "Mark task complete →").

**URL hash**: update the hash on every page transition so the browser back/forward buttons work (e.g. `#cover`, `#character`, `#primer-1`, `#task-1-intro`, `#task-1-concepts`, `#task-1-procedure`, `#complete`).

**Scroll**: scroll to the top of the content area on every page transition.

---

## Design requirements

### Interactive components
All interactive components (flip cards, accordions, step-through panels, expand/collapse toggles) must:
- Be implemented entirely in vanilla JS with embedded CSS
- Work correctly when multiple instances of the same component type appear on the same page (scoped state: no shared global variables per component type)
- Be keyboard accessible: Enter/Space to activate, Escape to collapse where applicable
- Animate smoothly (CSS transitions, not abrupt show/hide)
- Degrade gracefully: if JS is disabled, content must still be readable (use `<details>`/`<summary>` as a fallback pattern where appropriate)

#### Flip card implementation: mandatory pattern

Do not implement flip cards using CSS 3D transforms (`transform-style: preserve-3d`, `backface-visibility`, `rotateY`). Browser support for the 3D rendering context is inconsistent, vendor-prefix requirements vary, and fixed-height containers clip back-face content when it is taller than the front face.

Use the **CSS grid stack + crossfade** pattern instead:

- The card has two children: `.front` and `.back`
- Both children are given `grid-area: 1 / 1`, placing them in the same grid cell while remaining in normal document flow: the container naturally sizes to whichever face is taller, no fixed heights, no clipping
- Default state: `.front` has `opacity: 1`, `.back` has `opacity: 0; pointer-events: none`
- Flipped state (toggled by JS adding a `.flipped` class to the card): `.front` has `opacity: 0; pointer-events: none; transform: scaleX(-1)`, `.back` has `opacity: 1; pointer-events: auto`
- Apply `transition: opacity 200ms ease, transform 200ms ease` to both faces
- The mirror-scale on the front gives the visual impression of a card turning over without requiring a 3D context

#### Page layout: mandatory pattern

Every `.page` element (the discrete page views defined in the Page model) must use:

```css
.page { display: none; flex-direction: column; min-height: 100vh; }
.page.active { display: flex; }
.page .content-area { flex: 1; }
```

This ensures the page-nav bar and page footer are always pushed to the bottom of the viewport regardless of content height. Never use `display: block` on a `.page` element: a short-content page will leave the footer floating mid-screen. The cover page's mid-section content (task card row, progress bar) must be wrapped in a `.content-area` div so the same flex rule applies.

### Visual style
Apply the colour palette, typography scale, button specifications, and spacing
values defined in `styles.md` exactly. Do not introduce colours, font sizes,
or spacing values not listed there.

Key constraints specific to this HTML output format:

- All CSS and JS embedded inline in index.html. No external CDN dependencies,
  including fonts. Use the CSS font stack defined in `styles.md` with system
  font fallbacks only. Images and logos are referenced as relative paths —
  never base64-encoded: so they resolve correctly when the ZIP is extracted.
- Content column: max-width 800px, centred.
- Line-height: 1.7 for body copy.
- Fully responsive at all viewport widths. This is non-negotiable. Every element
  in the output: sidebar, content column, card grids, flip cards, accordions,
  procedure tables, carousels, navigation buttons, progress indicators: must
  reflow correctly from 320px (small phone) to 1440px (desktop). Use CSS
  flexbox or grid with wrapping for all multi-column layouts; never use fixed
  pixel widths on containers. Test every interactive component at mobile width:
  cards must stack to a single column, tables must scroll horizontally or
  reflow to a stacked layout, buttons must be tap-target sized (minimum 44px
  height). The sidebar collapses to a hamburger menu on narrow screens. Do not
  leave any layout that breaks, overflows, or becomes unusable below 480px.
- Print-friendly: @media print hides all navigation and renders content
  sequentially without page breaks inside procedure tables.
- Apply spacing values from the vertical spacing scale in `styles.md` for all
  margin and padding. Do not interpolate between scale values.
- Add `li em, td em { white-space: nowrap; }` to the global stylesheet. This
  prevents italic product name fragments (e.g. the *for Microsoft Windows*
  portion of a Veeam product name) from wrapping independently of the
  surrounding text inside list items or table cells.

### Semantic colour use
Use the colour roles defined in `styles.md` consistently:

- Completion signal table cells: Viridis green tint background
- Warning / irreversible banners: Suma orange or Ignis red tint
- Facts callout cards: visually distinct from Concepts prose sections
- Primer cards: visually distinct from task content pages
- Interactive elements (buttons, links, progress indicators): use the button
  colour variants from `styles.md` appropriate to the background context

---

## Writing and tone

Apply all rules in `editorial.md`. Pay particular attention to the LEARNING
OVERRIDE section, which supersedes the marketing tone of the original Veeam
editorial guide.

Rules specific to this HTML output that are not covered in `editorial.md`:

- Do not use em dashes (—) anywhere: not in headings, framing text, generated copy, or any other output. Covered in editorial.md. Hard rule, no exceptions
- Do not editorialize step text: render step.text exactly as authored in the
  data; do not paraphrase, improve, or reframe it
- Section headings must follow the exact wording defined in the Page structure
  section of this prompt; do not substitute synonyms
- Product names that contain `<em>` tags (e.g. Veeam Agent <em>for Microsoft
  Windows</em>) must never be split across flex children or broken by a line
  wrap. Wherever a product name appears inside a list item, table cell, or any
  flex or grid container, wrap the entire product name in
  `<span style="white-space: nowrap">` so the italicised portion cannot be
  separated from the preceding words by a line break or flex reflow

---

## Footer (every page)

Display: Workflow record_id | Version [version] | Generated by Blueprinted

---

## Data reference

    workflow.title
    workflow.objective
    workflow.record_id
    workflow.version
    workflow.tasks[]
    workflow.primers[]          // optional

    tasks[].title
    tasks[].outcome
    tasks[].facts[]
    tasks[].concepts[]
    tasks[].procedure_name
    tasks[].steps[]
    tasks[].dependencies[]
    tasks[].task_assets[]       // optional; objects with url, type, label
    tasks[].irreversible        // boolean, optional

    steps[].text
    steps[].actions[]           // optional
    steps[].completion
    steps[].notes               // optional
    steps[].screenshots[]       // optional; filenames only
""",

    "helpsheet": """\
# Export Package: Helpsheet

You are generating a helpsheet: a practical just-in-time reference document intended for learners who have already completed training and need a quick process reminder at the point of work. It is not a teaching document. Do not explain concepts, provide background, or include introductory material. The learner knows what the task is; they need to remember how to do it.

This ZIP contains:
- `data.json`: full structured workflow and task content
- `images/`: screenshots (include only the most essential ones, max 1 per task)
- `logos/`: client logo file (e.g. `client-logo.png`); use it if present
- `styles.md`: visual design specification (colours, typography, spacing, buttons)
- `editorial.md`: writing rules, product name spelling, and tone guidance

Read `styles.md` and `editorial.md` before generating output. They are authoritative. Where any instruction in this prompt conflicts with those files, the guidance in those files takes precedence.

If a logo file is present in `logos/`, place it in the header (top-left or top-right, maintaining native aspect ratio) and in the footer alongside the workflow record_id and version. If no logo is present, leave those positions empty.

---

## Understanding the data model

Before generating output, understand what each entity is.

### Workflow

A Workflow is a composite outcome made of ordered Tasks. It has:

- `title`: learner-facing name
- `objective`: the organisation-defined outcome the workflow produces
- `tasks[]`: ordered list of Task references

### Task

A Task is an atomic, self-contained unit of performance. It produces exactly one observable outcome. A Task has:

- `title`: verb-driven learner-facing name
- `outcome`: the observable result when the task is complete
- `facts[]`: background knowledge (do not render on the helpsheet; excluded by design)
- `concepts[]`: mental models (do not render on the helpsheet; excluded by design)
- `procedure_name`: the name of the step sequence
- `steps[]`: ordered atomic instructions
- `dependencies[]`: conditions or access requirements (render only if critical to execution)
- `task_assets[]`: optional media objects (see Media below)
- `irreversible`: boolean; true means the task cannot be undone; always flag this

### Step

Each Step inside a Task has:

- `text`: (required) the primary action in imperative form
- `actions[]`: (optional) sub-steps; condense to key CLI commands, menu paths, or keyboard shortcuts only
- `completion`: (required) the observable signal confirming the step is complete
- `notes`: (optional) warnings only; include if safety-critical; omit general context
- `screenshots[]`: (optional) image filenames; include at most one per task if it aids orientation

### Media (task_assets)

Each task asset object has:

- `url`: fully-qualified external URL
- `type`: one of: `video`, `demo`, `image`, `audio`, `module`, `link`
- `label`: short descriptive label

On a helpsheet, render media as a single labelled reference line at the end of the task section only if the asset is directly needed at the point of work (e.g. a link to a required tool or demo environment). Do not include videos or conceptual reading links.

---

## Your output

Produce a single-page (A4) quick-reference helpsheet in Markdown. It must be dense but scannable. Prioritise process over explanation at every decision point.

### Structure
- **Header**: workflow title, one-sentence objective, software name/version if present
- **For each task**: a compact section:
  - Task title (bold heading)
  - If `irreversible` is true: a brief caution note before the steps
  - Steps: numbered, condensed to the essential action only (omit explanatory text; keep the verb and the object)
  - Key commands/actions: inline code blocks for any CLI commands, menu paths, or keyboard shortcuts from the `actions` field
  - One screenshot if present and critical to understanding; render at native aspect ratio: never stretch or distort; scale down proportionally if wider than the column, never set height independently of width
  - "Done when:": one-line completion signal taken from the final step's `completion` field
  - Media reference line if applicable (see Media above)

- **Footer**: workflow record_id and version, export date

---

## Style rules
- Maximum two A4 pages when printed at 10pt
- No preamble, no introductory paragraphs, no concepts or facts sections
- Terse imperative sentences only ("Click Save", "Run `sudo systemctl restart`")
- Do not use em dashes (—) anywhere. Hard rule, no exceptions
- Use a two-column layout if there are more than 4 tasks (indicate this with a note at the top of the file)
- Bold the first word of each step (the verb)

---

## Data reference

    workflow.title
    workflow.objective
    workflow.record_id
    workflow.version
    workflow.tasks[]

    tasks[].title
    tasks[].procedure_name
    tasks[].steps[]
    tasks[].dependencies[]     // render only if critical
    tasks[].task_assets[]      // optional; reference-only on helpsheet
    tasks[].irreversible       // boolean; always flag if true

    steps[].text
    steps[].actions[]          // condense to key commands only
    steps[].completion         // "Done when:" line
    steps[].notes              // warnings only
    steps[].screenshots[]      // max 1 per task
""",
}


_BRAND_STYLE_GUIDE = """\
# Veeam Brand Style Guide
## For use by LLM-generated HTML outputs

This file defines the visual and typographic standards for all Veeam-branded
HTML content generated from Blueprinted workflow exports. Apply these rules
exactly. Do not introduce colours, fonts, or spacing values not listed here.

---

## Logotypes

Use the primary Veeam logo on white or light backgrounds.
Use the secondary logo when the rendered logo width is under 200px.
Variant logos exist for: Grey Mineral background, Viridis (green) background,
Black background.

Do not modify logo proportions, colours, or surrounding elements.

The Bounce Forward Mark (the green chevron/tick mark) may be used as a
standalone decorative element in marketing contexts at reduced opacity:
- Marketing use: 25% opacity
- GUI use: 20% opacity

---

## Colours

### Primary Brand Colours

| Name         | Hex       | Usage                                      |
|--------------|-----------|--------------------------------------------|
| Viridis      | #00D15F   | Primary brand green; CTAs, highlights      |
| Black        | #000000   | Primary text on light backgrounds          |
| White        | #FFFFFF   | Primary text on dark backgrounds; surfaces |

### Green Palette (light to dark)

| Hex       | Notes                        |
|-----------|------------------------------|
| #E1F4EC   | Lightest tint                |
| #9CFFA3   | Lime                         |
| #32F26F   | Mint                         |
| #00D15F   | Viridis (primary)            |
| #009277   | Hover state for primary CTAs |
| #007F49   | Pine                         |
| #02613F   | Darkest green                |

### Blue Palette

| Hex       | Notes                                  |
|-----------|----------------------------------------|
| #EEF4F6   | Lightest blue tint; tertiary hover bg  |
| #E3EEFE   | Light blue tint                        |
| #57E0FF   | Sky                                    |
| #3700FF   | Electric Azure; default button colour  |
| #283E8E   | Button hover state (primary/secondary) |
| #002833   | Darkest blue                           |

### Supplementary

| Hex       | Name   |
|-----------|--------|
| #8E71F4   | Casia (purple) |
| #CECBB8   | Warm grey      |

### System / Status Colours

| Hex       | Name   | Usage              |
|-----------|--------|--------------------|
| #ED2B3D   | Ignis  | Error, danger      |
| #FE8A25   | Suma   | Warning, orange    |
| #FFD839   | Sol    | Caution, yellow    |
| #FCF8EB   |:      | Warning background |

### Neutral Family

| Hex       | Name          |
|-----------|---------------|
| #000000   | Black         |
| #232323   | Near black    |
| #505861   | Dark Mineral  |
| #ADACAF   | French Grey   |
| #F0F0F0   | Fog           |
| #F9F9F9   | Off white     |
| #FFFFFF   | White         |

### Gradients

**Green gradients**: use for hero backgrounds and feature sections:
- Primary dark 135deg
- Primary light 135deg
- Secondary 180deg / 135deg
- Tertiary 180deg / 90deg
- Tertiary light 90deg
- Tertiary green 135deg

**Blue gradients:**
- Azure Blue to Sky
- Dark Grey to Dark Mineral
- Azure Blue to Casia

**White gradients:**
- White to Neutral
- White to Sky Pale
- Super white 180deg

---

## Typography

### Font Stack

**Primary (web and print):** ES Build Variable
- Request from brand@veeam.com if not available; do not substitute without fallback

**Web fallback / Google Fonts:** Source Sans Pro
- Source Sans Pro Light
- Source Sans Pro Regular
- Source Sans Pro Bold

**MS Office / plain text fallback:** Tahoma Regular, Tahoma Bold

**CSS font stack for generated HTML:**

    font-family: 'ES Build', 'Source Sans Pro', 'Tahoma', sans-serif;

### Type Scale

All sizes in pixels. Format: font-size / line-height.

| Style       | Size / Line-height | Notes                        |
|-------------|--------------------|------------------------------|
| Display 100 | 100 /:            | Hero display only            |
| Display 60  | 60 /:             | Large display                |
| H1          | 50 / 60            |                              |
| H2          | 55 / 52            |                              |
| H3          | 36 / 44            |                              |
| Text 28     | 28 / 36            |                              |
| Text 24     | 24 / 28            |                              |
| Text 20     | 20 / 24            |                              |
| Text 18     | 18 / 24            |                              |
| Paragraph   | 16 / 24            | Body copy default            |
| Caption     | 14 / 20            | Labels, footnotes            |
| Eyebrow     |: / 28 at size 24  | Section labels above headings|

### Paragraph Styles

| Style       | Size / Line-height | Weight  |
|-------------|--------------------|---------|
| Body        | 18 / 24            | Regular |
| Body bold   | 18 / 24            | Bold    |
| Body        | 16 / 24            | Regular |
| Body bold   | 16 / 24            | Bold    |
| Caption     | 14 / 20            | Regular |

---

## Buttons

### Corner radius
All buttons: 6px border-radius.

### Colour variants

**On dark backgrounds (default brand context):**

| Variant   | Default    | Hover      | Notes                         |
|-----------|------------|------------|-------------------------------|
| Primary   | #3700FF    | #283E8E    | Electric Azure                |
| Secondary | #3700FF    | #283E8E    |                               |
| Tertiary  | transparent| #EEF4F6    | Ghost/outline style           |

**On white backgrounds:**

| Variant   | Default    | Hover      |
|-----------|------------|------------|
| Primary   | #FFFFFF    | #232323    |
| Secondary | #FFFFFF    | #232323    |
| Tertiary  | transparent| #FFFFFF    |

**Green variant:**

| Variant   | Default    | Hover      |
|-----------|------------|------------|
| Primary   | #00D15F    | #009277    |
| Secondary | #00D15F    | #009277    |
| Tertiary  | transparent| #EEF4F6    |

### Button sizing

**Large CTA (default behaviour):**
- Padding: 12px top/bottom, 24px left/right
- Text: centred

**Small CTA:**
- Padding: 8px top/bottom, 20px left/right
- Text: centred

**Fixed width:**
- Large: 285px wide
- Small: 285px wide

**Mobile (full width):**
- Large and small: 100% width

**Minimum widths:**
- Large: 190px
- Small: 66px (rendered as 166px in spec)

---

## Grid System

### Breakpoints

| Label | Min-width | Layout     | Content padding |
|-------|-----------|------------|-----------------|
| L     | 1260px    | Centred    | 0 15px          |
| M     | 1024px    | Stretched  | 30px            |
| S     | 768px     | Stretched  | 30px            |
| xS    | < 767px   | Stretched  | 30px            |

**L breakpoint:** content area max-width 1260px, centred, 15px horizontal padding.

### Vertical spacing scale (pixels)

8, 16, 24, 32, 48, 64, 80, 120

Use only these values for margin and padding. Do not interpolate.

---

## Icons and Illustrations

- **Line icons:** Veeam line icons library (access via brand portal)
- **Illustrations:** Veeam new illustrations library (access via brand portal)
- Icon gradient alignment: use the provided gradient direction guide; do not
  apply arbitrary gradients to icons
- Marketing illustrations use an isometric approach with unified component
  libraries

---

## Things to Avoid

- Do not use 3D styling on book covers or document covers: use 2D only
- Container values with images or icons must be even numbers divisible by two
- Do not introduce colours outside the defined palette
- Do not use ES Build substitutes without falling back to Source Sans Pro

---

## Notes for LLM-generated HTML

When generating HTML from this style guide:

1. Use the CSS font stack defined above. Do not load external font files unless
   ES Build or Source Sans Pro are confirmed available in the delivery environment.
2. Apply the Viridis green (#00D15F) as the primary interactive accent.
3. Use Electric Azure (#3700FF) for primary CTA buttons on dark backgrounds.
4. Apply the neutral family for backgrounds: #000000 or #232323 for dark
   surfaces, #F0F0F0 or #F9F9F9 for light surfaces.
5. Use Ignis (#ED2B3D) for error/danger states and Suma (#FE8A25) for warnings.
6. Respect the vertical spacing scale strictly: 8px increments only.
7. All buttons must have 6px border-radius.
8. Body copy default: 16px / 24px line-height, Source Sans Pro Regular.
9. Do not introduce drop shadows, gradients, or decorative elements not
   specified in this guide.
"""


_EDITORIAL_STYLE_GUIDE = """\
# Veeam Editorial Style Guide: Learning Content Edition
## Extracted and adapted for Blueprinted HTML course output

This file is derived from the Veeam Editorial Style Guide 2024. It has been
filtered and annotated for use in self-paced learning content. Not all rules
from the original guide apply here.

Where the original guide produces marketing or technical language that is
hostile to new learners, this document overrides it. Those overrides are
marked clearly.

---

## Language

### English variant
Use American English throughout.
- virtualize (not virtualise)
- center (not centre)
- backup (not back-up)

### Tone: LEARNING OVERRIDE
The original style guide is written for marketing collateral and partner
communications. That tone is inappropriate for learning content.

Do not use:
- Superlatives and brand claims ("industry-leading," "#1," "radically resilient")
- Benefit-forward language that describes outcomes before explaining what
  something is ("Bounce forward with confidence" means nothing to a new learner)
- Passive constructions that obscure the action ("data can be protected" vs
  "back up your data")
- Acronym-first introductions without plain-language context

Do use:
- Second-person, present tense: "Select the dropdown. The list expands."
- Plain descriptions before product names: explain what the thing does before
  naming it
- Concrete actions with observable results
- The simplest word that is accurate

If you find yourself writing a sentence that could appear on a product
brochure, rewrite it.

---

## Capitalisation rules (apply these)

### Sentence case
Use sentence case (first word and proper nouns only) for:
- Section headers and sub-headers within learning content
- Step text and action text
- Body copy, bullet points, callout boxes
- Knowledge check questions and answer options
- Outcome statements and completion signals

### Title case
Use title case (all major words capitalised) only for:
- The workflow title (cover page heading)
- Task titles
- Primer titles
- Page-level navigation labels

### Product names
All Veeam product names are title-cased and must be written exactly as
specified in the Product Names section below. Never abbreviate on first use.

---

## Numbers

- Spell out numbers one through nine in body copy
- Use numerals for 10 and above
- Use numerals for all units of measurement: 2 CPUs, 10GB, 5-minute RTO
- No space between a numeral and its unit: 10GB (not 10 GB)
- Do not begin a sentence with a numeral; restructure the sentence

---

## Punctuation

### Oxford comma
Always use the Oxford (serial) comma in lists of three or more items.
- Correct: facts, concepts, and steps
- Incorrect: facts, concepts and steps

### Em dash
Do not use em dashes (—) anywhere. Not in headings, body copy, step text, callouts, or any other output. Replace with a comma, colon, or rewrite the sentence. This is a hard rule with no exceptions.

### En dash
Use to express ranges without spaces: 9–10 minutes, steps 3–5.

### Hyphen
Use for compound modifiers before a noun: cloud-native environment,
step-by-step procedure. No space around a hyphen.

### Ampersand
Do not use & in body copy or step text. Spell out "and."
Exception: Veeam Backup & Replication: the ampersand is part of the
official product name and must be retained.

### Symbols
- Use % not "percent"
- Use # only in the specific phrase "#1" when that claim is editorially
  required. Do not use in learning content.
- Do not use trademark symbols (®, ™) in learning materials

---

## Dates and times

- Date format: Jan. 12, 2023 (abbreviate all months except March, April, May,
  June, July when paired with a specific date)
- Time format: 11 a.m., 2:30 p.m. (12-hour clock; no zeroes on the hour)
- Do not reference time zones unless the content is event-specific

---

## Product names and spelling

Write all product names exactly as shown. Do not abbreviate on first use.
Do not add trademark or registered symbols.

### Platform products

- Veeam Data Platform (umbrella term)
  - Veeam Data Platform Premium
  - Veeam Data Platform Advanced
  - Veeam Data Platform Foundation
  - Veeam Data Platform Essentials (small business; max 50 workloads)
- Veeam Data Cloud
  - Veeam Data Cloud for Microsoft 365
  - Veeam Data Cloud for Microsoft Azure
  - Veeam Data Cloud Vault (second reference: Veeam Vault)

### Core products

- Veeam Backup & Replication (the ampersand is required)
- Veeam ONE
- Veeam Recovery Orchestrator
- Veeam Service Provider Console
- Veeam Kasten (first reference); Kasten (second reference)
  - Do not use Kasten K10 or K10

### Backup products (the "for" portion is always italicised)

- Veeam Backup *for Microsoft 365*
- Veeam Backup *for Salesforce*
- Veeam Backup *for AWS*
- Veeam Backup *for Microsoft Azure*
- Veeam Backup *for Google Cloud*
- Veeam Backup *for Nutanix AHV*
- Veeam Backup *for Red Hat Virtualization*

### Agent products (the "for" portion is always italicised)

- Veeam Agent *for Microsoft Windows*
- Veeam Agent *for Linux*
- Veeam Agent *for Mac*
- Veeam Agent *for IBM AIX*
- Veeam Agent *for Oracle Solaris*

When listing multiple agents together, conjunctions ("and") are not italicised:
Veeam Agents *for Microsoft Windows*, *Mac*, *Linux*, *IBM AIX* and *Oracle Solaris*

### Explorer products (the "for" portion is always italicised)

- Veeam Explorer *for Microsoft Active Directory*
- Veeam Explorer *for Microsoft Exchange*
- Veeam Explorer *for Microsoft SharePoint*
- Veeam Explorer *for Microsoft SQL Server*
- Veeam Explorer *for Microsoft Teams*
- Veeam Explorer *for Oracle*
- Veeam Explorer *for PostgreSQL*
- Veeam Explorer for Storage Snapshots (no italics: multi-partner exception)

### Plug-ins (the "for" portion is always italicised)

- Veeam Plug-in *for Oracle RMAN*
- Veeam Plug-in *for SAP HANA*

### Community and free editions

- Veeam Backup & Replication Community Edition
- Veeam ONE Community Edition
- Veeam Agent *for Windows* Free
- Veeam Agent *for Linux* Free

### Version numbers

When the product name is present: lowercase v, no trailing .0 for major versions.
- Veeam Backup & Replication v12
- Veeam Backup & Replication v12a

When the product name is absent: uppercase V.
- V12 introduces new recovery options.
- This feature was added in V12A.

Only products receive version numbers. Features do not.

---

## Common word usage

| Word | Correct form | Notes |
|------|-------------|-------|
| backup | backup (noun/adjective) | "Create a backup"; "backup job" |
| back up | back up (verb) | "Back up your data" |
| resilience | resilience | Never "resiliency" |
| resilient | resilient | Adjective form of resilience |
| cyberattack | cyberattack | One word, no hyphen |
| cybersecurity | cybersecurity | One word, no hyphen |
| cyber resilience | cyber resilience | Two words |
| cyberthreat | cyberthreat | One word, no hyphen |
| log-in | log-in (noun/adjective) | "Enter your log-in details" |
| log in | log in (verb) | "Log in to the console" |
| sign-in | sign-in (noun/adjective) | "Use your sign-in credentials" |
| sign in | sign in (verb) | "Sign in to continue" |

---

## Hypervisors

First reference: VMware vSphere, Microsoft Hyper-V (full names).
Subsequent references: vSphere, Hyper-V.
Do not abbreviate when both appear in the same document.

---

## Alliance partner names (commonly referenced in Veeam content)

Use full name on first reference, abbreviation thereafter where noted.

| Partner | First reference | Subsequent |
|---------|----------------|------------|
| Amazon Web Services | Amazon Web Services (AWS) | AWS |
| Hewlett Packard Enterprise | Hewlett Packard Enterprise (HPE) | HPE |
| VMware | VMware vSphere | vSphere |
| Google Cloud | Google Cloud | Google Cloud |
| Microsoft Azure | Microsoft Azure | Azure |

Never use "Google Cloud Platform." Use "Google Cloud" only.

---

## Learning content: additional rules

These rules apply specifically to Blueprinted-generated HTML output and are
not in the original Veeam style guide.

**Step text:** Write in imperative form. One action per step.
- Correct: "Select Save from the File menu."
- Incorrect: "The user should navigate to the File menu and then save."

**Completion signals:** State the observable result, not a feeling.
- Correct: "The backup job status shows Completed."
- Incorrect: "You will know the backup is done."

**Facts:** State as discrete, verifiable sentences. Do not editorialize.
- Correct: "Veeam Backup & Replication stores restore points as incremental
  backup files after the first full backup."
- Incorrect: "Veeam's industry-leading backup technology creates restore
  points so you can recover with confidence."

**Concepts:** Explain the mechanism in plain language. Avoid brand framing.
- Correct: "A restore point represents the state of a VM at a specific moment
  in time. Each incremental backup captures only the blocks that changed since
  the last backup, which reduces storage usage and backup duration."
- Incorrect: "Veeam's revolutionary instant recovery technology leverages
  restore points to deliver RTOs of minutes."

**Product claims in learning content:** Omit entirely. The learner is already
using the product. Claims are marketing, not instruction.

**Acronyms:** Define on first use, then use the acronym.
- "Recovery Time Objective (RTO)" on first use, "RTO" thereafter.
"""


def _build_package_data(conn, record_id: str, version: int) -> dict[str, Any]:
    """Build the full data.json payload for an export package.

    Image URLs are rewritten from /task-images/{rid}/{file} to images/{rid}/{file}
    so they resolve as relative paths within the ZIP.
    """
    wf = conn.execute(
        "SELECT * FROM workflows WHERE record_id=? AND version=?", (record_id, version)
    ).fetchone()
    if not wf:
        raise HTTPException(404)

    ref_rows = conn.execute(
        "SELECT order_index, task_record_id, task_version FROM workflow_task_refs "
        "WHERE workflow_record_id=? AND workflow_version=? ORDER BY order_index",
        (record_id, version),
    ).fetchall()

    primer_rows = conn.execute(_PRIMER_ATTACH_QUERY, (record_id,)).fetchall()

    def _rewrite_url(url: str) -> str:
        if url and url.startswith("/task-images/"):
            return "images/" + url[len("/task-images/"):]
        return url

    tasks = []
    for r in ref_rows:
        t = conn.execute(
            "SELECT * FROM tasks WHERE record_id=? AND version=?",
            (r["task_record_id"], int(r["task_version"])),
        ).fetchone()
        if not t:
            raise HTTPException(409, detail=f"Referenced task {r['task_record_id']} not found")
        steps = _normalize_steps(_json_load(t["steps_json"]) or [])
        for step in steps:
            step["screenshots"] = [_rewrite_url(u) for u in (step.get("screenshots") or [])]
        assets = [
            {**a, "url": _rewrite_url(a.get("url", ""))}
            for a in (_json_load(t["task_assets_json"]) or [])
        ]
        tasks.append({
            "order_index": int(r["order_index"]),
            "record_id": t["record_id"],
            "version": int(t["version"]),
            "title": t["title"],
            "outcome": t["outcome"],
            "facts": _json_load(t["facts_json"]) or [],
            "concepts": _json_load(t["concepts_json"]) or [],
            "dependencies": _json_load(t["dependencies_json"]) or [],
            "procedure_name": t["procedure_name"],
            "irreversible": bool(t["irreversible_flag"]),
            "steps": steps,
            "task_assets": assets,
            "media_url": (t["media_url"] if "media_url" in t.keys() else None),
            "software_name": (t["software_name"] if "software_name" in t.keys() else None),
            "software_version": (t["software_version"] if "software_version" in t.keys() else None),
        })

    return {
        "workflow": {
            "record_id": wf["record_id"],
            "version": int(wf["version"]),
            "title": wf["title"],
            "objective": wf["objective"],
            "domains": _normalize_domains(wf["domains_json"]),
        },
        "tasks": tasks,
        "primers": [
            {
                "title": p["title"],
                "summary": p["summary"],
                "explanation": p["explanation"],
                "analogies": p["analogies"],
            }
            for p in primer_rows
        ],
    }


def _collect_package_images(data: dict[str, Any]) -> list[tuple[str, str]]:
    """Return [(zip_path, abs_disk_path), ...] for all images referenced in the package data.

    Skips files that don't exist on disk.
    """
    seen: set[str] = set()
    pairs: list[tuple[str, str]] = []
    for task in data.get("tasks", []):
        for step in task.get("steps", []):
            for url in (step.get("screenshots") or []):
                if url and url.startswith("images/") and url not in seen:
                    seen.add(url)
                    rel = url[len("images/"):]  # {task_record_id}/{filename}
                    abs_path = os.path.join(TASK_IMAGES_DIR, rel)
                    if os.path.isfile(abs_path):
                        pairs.append((url, abs_path))
                    else:
                        logger.warning("Export package: image not found on disk: %s", abs_path)
        for asset in task.get("task_assets", []):
            url = asset.get("url", "")
            if url and url.startswith("images/") and url not in seen:
                seen.add(url)
                rel = url[len("images/"):]
                abs_path = os.path.join(TASK_IMAGES_DIR, rel)
                if os.path.isfile(abs_path):
                    pairs.append((url, abs_path))
                else:
                    logger.warning("Export package: image not found on disk: %s", abs_path)
    return pairs


@router.post("/workflows/{record_id}/{version}/export-package")
def workflow_export_package(
    request: Request,
    record_id: str,
    version: int,
    export_format: str = Form(...),
    logo_path: str | None = None,
):
    """Generate a ZIP export package for a confirmed workflow.

    The ZIP contains instructions.md (LLM prompt), data.json, and all referenced images.
    """
    require(request.state.role, "workflow:export")
    actor = request.state.user

    if export_format not in _EXPORT_FORMATS:
        raise HTTPException(400, detail=f"Unknown export format '{export_format}'")

    with db() as conn:
        wf = conn.execute(
            "SELECT * FROM workflows WHERE record_id=? AND version=?", (record_id, version)
        ).fetchone()
        if not wf:
            raise HTTPException(404)
        if wf["status"] != "confirmed":
            raise HTTPException(409, detail="Export is allowed for confirmed workflows only")

        ref_rows = conn.execute(
            "SELECT order_index, task_record_id, task_version FROM workflow_task_refs "
            "WHERE workflow_record_id=? AND workflow_version=? ORDER BY order_index",
            (record_id, version),
        ).fetchall()
        readiness = workflow_readiness(conn, [(r["task_record_id"], int(r["task_version"])) for r in ref_rows])
        if readiness != "ready":
            raise HTTPException(409, detail="Export is allowed only when all referenced task versions are confirmed")

        data = _build_package_data(conn, record_id, version)

    images = _collect_package_images(data)

    # Build ZIP in memory
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("instructions.md", _FORMAT_PROMPTS[export_format])
        zf.writestr("style.md", _BRAND_STYLE_GUIDE)
        zf.writestr("editorial.md", _EDITORIAL_STYLE_GUIDE)
        zf.writestr("data.json", json.dumps(data, ensure_ascii=False, indent=2))
        if logo_path and os.path.isfile(logo_path):
            logo_ext = Path(logo_path).suffix.lower()
            zf.write(logo_path, f"logos/client-logo{logo_ext}")
        else:
            zf.writestr("logos/place-logo-here.txt", "Place the client logo file in this folder before uploading to the LLM.\n")
        for zip_path, abs_path in images:
            zf.write(abs_path, zip_path)
    zip_bytes = buf.getvalue()

    # Save to EXPORTS_DIR and record artifact
    short = _short_code("WF", record_id)
    ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    filename = f"workflow__{short}__v{version}__{export_format}__{ts}.zip"
    out_path = os.path.join(EXPORTS_DIR, filename)
    os.makedirs(EXPORTS_DIR, exist_ok=True)
    with open(out_path, "wb") as f:
        f.write(zip_bytes)

    sha = _sha256_bytes(zip_bytes)
    task_refs_json = json.dumps([
        {"record_id": r["task_record_id"], "version": int(r["task_version"])} for r in ref_rows
    ])
    with db() as conn:
        conn.execute(
            """INSERT INTO export_artifacts(id, kind, filename, path, sha256,
               workflow_record_id, workflow_version, task_refs_json,
               exported_at, exported_by, retention_days)
               VALUES (?,?,?,?,?,?,?,?,?,?,?)""",
            (str(uuid.uuid4()), "zip_package", filename, out_path, sha,
             record_id, version, task_refs_json,
             utc_now_iso(), actor, 30),
        )
    audit("workflow", record_id, version, "export", actor, note=f"kind=zip_package format={export_format}")

    logger.info("Export package generated: %s (%d images)", filename, len(images))
    return FileResponse(
        path=out_path,
        media_type="application/zip",
        filename=filename,
    )


@router.get("/api/present/{token_id}")
def present_fetch(token_id: str):
    """Public endpoint: authenticated by the one-time token, no session required."""
    with db() as conn:
        tok = conn.execute(
            "SELECT * FROM presentation_tokens WHERE id=?", (token_id,)
        ).fetchone()
        if not tok:
            raise HTTPException(status_code=404, detail="Token not found")

        now_iso = utc_now_iso()
        now_dt = datetime.fromisoformat(now_iso.replace("Z", "+00:00"))
        exp_dt = datetime.fromisoformat(tok["expires_at"].replace("Z", "+00:00"))
        if now_dt > exp_dt:
            raise HTTPException(status_code=410, detail="Token expired")
        if tok["consumed_at"]:
            raise HTTPException(status_code=410, detail="Token already used")

        payload = _build_presentation_payload(conn, tok["workflow_record_id"], tok["workflow_version"])

        conn.execute(
            "UPDATE presentation_tokens SET consumed_at=? WHERE id=?",
            (now_iso, token_id),
        )
        audit("presentation_token", token_id, 1, "consume", "anon", conn=conn)

    return {**payload, "meta": {"generated_at": now_iso, "expires_at": tok["expires_at"]}}
